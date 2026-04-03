"""
Generate competitive analysis .docx: Instantly.ai vs NeuraLeads (ra.partnerwithus.tech)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style Definitions ───────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)  # Dark navy

h1 = doc.styles['Heading 1']
h1.font.size = Pt(22)
h1.font.bold = True
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(8)

h2 = doc.styles['Heading 2']
h2.font.size = Pt(16)
h2.font.bold = True
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(6)

h3 = doc.styles['Heading 3']
h3.font.size = Pt(13)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x2D, 0x5F, 0x9A)
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(4)

# ── Helper Functions ────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(headers, rows, col_widths=None, header_color="1A1A2E"):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.name = 'Calibri'
            # Alternate row shading
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F5F5FA")

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table

def add_status_table(headers, rows, status_col=2):
    """Add a table with color-coded status column."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
        set_cell_shading(cell, "1A1A2E")

    # Data
    STATUS_COLORS = {
        "FULL": "D4EDDA",       # green
        "PARTIAL": "FFF3CD",    # yellow
        "MISSING": "F8D7DA",    # red
        "SUPERIOR": "CCE5FF",   # blue
    }
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
            if c_idx == status_col:
                status_key = str(val).strip().upper().split(" ")[0]
                if status_key in ("FULL", "FULLY"):
                    set_cell_shading(cell, STATUS_COLORS["FULL"])
                elif status_key in ("PARTIAL", "PARTIALLY"):
                    set_cell_shading(cell, STATUS_COLORS["PARTIAL"])
                elif status_key in ("MISSING", "NOT"):
                    set_cell_shading(cell, STATUS_COLORS["MISSING"])
                elif status_key in ("SUPERIOR", "BETTER", "EXCEEDS"):
                    set_cell_shading(cell, STATUS_COLORS["SUPERIOR"])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8F9FA")

    return table

def add_bullet(text, level=0, bold_prefix=""):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(f" {text}").font.size = Pt(10)
    else:
        p.runs[0].font.size = Pt(10) if p.runs else None
        if not p.runs:
            p.add_run(text).font.size = Pt(10)

def add_info_box(text, color="2D5F9A"):
    """Add a colored info box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add a border effect via shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8C)
    run.font.name = 'Calibri'


# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(0)
run = title.add_run("COMPETITIVE ANALYSIS")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run("NeuraLeads vs. Instantly.ai")
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2D, 0x5F, 0x9A)
run.font.name = 'Calibri'

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.paragraph_format.space_after = Pt(40)
run = tagline.add_run("Feature Gap Analysis & Implementation Roadmap")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

# Metadata
meta_items = [
    ("Platform:", "NeuraLeads AI Agent — ra.partnerwithus.tech"),
    ("Competitor:", "Instantly.ai"),
    ("Date:", "April 3, 2026"),
    ("Version:", "1.0"),
    ("Prepared By:", "Exzelon AI Engineering Team"),
    ("Classification:", "Confidential — Internal Use Only"),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + " ")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Platform Overview Comparison",
    "3. Feature-by-Feature Analysis",
    "    3.1 Lead Generation & Database",
    "    3.2 Email Campaigns & Sequences",
    "    3.3 Email Warmup & Deliverability",
    "    3.4 Unified Inbox & Reply Management",
    "    3.5 CRM & Deal Pipeline",
    "    3.6 AI-Powered Features",
    "    3.7 Analytics & Reporting",
    "    3.8 Integrations & API",
    "    3.9 Team Collaboration & Multi-Tenancy",
    "    3.10 Billing & Pricing Infrastructure",
    "    3.11 Multichannel Outreach",
    "    3.12 Website Visitor Tracking",
    "    3.13 Done-For-You Services",
    "    3.14 Onboarding & UX",
    "4. Competitive Scorecard",
    "5. Gap Analysis Summary",
    "6. Implementation Roadmap",
    "    Phase 1: Quick Wins (Weeks 1-3)",
    "    Phase 2: Core Differentiators (Weeks 4-8)",
    "    Phase 3: Market Leadership (Weeks 9-16)",
    "    Phase 4: Platform Dominance (Weeks 17-24)",
    "7. UI/UX Enhancement Strategy",
    "8. Conclusion & Strategic Recommendations",
]
for item in toc_items:
    p = doc.add_paragraph()
    indent_level = item.count("    ")
    p.paragraph_format.left_indent = Cm(indent_level * 1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item.strip())
    run.font.size = Pt(11)
    if indent_level == 0:
        run.font.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)

p = doc.add_paragraph()
p.add_run("This document provides a comprehensive feature-by-feature analysis comparing ").font.size = Pt(10.5)
run = p.add_run("NeuraLeads AI Agent")
run.font.bold = True
run.font.size = Pt(10.5)
p.add_run(" (ra.partnerwithus.tech) against ").font.size = Pt(10.5)
run = p.add_run("Instantly.ai")
run.font.bold = True
run.font.size = Pt(10.5)
p.add_run(", the market-leading cold email outreach platform. The analysis identifies feature gaps, areas where NeuraLeads already exceeds the competition, and a phased roadmap to achieve market leadership.").font.size = Pt(10.5)

doc.add_heading("Key Findings", level=3)

findings = [
    ("200+ features already built", "across 24 dashboard pages, 36 API endpoint groups, 50+ adapters/integrations — a strong foundation that rivals Instantly.ai's core"),
    ("11 job source adapters vs. Instantly's ~3", "— NeuraLeads has a SUPERIOR lead sourcing pipeline with 3-layer deduplication"),
    ("7 email validation providers", "vs. Instantly's 1 built-in — NeuraLeads offers waterfall validation across NeverBounce, ZeroBounce, Hunter, Clearout, Emailable, MailboxValidator, and Reacher"),
    ("7 contact discovery providers", "vs. Instantly's 1 (SuperSearch) — NeuraLeads uses Apollo, Seamless, Hunter, Snov.io, RocketReach, PDL, and Proxycurl"),
    ("Key gaps to address:", "AI Copilot frontend UI, LinkedIn/multichannel automation, website visitor tracking UI, Done-for-You services, and advanced AI Reply Agent (auto-pilot mode)"),
]
for bold_part, rest in findings:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_part)
    run.font.bold = True
    run.font.size = Pt(10)
    p.add_run(rest).font.size = Pt(10)

doc.add_heading("Strategic Position", level=3)
p = doc.add_paragraph(
    "NeuraLeads is NOT behind Instantly.ai — it is architecturally superior in several dimensions "
    "(multi-provider waterfall enrichment, 11-source job aggregation, 4 AI providers, 7 email validators). "
    "The gaps are primarily in UI polish, AI agent automation (auto-reply), multichannel expansion, "
    "and self-service onboarding. This roadmap closes those gaps while building unique differentiators "
    "that no competitor offers."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. PLATFORM OVERVIEW COMPARISON
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Platform Overview Comparison", level=1)

add_styled_table(
    ["Dimension", "Instantly.ai", "NeuraLeads"],
    [
        ["Founded / Launched", "2021, SaaS product", "2024, Enterprise platform"],
        ["Target Market", "SMBs, Agencies, Sales teams", "Staffing agencies, B2B outreach teams"],
        ["Architecture", "Monolithic SaaS", "FastAPI + Next.js 14, multi-tenant"],
        ["Multi-Tenancy", "Workspace-based", "Full tenant isolation (38 models)"],
        ["Deployment", "Cloud-only (hosted)", "Self-hosted VPS + Cloud-ready"],
        ["Pricing Model", "Per-plan, per-credit", "Per-tenant, configurable plans"],
        ["Email Accounts", "Unlimited (all plans)", "Per-plan limits (configurable)"],
        ["Lead Database", "450M+ contacts (credit-based)", "11-source aggregation (own pipeline)"],
        ["Contact Discovery", "1 provider (SuperSearch)", "7 providers (waterfall enrichment)"],
        ["Email Validation", "Built-in (1 provider)", "7 providers (NeverBounce, ZeroBounce, etc.)"],
        ["AI Providers", "1 (proprietary)", "4 (Groq, OpenAI, Anthropic, Gemini)"],
        ["CRM Integration", "HubSpot, Salesforce, Pipedrive", "HubSpot, Salesforce (bidirectional)"],
        ["Job Source Adapters", "~3 (Apollo, LinkedIn)", "11 (JSearch, Apollo, TheirStack, SerpAPI, Adzuna, SearchAPI, USAJOBS, Jooble, JobDataFeeds, Coresignal, Mock)"],
        ["RBAC System", "Workspace roles (basic)", "4-role, 3-layer permissions with per-tab control"],
        ["Billing System", "Stripe (external)", "Built-in invoicing + Stripe + manual payments"],
        ["Open Source", "No", "Self-hosted, full code ownership"],
    ],
    col_widths=[4, 5.5, 5.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. FEATURE-BY-FEATURE ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Feature-by-Feature Analysis", level=1)
p = doc.add_paragraph(
    "Each sub-section below compares a specific feature area. The Status column uses: "
)
run = p.add_run("FULL")
run.font.bold = True
run.font.color.rgb = RGBColor(0x15, 0x58, 0x24)
p.add_run(" (fully implemented), ")
run = p.add_run("PARTIAL")
run.font.bold = True
run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
p.add_run(" (partially built), ")
run = p.add_run("MISSING")
run.font.bold = True
run.font.color.rgb = RGBColor(0x72, 0x1C, 0x24)
p.add_run(" (not yet built), ")
run = p.add_run("SUPERIOR")
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x4E, 0x92)
p.add_run(" (exceeds competitor).")

# ─── 3.1 Lead Generation ────────────────────────────────────────────────
doc.add_heading("3.1 Lead Generation & Database", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["B2B Contact Database", "450M+ contacts (SuperSearch)", "Superior", "11 aggregated sources, 3-layer dedup, sub-source tracking"],
        ["Natural Language Search", "AI prompt-based lead search", "Full", "AI Lead Search: NLP query -> SQL filters"],
        ["Advanced Filters", "Job title, industry, size, revenue, tech stack, funding", "Full", "States, titles, industries, company size, exclusion keywords"],
        ["Saved Searches / Smart Lists", "Saved filters", "Full", "SavedSearch model with sharing support"],
        ["Lookalike Search", "Domain lookalike matching", "Missing", "Not implemented — high priority"],
        ["Intent Data / Buying Signals", "Hiring intent, funding signals", "Missing", "Not implemented"],
        ["Technology Stack Filtering", "TechStack filter via TheirStack", "Partial", "TheirStack adapter exists but no dedicated UI filter"],
        ["Waterfall Email Enrichment", "5+ provider waterfall", "Superior", "7-provider waterfall (Apollo, Seamless, Hunter, Snov.io, RocketReach, PDL, Proxycurl)"],
        ["ICP Wizard", "Basic ICP definition", "Superior", "AI-powered ICP generation with industry/title/location recommendations"],
        ["Lead Scoring", "Basic scoring", "Full", "AI lead scoring (salary, location, fit)"],
        ["CSV Import / Export", "CSV upload with verification", "Full", "CSV export in outreach, lead management"],
        ["Credit System", "Credit-based (pay per lead)", "Missing", "No credit system — unlimited within plan limits"],
        ["Job Board Aggregation", "Not a primary feature", "Superior", "11 job source adapters with parallel fetching"],
    ],
    status_col=2
)

# ─── 3.2 Email Campaigns ────────────────────────────────────────────────
doc.add_heading("3.2 Email Campaigns & Sequences", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Multi-Step Sequences", "Unlimited steps with delays", "Full", "SequenceStep model with email/wait/condition types"],
        ["A/B Testing", "A/Z testing (up to 26 variants)", "Full", "Multi-variant with weighted distribution + chi-squared auto-optimize"],
        ["Spintax Support", "{option1|option2} syntax", "Full", "Nested spintax support in campaign engine"],
        ["Conditional Logic (If/Then)", "Liquid syntax conditions", "Full", "Condition steps with if/then branching"],
        ["AI Sequence Generator", "Copilot generates sequences", "Full", "AI sequence generator with template fallback"],
        ["AI Subject Line Generator", "AI-powered subjects", "Partial", "Part of AI content generation, no standalone UI"],
        ["Personalization Variables", "First name, company, custom vars", "Full", "Dynamic variable insertion in campaigns"],
        ["Inbox Rotation", "Rotate across sending accounts", "Full", "Round-robin mailbox selection per campaign"],
        ["Smart Send Windows", "Schedule by time/day", "Full", "Timezone-aware (US state -> timezone), business hours"],
        ["Slow Ramp (Campaign)", "+2 emails/day gradual increase", "Missing", "No campaign-level slow ramp — warmup has it"],
        ["Campaign Templates Library", "Pre-built templates", "Partial", "Email templates exist but no pre-built campaign library"],
        ["Deliverability Test (Pre-Send)", "Test email + spam score", "Full", "Spam checker (100+ trigger words, 0-100 score)"],
        ["Auto-Pause Underperformers", "Auto-pause low-performing campaigns", "Missing", "No auto-pause based on performance metrics"],
        ["Campaign Cloning", "Clone campaigns", "Partial", "Not explicit clone feature — manual recreation"],
        ["Unsubscribe Management", "Built-in unsubscribe", "Full", "Unsubscribe footer + per-contact tracking"],
    ],
    status_col=2
)

# ─── 3.3 Email Warmup ────────────────────────────────────────────────────
doc.add_heading("3.3 Email Warmup & Deliverability", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Peer-to-Peer Warmup", "1M+ account warmup pool", "Full", "Peer warmup between mailboxes within tenant"],
        ["Slow Ramp (Warmup)", "Gradual daily increase (2→4→6...)", "Full", "3 warmup profiles: Conservative 45d, Standard 30d, Aggressive 20d"],
        ["Read Emulation", "Headless browser scrolling", "Full", "IMAP-based read marking (imap_reader.py)"],
        ["Health Score", "7-day inbox vs spam ratio", "Full", "Warmup health scores with daily tracking"],
        ["DNS Checking (SPF/DKIM/DMARC)", "Setup guidance", "Full", "Automated DNS checker with per-record validation"],
        ["Blacklist Monitoring", "Blacklist alerts", "Full", "IP/domain blacklist monitoring service"],
        ["Custom Tracking Domain", "Custom domain warmup", "Full", "TrackingDomain model with CNAME verification"],
        ["Inbox Placement Testing", "Automated inbox placement tests", "Partial", "Inbox placement estimation (not seed-based)"],
        ["Open/Click Tracking", "Pixel + link redirect", "Full", "/t/{id}/px.gif pixel + /t/{id}/l redirect"],
        ["AI Auto-Reply (Warmup)", "Automated warmup replies", "Full", "AI-generated warmup auto-replies via Groq"],
        ["Server & IP Rotation (SISR)", "Dedicated IP pools on Light Speed", "Missing", "No IP rotation — single sending infrastructure"],
        ["Multi-Tier Warmup Pools", "Basic/Standard/Premium pools", "Missing", "Single warmup pool (peer-to-peer within tenant)"],
        ["Auto-Recovery", "Auto-pause and re-warm", "Full", "auto_recovery.py service"],
        ["Domain Reputation Score", "Built-in reputation tracking", "Full", "domain_reputation.py with scoring"],
        ["Warmup Report Export", "Analytics export", "Full", "PDF/CSV report export (report_exporter.py)"],
    ],
    status_col=2
)

# ─── 3.4 Unified Inbox ──────────────────────────────────────────────────
doc.add_heading("3.4 Unified Inbox & Reply Management", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Unified Inbox (Unibox)", "All replies in one view", "Full", "InboxMessage model, thread grouping, multi-mailbox"],
        ["AI Sentiment Analysis", "AI Custom Reply Labels", "Full", "Rule-based + LLM fallback sentiment analysis"],
        ["Auto-Categorization", "Interested/Not Interested/OOO/Referral", "Full", "6 categories: interested, not_interested, ooo, question, referral, do_not_contact"],
        ["AI Reply Suggestions", "AI drafts contextual replies", "Full", "AI reply suggestions from conversation context"],
        ["AI Reply Agent (Auto-Send)", "Autopilot mode — sends within 5 min", "Missing", "Suggestions only — no auto-send mode"],
        ["AI Reply Agent (HITL)", "Human-in-the-loop approval", "Missing", "No approve-and-send workflow"],
        ["Objection Handling AI", "AI handles objections automatically", "Missing", "No objection handling logic"],
        ["Calendar Link Auto-Share", "Auto-shares calendar on interest", "Missing", "No calendar integration"],
        ["Reply Macros / Templates", "Quick reply templates", "Missing", "No reply template/macro system"],
        ["Slack Notifications", "Reply alerts to Slack channels", "Partial", "Slack webhook adapter exists, not wired to inbox"],
        ["Thread History View", "Full email thread in inbox", "Full", "Thread grouping via Message-ID chain or hash"],
        ["Multi-Channel Inbox", "Email + SMS + Call in one view", "Partial", "Email only — SMS/Call not integrated"],
    ],
    status_col=2
)

# ─── 3.5 CRM ────────────────────────────────────────────────────────────
doc.add_heading("3.5 CRM & Deal Pipeline", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Deal Pipeline (Kanban)", "Visual pipeline stages", "Full", "7-stage Kanban: New Lead -> Won/Lost"],
        ["Deal Activity Timeline", "Interaction history per deal", "Full", "Activity timeline per deal"],
        ["Deal Statistics", "Win rate, pipeline value", "Full", "Win rate, avg deal size, pipeline value"],
        ["Deal Task Management", "Task assignment within deals", "Full", "DealTask model: assignee, due date, priority, status"],
        ["Lead Assignment (Round Robin)", "Automatic lead distribution", "Missing", "No round-robin lead assignment"],
        ["Goal Tracking", "Monthly/quarterly targets", "Missing", "No goal/quota tracking"],
        ["Salesflows", "Custom automated views", "Missing", "No salesflow automation"],
        ["CRM Sync (HubSpot)", "Native HubSpot integration", "Full", "Bidirectional sync (contacts, deals)"],
        ["CRM Sync (Salesforce)", "Native Salesforce integration", "Full", "Bidirectional sync (contacts, opportunities)"],
        ["CRM Sync (Pipedrive)", "Native Pipedrive integration", "Missing", "Not implemented"],
        ["Auto-Forward to CRM", "Sync interested replies", "Full", "crm_auto_forward.py service"],
        ["Opportunities View", "Opportunity tracking", "Full", "Deal pipeline with value and probability"],
        ["Shared Context", "Team-wide lead visibility", "Full", "Multi-tenant with role-based data access"],
    ],
    status_col=2
)

# ─── 3.6 AI Features ────────────────────────────────────────────────────
doc.add_heading("3.6 AI-Powered Features", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["AI Email Content Generation", "AI personalization in sequences", "Full", "4 AI providers: Groq, OpenAI, Anthropic, Gemini"],
        ["AI Sequence Generator", "Copilot creates full sequences", "Full", "AI sequence generation with template fallback"],
        ["AI ICP Wizard", "Basic ICP definition", "Superior", "AI-generated profiles with industry/title/state/size recommendations"],
        ["AI Lead Scoring", "Basic scoring", "Full", "AI scoring by salary, location, company fit"],
        ["AI Natural Language Search", "Natural language lead finder", "Full", "NLP query parsing -> SQL filter dict"],
        ["AI Copilot (Chat Assistant)", "Copilot for research, campaigns, analytics", "Partial", "Backend endpoint exists, NO frontend chat widget"],
        ["AI Reply Agent", "Auto-reply with objection handling", "Missing", "AI reply suggestions exist but no auto-send agent"],
        ["AI Company Enrichment", "Enrichment via database", "Full", "3-layer: leads data -> Clearbit -> AI fallback"],
        ["AI Spam Checker", "SpamAssassin scoring", "Full", "100+ trigger words, pattern matching, 0-100 score"],
        ["AI Sentiment Analysis", "Reply label classification", "Full", "Rule-based + LLM fallback"],
        ["Multi-LLM Support", "Single proprietary model", "Superior", "4 providers with model selection per tenant"],
        ["AI Web Research Agent", "SuperSearch AI researcher", "Missing", "No web research agent"],
    ],
    status_col=2
)

# ─── 3.7 Analytics ──────────────────────────────────────────────────────
doc.add_heading("3.7 Analytics & Reporting", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Campaign Analytics", "Opens, clicks, replies, bounces", "Full", "Per-campaign and per-step analytics"],
        ["Global Analytics Dashboard", "Cross-campaign performance", "Full", "Analytics page with campaign comparison"],
        ["Team Leaderboard", "Team performance tracking", "Full", "Team leaderboard with per-user metrics"],
        ["Revenue / ROI Tracking", "Pipeline value, conversions", "Full", "CostEntry model, revenue metrics, ROI analytics"],
        ["A/B Test Analytics", "Variant performance comparison", "Full", "Per-variant stats with auto-optimization"],
        ["Custom Date Ranges", "Flexible date filtering", "Full", "Date range filters on analytics"],
        ["Export Analytics", "CSV/PDF export", "Partial", "Some exports available, not comprehensive"],
        ["Real-Time Activity Feed", "Live campaign activity", "Full", "AutomationEvent log with live updates"],
        ["Meeting Booking Rate", "Track meetings booked", "Missing", "No meeting tracking (no calendar integration)"],
        ["Pipeline Forecast", "Revenue forecasting", "Missing", "No forecast/prediction engine"],
        ["Heatmap (Best Send Times)", "Engagement by time of day", "Missing", "No heatmap visualization"],
    ],
    status_col=2
)

# ─── 3.8 Integrations ───────────────────────────────────────────────────
doc.add_heading("3.8 Integrations & API", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["REST API", "API V2 (comprehensive)", "Full", "Full REST API with 36 endpoint groups"],
        ["Webhook System", "8+ event types", "Full", "8 events, HMAC-SHA256 signed, exponential backoff"],
        ["API Key Management", "API key auth", "Full", "SHA-256 hashed, scoped, with expiry"],
        ["Zapier Integration", "Native Zapier app", "Partial", "REST hook subscribe/unsubscribe endpoints exist"],
        ["Make.com Integration", "Native Make module", "Partial", "Same webhook endpoints work"],
        ["Slack Notifications", "Native Slack + webhook", "Full", "Slack webhook adapter"],
        ["Microsoft Teams", "Not native", "Full", "Teams webhook adapter"],
        ["OAuth 2.0 (M365/Gmail)", "Email provider OAuth", "Full", "OAuth helper for Microsoft 365 and Gmail"],
        ["Calendly Integration", "Calendar booking links", "Missing", "No calendar integration"],
        ["Zapier App (Published)", "Listed on Zapier marketplace", "Missing", "Endpoints exist but no published Zapier app"],
        ["n8n / Pabbly", "Pabbly integration", "Missing", "Mentioned in docs, not in code"],
        ["Clay Integration", "Native Clay connector", "Missing", "Not implemented"],
    ],
    status_col=2
)

# ─── 3.9 Team & Multi-Tenancy ───────────────────────────────────────────
doc.add_heading("3.9 Team Collaboration & Multi-Tenancy", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Multi-Tenancy", "Workspace-based", "Superior", "Full tenant isolation, 38 models, plan limits"],
        ["Role-Based Access (RBAC)", "Basic workspace roles", "Superior", "4-role, 3-layer permissions (default->role->user override)"],
        ["Per-Module Permissions", "Limited", "Superior", "Independent per-tab settings permissions"],
        ["Super Admin Impersonation", "Not available", "Superior", "X-Tenant-ID header, admin panel"],
        ["Unlimited Seats", "Unlimited on CRM plans", "Full", "Configurable per plan"],
        ["User Management", "Basic user admin", "Full", "CRUD, role assignment, email verification, lockout"],
        ["Audit Logging", "Basic activity log", "Superior", "Login history, IP tracking, auth audit, activity log"],
        ["Deal Task Assignment", "Team task management", "Full", "Assignee, due date, priority, status per deal"],
        ["Team Chat / @Mentions", "Not available", "Missing", "Not implemented"],
        ["Shared Notes / Comments", "Basic notes per lead", "Missing", "Not implemented"],
        ["Real-Time Collaboration", "Not available", "Missing", "No WebSockets (planned)"],
    ],
    status_col=2
)

# ─── 3.10 Billing ───────────────────────────────────────────────────────
doc.add_heading("3.10 Billing & Pricing Infrastructure", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Stripe Payments", "Stripe for all plans", "Full", "StripeGateway with checkout sessions"],
        ["Invoice Generation", "Stripe-managed invoicing", "Superior", "Built-in INV-YYYY-NNNN, auto-monthly, PDF generation"],
        ["Manual Payments", "Not available", "Superior", "ManualGateway + bank transfer/check support"],
        ["Overdue Reminders", "Stripe dunning", "Full", "Auto reminders every 3 days, max 5"],
        ["Payment Records", "Stripe dashboard", "Full", "PaymentRecord model with full audit trail"],
        ["Tenant Plan Limits", "Tier-based limits", "Full", "max_users, max_mailboxes, max_contacts, max_campaigns, max_leads"],
        ["Credit System", "Per-lead credit pricing", "Missing", "No credit-based billing"],
        ["Usage Metering", "API-based metering", "Missing", "No per-API-call metering"],
    ],
    status_col=2
)

# ─── 3.11 Multichannel ──────────────────────────────────────────────────
doc.add_heading("3.11 Multichannel Outreach", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Email (SMTP/OAuth)", "Primary channel", "Full", "SMTP, M365 OAuth, Gmail OAuth"],
        ["SMS Outreach", "Available on Hyper CRM", "Partial", "Twilio adapter exists, NO UI/endpoints"],
        ["Voice Calling (Click-to-Call)", "Available on Hyper CRM", "Partial", "Twilio adapter exists, NO UI/endpoints"],
        ["LinkedIn Messaging", "Not native (via Zapier)", "Missing", "LinkedIn URLs tracked but no outreach automation"],
        ["WhatsApp", "Not native", "Missing", "Not implemented"],
        ["Social Media (Twitter/FB/IG)", "Not available", "Missing", "Not implemented"],
        ["Multichannel Sequences", "Email + SMS + Call in one flow", "Missing", "Email-only sequences"],
    ],
    status_col=2
)

# ─── 3.12 Website Visitor Tracking ───────────────────────────────────────
doc.add_heading("3.12 Website Visitor Tracking", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Tracking Pixel / Script", "JavaScript pixel for websites", "Missing", "VisitorEvent model exists, NO pixel/script"],
        ["Visitor Identification", "Business email via reverse IP", "Missing", "Model has company_name/domain fields, NO implementation"],
        ["Page Visit Tracking", "Track which pages visited", "Missing", "page_url field in model, NO tracking logic"],
        ["Company Enrichment", "Company details from visitor data", "Missing", "Not implemented"],
        ["Lead Creation from Visitors", "Auto-create leads from visitors", "Missing", "Not implemented"],
        ["CRM Integration (Visitors)", "Add visitors to CRM", "Missing", "Not implemented"],
    ],
    status_col=2
)

# ─── 3.13 DFY Services ──────────────────────────────────────────────────
doc.add_heading("3.13 Done-For-You (DFY) Services", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Automated Domain Registration", "DFY domain setup service", "Missing", "Not implemented"],
        ["DNS Auto-Configuration", "SPF/DKIM/DMARC auto-setup", "Missing", "DNS checker validates but doesn't create records"],
        ["Email Account Provisioning", "Create email accounts automatically", "Missing", "Manual OAuth connection only"],
        ["Secondary Domain Suggestions", "AI domain name suggestions", "Missing", "Not implemented"],
        ["White-Label Setup", "Agency white-labeling", "Missing", "Not implemented"],
    ],
    status_col=2
)

# ─── 3.14 Onboarding & UX ───────────────────────────────────────────────
doc.add_heading("3.14 Onboarding & UX", level=2)

add_status_table(
    ["Feature", "Instantly.ai", "NeuraLeads Status", "Notes"],
    [
        ["Getting Started Wizard", "Setup flow for new users", "Full", "6-step auto-detected checklist"],
        ["Interactive Tour", "Guided product tour", "Full", "driver.js spotlight (~10 steps), help button replay"],
        ["Email Verification", "Account verification", "Full", "JWT-based email verification on signup"],
        ["Demo Data Seeding", "Not available", "Superior", "Auto-seeds sample data for new starter tenants"],
        ["Dark Mode", "Not available", "Full", "Theme provider with dark/light mode"],
        ["Mobile Responsive", "Responsive design", "Full", "Tailwind CSS mobile-first"],
        ["Offline Indicator", "Not available", "Full", "Offline connection banner"],
        ["Impersonation Banner", "Not available", "Superior", "Visual indicator when super admin impersonates"],
        ["Per-User Dismiss", "Not available", "Superior", "Each user controls their onboarding state"],
        ["Masterclass / Academy", "Free cold email masterclass", "Missing", "No educational content / academy"],
    ],
    status_col=2
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. COMPETITIVE SCORECARD
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Competitive Scorecard", level=1)

p = doc.add_paragraph("Score out of 10 for each category, based on feature completeness and quality.")

add_styled_table(
    ["Category", "Instantly.ai", "NeuraLeads", "Gap", "Verdict"],
    [
        ["Lead Generation & Database", "9", "8", "+1", "NeuraLeads superior in sources; needs intent data"],
        ["Email Campaigns & Sequences", "9", "8", "+1", "Missing slow ramp and auto-pause"],
        ["Email Warmup & Deliverability", "9", "8", "+1", "Missing SISR and multi-tier pools"],
        ["Unified Inbox", "9", "6", "+3", "Missing AI auto-reply agent and calendar"],
        ["CRM & Deal Pipeline", "7", "7", "0", "At parity — both have Kanban + CRM sync"],
        ["AI Features", "8", "7", "+1", "Missing Copilot UI and AI Reply Agent"],
        ["Analytics & Reporting", "8", "7", "+1", "Missing forecasting and heatmaps"],
        ["Integrations & API", "8", "6", "+2", "Missing Zapier app, Calendly, Clay"],
        ["Team & Multi-Tenancy", "5", "9", "-4", "NeuraLeads FAR superior"],
        ["Billing Infrastructure", "7", "8", "-1", "NeuraLeads superior (built-in invoicing)"],
        ["Multichannel Outreach", "6", "2", "+4", "Biggest gap — SMS/LinkedIn/Call UI needed"],
        ["Website Visitor Tracking", "6", "1", "+5", "Model exists but no implementation"],
        ["DFY Services", "7", "0", "+7", "Not implemented — premium feature opportunity"],
        ["Onboarding & UX", "7", "8", "-1", "NeuraLeads superior (tour, demo data, dark mode)"],
        ["", "", "", "", ""],
        ["TOTAL (out of 140)", "105", "85", "+20", "Gap closable in 16-24 weeks"],
    ],
    col_widths=[4.5, 2.5, 2.5, 1.5, 5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. GAP ANALYSIS SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Gap Analysis Summary", level=1)

doc.add_heading("Critical Gaps (Must Fix)", level=3)
critical_gaps = [
    "AI Reply Agent (Auto-Pilot mode with approve-and-send workflow)",
    "AI Copilot Frontend Chat Widget (backend exists, needs chat UI)",
    "Campaign Slow Ramp (gradual daily limit increase per campaign)",
    "Multichannel Sequences (SMS + LinkedIn steps in campaign builder)",
    "Website Visitor Tracking (implement pixel, API, and dashboard)",
]
for gap in critical_gaps:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(gap)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

doc.add_heading("Important Gaps (Should Fix)", level=3)
important_gaps = [
    "Calendar Integration (Calendly/Cal.com — auto-share links on interest)",
    "Lookalike Search (find companies similar to a given domain)",
    "Lead Round-Robin Assignment (distribute leads across team members)",
    "Reply Macros / Quick Templates (canned responses in Unibox)",
    "Auto-Pause Underperforming Campaigns (based on bounce/spam thresholds)",
    "Zapier App Publication (list on Zapier marketplace for visibility)",
    "Engagement Heatmap (best send times by open/reply rates)",
    "Pipeline Forecasting (AI-based revenue predictions)",
]
for gap in important_gaps:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(gap)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xD4, 0x8A, 0x0B)

doc.add_heading("Nice-to-Have Gaps (Can Wait)", level=3)
nice_gaps = [
    "DFY Domain/Email Setup Service (premium upsell opportunity)",
    "Pipedrive CRM Integration (third CRM connector)",
    "Clay Integration (data enrichment workflow)",
    "White-Label / Agency Mode (custom branding per agency)",
    "SISR (Server & IP Sharding and Rotation for high volume)",
    "Multi-Tier Warmup Pools (Basic/Standard/Premium)",
    "Academy / Masterclass Content (educational marketing)",
    "Intent Data / Buying Signals (hiring, funding, tech adoption signals)",
    "Credit System for API Metering (usage-based billing add-on)",
]
for gap in nice_gaps:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(gap)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

doc.add_heading("Areas Where NeuraLeads Already Exceeds Instantly.ai", level=3)
superior = [
    "11 Job Source Adapters vs. ~3 — massive lead sourcing advantage",
    "7 Contact Discovery Providers (waterfall) vs. 1 — deeper enrichment",
    "7 Email Validation Providers vs. 1 — higher validation accuracy",
    "4 AI Providers (Groq/OpenAI/Anthropic/Gemini) vs. 1 — flexibility and cost control",
    "Full Multi-Tenancy with 38 tenant-scoped models vs. basic workspaces",
    "3-Layer RBAC (default → role → user override) with per-tab permissions vs. basic roles",
    "Built-in Invoicing System with PDF generation vs. Stripe-only billing",
    "Super Admin Impersonation with visual banner vs. no equivalent",
    "Demo Data Auto-Seeding for new tenants vs. empty workspace",
    "3-Layer Lead Deduplication (external_job_id → employer_linkedin → company+title+state+city)",
    "Database Backup/Restore System with audit trail vs. no equivalent",
    "Login History & Security Audit with IP tracking vs. basic logs",
]
for item in superior:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x4E, 0x92)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. IMPLEMENTATION ROADMAP
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Implementation Roadmap", level=1)

p = doc.add_paragraph(
    "The following phased roadmap prioritizes features by competitive impact, "
    "technical feasibility, and revenue potential. Each phase builds on the previous one."
)

# ─── Phase 1 ─────────────────────────────────────────────────────────────
doc.add_heading("Phase 1: Quick Wins (Weeks 1-3)", level=2)
p = doc.add_paragraph()
run = p.add_run("Goal: ")
run.font.bold = True
p.add_run("Close the most visible gaps with minimal architectural changes.")

phase1 = [
    ["1.1", "AI Copilot Chat Widget", "Build floating chat widget on frontend that calls existing /copilot/chat endpoint. Draggable, resizable, context-aware (passes current page URL). Include quick-action buttons for common tasks.", "Frontend only — backend exists"],
    ["1.2", "Campaign Slow Ramp", "Add slow_ramp_enabled and slow_ramp_increment fields to Campaign model. Campaign engine checks current day count and caps sends accordingly. Default: +2/day/account.", "Backend model + engine logic"],
    ["1.3", "Reply Macros / Quick Templates", "Add ReplyMacro model (title, body, variables, category). Inbox UI shows macro picker dropdown when composing replies. Auto-insert with variable substitution.", "New model + inbox UI update"],
    ["1.4", "Auto-Pause Underperformers", "Add bounce_threshold and spam_threshold fields to Campaign. Scheduler job checks campaigns every hour — auto-pauses if thresholds exceeded. Alert in dashboard.", "Backend scheduler + model fields"],
    ["1.5", "Lookalike Company Search", "Add /leads/lookalike endpoint that takes a domain, enriches it (industry, size, tech stack), then queries existing leads for similar companies. Surface in leads page.", "New endpoint + UI button"],
    ["1.6", "Engagement Heatmap", "Add analytics endpoint that aggregates opens/replies by hour-of-day and day-of-week. Frontend renders as a color-coded grid (green = high engagement). Include on analytics page.", "New endpoint + frontend component"],
]

add_styled_table(
    ["#", "Feature", "Implementation Details", "Scope"],
    phase1,
    col_widths=[1, 3.5, 8, 3.5]
)

# ─── Phase 2 ─────────────────────────────────────────────────────────────
doc.add_heading("Phase 2: Core Differentiators (Weeks 4-8)", level=2)
p = doc.add_paragraph()
run = p.add_run("Goal: ")
run.font.bold = True
p.add_run("Build the AI Reply Agent and multichannel foundation — the two biggest competitive gaps.")

phase2 = [
    ["2.1", "AI Reply Agent (HITL Mode)", "Build human-in-the-loop auto-reply system. When inbox receives a reply: 1) AI analyzes intent (interested/objection/question/OOO). 2) AI generates contextual response. 3) Response appears in inbox as 'Draft — Approve or Edit'. 4) User clicks 'Send' to approve. Track approval rate for optimization.", "New service + inbox UI overhaul"],
    ["2.2", "AI Reply Agent (Autopilot Mode)", "Extension of 2.1 — add auto_reply_enabled flag per campaign. When enabled, approved drafts auto-send after configurable delay (default: 5 min). Include safety rails: max 3 auto-replies per thread, escalation on negative sentiment, pause on unknown intent.", "Extends 2.1 + scheduler job"],
    ["2.3", "Calendar Integration", "Integrate with Calendly and Cal.com APIs. Add calendar_link field to User model. AI Reply Agent auto-includes calendar link when lead expresses interest. Track meeting bookings as deal activities.", "New adapter + user field + AI prompt update"],
    ["2.4", "SMS Outreach UI", "Build SMS step type in campaign sequences. Activate existing Twilio adapter. Add SMS compose in inbox. Create SMS templates. Track delivery/response.", "Frontend pages + API endpoints (adapter exists)"],
    ["2.5", "Lead Round-Robin Assignment", "Add assignment_mode to Campaign (manual/round_robin/weighted). When new leads enter pipeline, auto-assign to team members based on mode. Show assignment in lead detail.", "Campaign model + assignment service"],
    ["2.6", "Pipeline Forecasting", "Build AI-powered deal forecasting. Analyze historical win rates, deal velocity, and stage duration to predict monthly/quarterly revenue. Show forecast chart on deals page and analytics.", "New service + analytics endpoint + chart component"],
]

add_styled_table(
    ["#", "Feature", "Implementation Details", "Scope"],
    phase2,
    col_widths=[1, 3.5, 8, 3.5]
)

# ─── Phase 3 ─────────────────────────────────────────────────────────────
doc.add_heading("Phase 3: Market Leadership (Weeks 9-16)", level=2)
p = doc.add_paragraph()
run = p.add_run("Goal: ")
run.font.bold = True
p.add_run("Build features that go BEYOND Instantly.ai — unique differentiators that no competitor offers.")

phase3 = [
    ["3.1", "Website Visitor Tracking", "Build JS tracking pixel that captures page visits, referrer, session data. Implement reverse-IP lookup to identify companies. Create visitor dashboard with: visitor timeline, company identification, page heatmap, auto-lead creation. Deploy via embeddable script tag.", "New frontend page + tracking script + API endpoints + IP lookup service"],
    ["3.2", "LinkedIn Connection Automation", "Build LinkedIn integration via browser extension or API partner (Phantombuster/Dripify). Add LinkedIn step types to campaign sequences: connection request, message, profile view. Track LinkedIn engagement alongside email.", "New adapter + campaign step type + frontend"],
    ["3.3", "Voice Calling (Click-to-Call)", "Activate existing Twilio adapter for voice. Build click-to-call button on contact cards. Auto-log call recordings and transcriptions. Add call step type to campaign sequences. Show call history in contact timeline.", "Frontend UI + API endpoints (adapter exists)"],
    ["3.4", "AI Objection Handling Library", "Build library of common sales objections with AI-generated responses. Train per-tenant models on approved replies. AI Reply Agent uses this library for intelligent objection responses. Include objection analytics.", "New model + AI prompt engineering + analytics"],
    ["3.5", "Advanced A/Z Testing", "Extend current A/B to support up to 26 variants per step. Add: send time testing, subject line only testing, body only testing. Real-time winner visualization. Auto-graduate winners after statistical significance.", "Extend SequenceStep + analytics UI"],
    ["3.6", "Zapier App Publication", "Build and publish official Zapier app with: triggers (reply received, lead created, deal won, email bounced), actions (add lead, create campaign, update deal). Submit to Zapier marketplace for public listing.", "Zapier app build + marketplace submission"],
    ["3.7", "Intent Data & Buying Signals", "Integrate hiring intent (from job sources), funding signals (Crunchbase API), tech stack changes (BuiltWith/Wappalyzer). Add intent score to leads. Filter by buying signals.", "New adapters + lead model fields + UI filters"],
]

add_styled_table(
    ["#", "Feature", "Implementation Details", "Scope"],
    phase3,
    col_widths=[1, 3.5, 8, 3.5]
)

# ─── Phase 4 ─────────────────────────────────────────────────────────────
doc.add_heading("Phase 4: Platform Dominance (Weeks 17-24)", level=2)
p = doc.add_paragraph()
run = p.add_run("Goal: ")
run.font.bold = True
p.add_run("Premium features, agency tools, and self-service infrastructure for market dominance.")

phase4 = [
    ["4.1", "Done-For-You (DFY) Service", "Build automated domain setup: 1) Domain registration via registrar API (Namecheap/Cloudflare). 2) Auto-configure DNS (SPF/DKIM/DMARC). 3) Create email accounts via provider API. 4) Auto-warm all accounts. Package as premium add-on with setup wizard.", "New service + registrar API + provider APIs + wizard UI"],
    ["4.2", "White-Label / Agency Mode", "Add tenant-level branding: custom logo, colors, domain (CNAME). Build agency dashboard with client overview, cross-client analytics, and managed services view. Per-client billing.", "Tenant model extension + theme system + agency dashboard"],
    ["4.3", "Real-Time Collaboration (WebSockets)", "Add WebSocket layer for: real-time inbox updates, live campaign activity feed, team member presence indicators, collaborative deal notes. Use Socket.IO or native WebSocket.", "Infrastructure change + frontend real-time components"],
    ["4.4", "Multichannel Sequences Builder", "Unified visual sequence builder supporting email, SMS, LinkedIn, call, wait, and condition steps in a single drag-and-drop flow. Branch logic based on channel engagement.", "Major frontend feature + backend orchestration"],
    ["4.5", "Credit & Usage Metering System", "Build credit-based metering for API calls, AI generations, email verifications, and lead lookups. Usage dashboard per tenant. Overage alerts. Auto-upgrade suggestions.", "New billing module + metering middleware"],
    ["4.6", "Academy & Knowledge Base", "Build in-app learning center: video tutorials, written guides, best practices, certification program. Track user progress. Context-sensitive help links throughout the app.", "Content + frontend pages + progress tracking"],
    ["4.7", "SISR (IP Sharding & Rotation)", "Build dedicated IP pool system for high-volume senders. Auto-rotate sending IPs across campaigns. IP warmup automation. Per-IP deliverability tracking.", "Infrastructure + new service + mailbox model extension"],
]

add_styled_table(
    ["#", "Feature", "Implementation Details", "Scope"],
    phase4,
    col_widths=[1, 3.5, 8, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. UI/UX ENHANCEMENT STRATEGY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. UI/UX Enhancement Strategy", level=1)

p = doc.add_paragraph(
    "To surpass Instantly.ai in user experience, NeuraLeads must establish a distinctive visual identity "
    "and interaction paradigm. The following enhancements define our unique style."
)

doc.add_heading("7.1 Design System Overhaul", level=2)
design_items = [
    ("Glassmorphism Cards:", "Replace flat cards with subtle glass-morphism effects (semi-transparent backgrounds, blur, border glow). This creates a premium, modern feel distinct from Instantly's flat Material-inspired design."),
    ("Micro-Animations:", "Add subtle transitions on all state changes — card hover lifts, button pulse on click, skeleton shimmer on load, slide-in notifications. Use Framer Motion for 60fps animations."),
    ("Data Visualization Upgrade:", "Replace basic Recharts with custom-styled charts featuring gradient fills, animated data points, and interactive tooltips. Add sparkline mini-charts in table rows for instant trend visibility."),
    ("Command Palette (Cmd+K):", "Build a global command palette (like VS Code / Linear) for instant navigation, search, and actions. Type to find any lead, campaign, mailbox, or setting. Power users will love this."),
    ("Contextual Side Panels:", "Replace full-page navigation with slide-out detail panels. Click a lead in a table → side panel opens with full details, contacts, and actions. No page navigation needed for quick tasks."),
    ("Smart Notifications Center:", "Build a notification bell with categorized alerts (campaigns, inbox, warmup, system). Group by priority. Mark as read/unread. Deep-link to relevant page."),
]
for title, desc in design_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + " ")
    run.font.bold = True
    run.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

doc.add_heading("7.2 Dashboard Revolution", level=2)
dashboard_items = [
    ("Customizable Widget Grid:", "Let users drag-and-drop dashboard widgets to create their own layout. Persist per-user. Offer pre-built layouts for different roles (Sales Rep, Manager, Admin)."),
    ("AI-Powered Dashboard Insights:", "Show AI-generated daily briefings: 'Your best campaign this week had 24% reply rate. 5 leads are ready to close. 2 mailboxes need attention.' Actionable, not just data."),
    ("Pipeline Funnel Visualization:", "Animated funnel showing leads flowing through sourcing → enrichment → validation → outreach → reply → deal. Click any stage to drill down. Real-time numbers."),
    ("Goal Progress Rings:", "Circular progress rings for monthly targets (leads sourced, emails sent, deals closed, revenue). Color changes from blue → green as targets approach."),
]
for title, desc in dashboard_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + " ")
    run.font.bold = True
    run.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

doc.add_heading("7.3 Inbox Experience", level=2)
inbox_items = [
    ("Gmail-Like Thread View:", "Show full email threads with inline replies, quoted text collapsing, and rich formatting. Currently threads are listed — should feel like reading a real email conversation."),
    ("AI Reply Confidence Score:", "Show a confidence badge on AI-generated reply suggestions (High/Medium/Low). Users learn to trust AI gradually. Track which confidence levels get approved most."),
    ("Quick Actions Bar:", "Horizontal action bar in inbox: Mark Interested, Schedule Follow-Up, Add to Deal, Forward to CRM, Archive. One-click actions without opening menus."),
    ("Snooze & Reminders:", "Snooze a reply to resurface later (1hr, tomorrow, next week, custom). Set reminders on threads that need follow-up. Snoozed items reappear at the top."),
]
for title, desc in inbox_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + " ")
    run.font.bold = True
    run.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

doc.add_heading("7.4 Campaign Builder UX", level=2)
campaign_items = [
    ("Visual Sequence Builder:", "Replace the current step-list UI with a visual flow diagram (like n8n or Zapier). Drag-and-drop nodes for email, wait, condition, SMS, LinkedIn. Lines show flow with branching paths."),
    ("Live Preview Panel:", "Side-by-side preview of the email as you compose. Shows how it looks in Gmail, Outlook, and mobile. Includes spam score in real-time as you type."),
    ("Template Gallery:", "Curated library of proven campaign templates organized by industry, use case, and goal. One-click import with variable mapping. Community-contributed templates (future)."),
    ("A/B Test Visualizer:", "Real-time chart showing variant performance as data comes in. Confidence interval bands. Clear 'Winner' badge when statistical significance reached. One-click promote winner."),
]
for title, desc in campaign_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + " ")
    run.font.bold = True
    run.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. CONCLUSION & STRATEGIC RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Conclusion & Strategic Recommendations", level=1)

p = doc.add_paragraph(
    "NeuraLeads is not a follower — it is architecturally positioned to SURPASS Instantly.ai. "
    "The platform already has deeper data infrastructure (11 job sources, 7 contact providers, 7 validators, "
    "4 AI models), stronger multi-tenancy, and more granular access control than any competitor in the market."
)

doc.add_heading("Strategic Priorities", level=3)
priorities = [
    ("1. AI Reply Agent is the #1 priority.", "This is Instantly.ai's most touted feature and our biggest functional gap. Building HITL mode first (Phase 2.1) then Autopilot (Phase 2.2) gives us parity within 8 weeks."),
    ("2. Multichannel is the market direction.", "Cold email alone is saturating. Adding SMS (we have Twilio), LinkedIn, and Click-to-Call puts us ahead of Instantly which is primarily email-focused."),
    ("3. Website Visitor Tracking is a premium differentiator.", "Instantly offers this on their CRM plans. Building it gives us a competitive moat, especially for B2B SaaS prospects who visit websites before buying."),
    ("4. UI/UX is the perception battleground.", "Users choose tools that FEEL powerful. The Copilot chat widget, command palette, glassmorphism design, and visual sequence builder will create a perception of premium quality."),
    ("5. Don't copy — innovate beyond.", "NeuraLeads should offer things Instantly.ai doesn't: multi-provider AI flexibility, enterprise-grade multi-tenancy, built-in invoicing, and a visual pipeline that Instantly lacks entirely."),
]
for title, desc in priorities:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + " ")
    run.font.bold = True
    run.font.size = Pt(10.5)
    p.add_run(desc).font.size = Pt(10)

doc.add_heading("Bottom Line", level=3)
p = doc.add_paragraph()
run = p.add_run(
    "With focused execution of this 24-week roadmap, NeuraLeads will not just match Instantly.ai — "
    "it will surpass it in feature depth, AI capabilities, data infrastructure, and user experience. "
    "The competitive advantage of owning the full stack (self-hosted, multi-tenant, multi-provider) "
    "means we can move faster, customize deeper, and deliver value that no SaaS-only competitor can replicate."
)
run.font.size = Pt(11)
run.font.italic = True

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Document —")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Confidential — Exzelon Technologies — April 2026")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, "NeuraLeads_vs_Instantly_Competitive_Analysis.docx")
doc.save(output_path)
print(f"Document saved: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
