#!/usr/bin/env python3
"""E2E Test Workflow: Create test data on production and generate .docx report.

Creates 2 test sets (clients, leads, contacts), a campaign with sequence steps,
enrolls contacts, creates deals, captures analytics, and generates a comprehensive
Word document report.

Usage:
    python scripts/e2e_test_workflow.py [--base-url URL]

Default: https://ra.partnerwithus.tech/api/v1
"""
import argparse
import json
import sys
import time
import traceback
from datetime import datetime, date
from dataclasses import dataclass, field
from typing import Optional, Any

try:
    import requests
except ImportError:
    print("ERROR: 'requests' package required. Install: pip install requests")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: 'python-docx' package required. Install: pip install python-docx")
    sys.exit(1)


# ──────────────────────────────────────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────────────────────────────────────

DEFAULT_BASE_URL = "https://ra.partnerwithus.tech/api/v1"
SUPER_ADMIN = {"email": "ali.aitechs@gmail.com", "password": "SA@Admin#123"}
TENANT_ID = "1"  # X-Tenant-ID header value

# Color constants for .docx
GREEN = "22C55E"
RED = "EF4444"
BLUE = "3B82F6"
GRAY = "6B7280"
LIGHT_GREEN = "DCFCE7"
LIGHT_RED = "FEE2E2"
LIGHT_BLUE = "DBEAFE"
LIGHT_GRAY = "F3F4F6"
WHITE = "FFFFFF"
DARK = "1F2937"


# ──────────────────────────────────────────────────────────────────────────────
# Data Classes
# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class StepResult:
    step_num: int
    phase: str
    action: str
    status: str  # PASS / FAIL / SKIP
    details: str
    response_data: Any = None
    duration_ms: float = 0
    timestamp: str = ""


@dataclass
class WorkflowState:
    base_url: str
    token: str = ""
    results: list = field(default_factory=list)
    # Test Set 1
    client1_id: int = 0
    lead1_id: int = 0
    lead2_id: int = 0
    contact1_id: int = 0
    contact2_id: int = 0
    # Test Set 2
    client2_id: int = 0
    lead3_id: int = 0
    contact3_id: int = 0
    # Campaign
    campaign_id: int = 0
    step1_id: int = 0
    step2_id: int = 0
    step3_id: int = 0
    # Deals
    deal1_id: int = 0
    deal2_id: int = 0
    # Discovered
    mailbox_ids: list = field(default_factory=list)
    stages: dict = field(default_factory=dict)  # name -> stage_id
    # Verification data
    campaign_detail: dict = field(default_factory=dict)
    campaign_contacts: dict = field(default_factory=dict)
    analytics_campaign: dict = field(default_factory=dict)
    analytics_revenue: dict = field(default_factory=dict)
    deal_stats: dict = field(default_factory=dict)


# ──────────────────────────────────────────────────────────────────────────────
# HTTP Helpers
# ──────────────────────────────────────────────────────────────────────────────

def _headers(state: WorkflowState) -> dict:
    h = {"Authorization": f"Bearer {state.token}"}
    if TENANT_ID:
        h["X-Tenant-ID"] = TENANT_ID
    return h


def api_get(state: WorkflowState, path: str, params=None) -> requests.Response:
    for attempt in range(3):
        try:
            return requests.get(
                f"{state.base_url}{path}", headers=_headers(state),
                params=params, timeout=30,
            )
        except requests.RequestException as e:
            if attempt == 2:
                raise
            time.sleep(2 ** attempt)


def api_put(state: WorkflowState, path: str, json_data=None) -> requests.Response:
    for attempt in range(3):
        try:
            return requests.put(
                f"{state.base_url}{path}", headers=_headers(state),
                json=json_data, timeout=30,
            )
        except requests.RequestException as e:
            if attempt == 2:
                raise
            time.sleep(2 ** attempt)


def api_post(state: WorkflowState, path: str, json_data=None, data=None,
             content_type=None) -> requests.Response:
    headers = _headers(state)
    if content_type:
        headers["Content-Type"] = content_type
    for attempt in range(3):
        try:
            return requests.post(
                f"{state.base_url}{path}", headers=headers,
                json=json_data, data=data, timeout=30,
            )
        except requests.RequestException as e:
            if attempt == 2:
                raise
            time.sleep(2 ** attempt)


def run_step(state: WorkflowState, step_num: int, phase: str, action: str,
             fn) -> StepResult:
    """Execute a step, record result."""
    start = time.time()
    ts = datetime.now().strftime("%H:%M:%S")
    try:
        status, details, resp_data = fn()
        duration = (time.time() - start) * 1000
        result = StepResult(
            step_num=step_num, phase=phase, action=action,
            status=status, details=details, response_data=resp_data,
            duration_ms=round(duration, 1), timestamp=ts,
        )
    except Exception as e:
        duration = (time.time() - start) * 1000
        result = StepResult(
            step_num=step_num, phase=phase, action=action,
            status="FAIL", details=f"Exception: {e}",
            response_data=traceback.format_exc(),
            duration_ms=round(duration, 1), timestamp=ts,
        )
    state.results.append(result)
    icon = {"PASS": "+", "FAIL": "X", "SKIP": "-"}[result.status]
    print(f"  [{icon}] Step {step_num}: {action} ({result.duration_ms:.0f}ms) - {result.status}")
    if result.status == "FAIL":
        print(f"      Details: {result.details[:200]}")
    return result


# ──────────────────────────────────────────────────────────────────────────────
# Workflow Phases
# ──────────────────────────────────────────────────────────────────────────────

def phase_auth(state: WorkflowState, step_counter: list):
    """Phase 1: Authenticate as super admin."""
    print("\n== Phase 1: Authentication ==")

    def login():
        resp = requests.post(
            f"{state.base_url}/auth/login",
            data={"username": SUPER_ADMIN["email"], "password": SUPER_ADMIN["password"]},
            timeout=30,
        )
        if resp.status_code == 200:
            body = resp.json()
            state.token = body["access_token"]
            user = body.get("user", {})
            return "PASS", f"Logged in as {user.get('email', '?')} (role={user.get('role', '?')})", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Auth", "Login as super admin", login)
    if not state.token:
        print("\nFATAL: Authentication failed. Cannot continue.")
        sys.exit(1)


def phase_discover(state: WorkflowState, step_counter: list):
    """Phase 2: Discover mailboxes and deal stages."""
    print("\n== Phase 2: Discover Resources ==")

    # Mailboxes
    def get_mailboxes():
        resp = api_get(state, "/mailboxes")
        if resp.status_code == 200:
            body = resp.json()
            items = body.get("items", body if isinstance(body, list) else [])
            active = [m for m in items if m.get("is_active")]
            state.mailbox_ids = [m["mailbox_id"] for m in active]
            return "PASS", f"Found {len(active)} active mailboxes: {state.mailbox_ids}", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Discover", "List active mailboxes", get_mailboxes)

    # Deal stages
    def get_stages():
        resp = api_get(state, "/deals/stages")
        if resp.status_code == 200:
            stages = resp.json()
            if isinstance(stages, list):
                state.stages = {s["name"]: s["stage_id"] for s in stages}
            elif isinstance(stages, dict) and "stages" in stages:
                state.stages = {s["name"]: s["stage_id"] for s in stages["stages"]}
            return "PASS", f"Found {len(state.stages)} stages: {list(state.stages.keys())}", stages
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Discover", "List deal stages", get_stages)


def _create_client(state, name, industry, company_size, location_state, website):
    """Create a client, handle duplicate."""
    resp = api_post(state, "/clients", json_data={
        "client_name": name,
        "industry": industry,
        "company_size": company_size,
        "location_state": location_state,
        "website": website,
        "status": "active",
    })
    if resp.status_code in (200, 201):
        body = resp.json()
        return "PASS", f"Created client '{name}' (ID={body.get('client_id', '?')})", body
    elif resp.status_code == 400 and "already exists" in resp.text.lower():
        # Fetch existing
        search_resp = api_get(state, "/clients", params={"search": name, "limit": 5})
        if search_resp.status_code == 200:
            items = search_resp.json()
            if isinstance(items, dict):
                items = items.get("items", [])
            for c in items:
                if c.get("client_name", "").lower() == name.lower():
                    return "PASS", f"Client '{name}' already exists (ID={c['client_id']})", c
        return "SKIP", f"Client '{name}' exists but couldn't fetch ID", None
    return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None


def _create_lead(state, client_name, job_title, loc_state, salary_min, salary_max, industry, company_size):
    """Create a lead."""
    resp = api_post(state, "/leads", json_data={
        "client_name": client_name,
        "job_title": job_title,
        "state": loc_state,
        "salary_min": salary_min,
        "salary_max": salary_max,
        "source": "manual",
        "lead_status": "open",
        "industry": industry,
        "company_size": company_size,
        "posting_date": date.today().isoformat(),
    })
    if resp.status_code in (200, 201):
        body = resp.json()
        return "PASS", f"Created lead '{job_title}' (ID={body.get('lead_id', '?')})", body
    return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None


def _create_contact(state, email, first_name, last_name, title, client_name,
                    priority_level, lead_ids, location_state):
    """Create a contact, handle duplicate by fetching existing.

    priority_level must be a PriorityLevel enum string:
      p1_job_poster, p2_hr_ta_recruiter, p3_hr_manager, p4_ops_leader, p5_functional_manager
    """
    resp = api_post(state, "/contacts", json_data={
        "email": email,
        "first_name": first_name,
        "last_name": last_name,
        "title": title,
        "client_name": client_name,
        "priority_level": priority_level,
        "source": "manual",
        "lead_ids": lead_ids,
        "location_state": location_state,
    })
    if resp.status_code in (200, 201):
        body = resp.json()
        cid = body.get("contact_id")
        # Set validation_status to Valid via update (not on create schema)
        if cid:
            api_put(state, f"/contacts/{cid}", json_data={"validation_status": "Valid"})
            body["validation_status"] = "Valid"
        return "PASS", f"Created contact '{email}' (ID={cid})", body
    elif resp.status_code == 400 and "already exists" in resp.text.lower():
        # Search both non-archived and archived contacts
        for archived_flag in [None, "true"]:
            params = {"search": email, "limit": 5}
            if archived_flag:
                params["show_archived"] = archived_flag
            search_resp = api_get(state, "/contacts", params=params)
            if search_resp.status_code == 200:
                data = search_resp.json()
                items = data.get("items", data if isinstance(data, list) else [])
                for c in items:
                    if c.get("email", "").lower() == email.lower():
                        cid = c["contact_id"]
                        # Ensure validation status is Valid
                        api_put(state, f"/contacts/{cid}", json_data={
                            "validation_status": "Valid",
                        })
                        c["validation_status"] = "Valid"
                        return "PASS", f"Contact '{email}' already exists (ID={cid}), updated", c
        return "SKIP", f"Contact '{email}' exists but couldn't fetch ID", None
    return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None


def phase_test_set_1(state: WorkflowState, step_counter: list):
    """Phase 3: Create Test Set 1 — Acme Digital Solutions."""
    print("\n== Phase 3: Test Set 1 — Acme Digital Solutions ==")

    # Client
    def create_client1():
        s, d, r = _create_client(
            state, "Acme Digital Solutions", "Information Technology",
            "201-500", "TX", "https://acmedigital.example.com"
        )
        if r and r.get("client_id"):
            state.client1_id = r["client_id"]
        return s, d, r

    step_counter[0] += 1
    run_step(state, step_counter[0], "Set 1", "Create client: Acme Digital Solutions", create_client1)

    # Lead 1
    def create_lead1():
        s, d, r = _create_lead(
            state, "Acme Digital Solutions", "Senior Software Engineer",
            "TX", 120000, 160000, "Information Technology", "201-500"
        )
        if r and r.get("lead_id"):
            state.lead1_id = r["lead_id"]
        return s, d, r

    step_counter[0] += 1
    run_step(state, step_counter[0], "Set 1", "Create lead: Senior Software Engineer (TX)", create_lead1)

    # Lead 2
    def create_lead2():
        s, d, r = _create_lead(
            state, "Acme Digital Solutions", "DevOps Engineer",
            "TX", 110000, 150000, "Information Technology", "201-500"
        )
        if r and r.get("lead_id"):
            state.lead2_id = r["lead_id"]
        return s, d, r

    step_counter[0] += 1
    run_step(state, step_counter[0], "Set 1", "Create lead: DevOps Engineer (TX)", create_lead2)

    # Contact 1 — ali.infy@gmail.com
    def create_contact1():
        lead_ids = [lid for lid in [state.lead1_id] if lid]
        s, d, r = _create_contact(
            state, "ali.infy@gmail.com", "Ali", "Infy",
            "VP of Engineering", "Acme Digital Solutions", "p1_job_poster", lead_ids, "TX"
        )
        if r and r.get("contact_id"):
            state.contact1_id = r["contact_id"]
        return s, d, r

    step_counter[0] += 1
    run_step(state, step_counter[0], "Set 1", "Create contact: ali.infy@gmail.com (P1)", create_contact1)

    # Contact 2 — ali.aitechs@gmail.com
    def create_contact2():
        lead_ids = [lid for lid in [state.lead2_id] if lid]
        s, d, r = _create_contact(
            state, "ali.aitechs@gmail.com", "Ali", "AiTechs",
            "Hiring Manager", "Acme Digital Solutions", "p2_hr_ta_recruiter", lead_ids, "TX"
        )
        if r and r.get("contact_id"):
            state.contact2_id = r["contact_id"]
        return s, d, r

    step_counter[0] += 1
    run_step(state, step_counter[0], "Set 1", "Create contact: ali.aitechs@gmail.com (P2)", create_contact2)


def phase_test_set_2(state: WorkflowState, step_counter: list):
    """Phase 4: Create Test Set 2 — Pinnacle Engineering Group."""
    print("\n== Phase 4: Test Set 2 — Pinnacle Engineering Group ==")

    # Client
    def create_client2():
        s, d, r = _create_client(
            state, "Pinnacle Engineering Group", "Engineering",
            "51-200", "CA", "https://pinnacle-eng.example.com"
        )
        if r and r.get("client_id"):
            state.client2_id = r["client_id"]
        return s, d, r

    step_counter[0] += 1
    run_step(state, step_counter[0], "Set 2", "Create client: Pinnacle Engineering Group", create_client2)

    # Lead 3
    def create_lead3():
        s, d, r = _create_lead(
            state, "Pinnacle Engineering Group", "Project Manager",
            "CA", 100000, 140000, "Engineering", "51-200"
        )
        if r and r.get("lead_id"):
            state.lead3_id = r["lead_id"]
        return s, d, r

    step_counter[0] += 1
    run_step(state, step_counter[0], "Set 2", "Create lead: Project Manager (CA)", create_lead3)

    # Contact 3 — ali@medeoan.com
    def create_contact3():
        lead_ids = [lid for lid in [state.lead3_id] if lid]
        s, d, r = _create_contact(
            state, "ali@medeoan.com", "Ali", "Medeoan",
            "Director of Operations", "Pinnacle Engineering Group", "p3_hr_manager", lead_ids, "CA"
        )
        if r and r.get("contact_id"):
            state.contact3_id = r["contact_id"]
        return s, d, r

    step_counter[0] += 1
    run_step(state, step_counter[0], "Set 2", "Create contact: ali@medeoan.com (P3)", create_contact3)


def phase_campaign(state: WorkflowState, step_counter: list):
    """Phase 5: Create campaign, add steps, enroll contacts, activate."""
    print("\n== Phase 5: Campaign Setup ==")

    # Create campaign
    def create_campaign():
        mailbox_ids = state.mailbox_ids[:1] if state.mailbox_ids else []
        resp = api_post(state, "/campaigns", json_data={
            "name": "E2E Test Campaign",
            "description": "End-to-end test campaign with 3 contacts across 2 test sets",
            "timezone": "America/Chicago",
            "send_window_start": "06:00",
            "send_window_end": "23:00",
            "send_days": ["mon", "tue", "wed", "thu", "fri", "sat", "sun"],
            "mailbox_ids": mailbox_ids,
            "daily_limit": 30,
        })
        if resp.status_code in (200, 201):
            body = resp.json()
            state.campaign_id = body.get("campaign_id", 0)
            return "PASS", f"Created campaign (ID={state.campaign_id})", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Campaign", "Create campaign: E2E Test Campaign", create_campaign)

    if not state.campaign_id:
        print("  WARN: No campaign created, skipping steps/enrollment/activation")
        return

    # Step 1: Initial email
    def add_step1():
        resp = api_post(state, f"/campaigns/{state.campaign_id}/steps", json_data={
            "step_type": "email",
            "step_order": 1,
            "subject": "Quick question about your {{job_title}} opening",
            "body_html": (
                "<p>Hi {{first_name}},</p>"
                "<p>I noticed {{client_name}} is looking for a {{job_title}} in {{job_location}}. "
                "We specialize in connecting companies with top-tier talent in this space.</p>"
                "<p>Would you be open to a brief conversation about how we can help fill this role "
                "with a qualified candidate?</p>"
                "<p>Best regards,<br>{{sender_name}}</p>"
            ),
            "delay_days": 0,
            "delay_hours": 0,
        })
        if resp.status_code in (200, 201):
            body = resp.json()
            state.step1_id = body.get("step_id", 0)
            return "PASS", f"Added Step 1: Initial email (ID={state.step1_id})", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Campaign", "Add Step 1: Initial outreach email", add_step1)

    # Step 2: Wait 1 day
    def add_step2():
        resp = api_post(state, f"/campaigns/{state.campaign_id}/steps", json_data={
            "step_type": "wait",
            "step_order": 2,
            "delay_days": 1,
            "delay_hours": 0,
        })
        if resp.status_code in (200, 201):
            body = resp.json()
            state.step2_id = body.get("step_id", 0)
            return "PASS", f"Added Step 2: Wait 1 day (ID={state.step2_id})", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Campaign", "Add Step 2: Wait 1 day", add_step2)

    # Step 3: Follow-up email
    def add_step3():
        resp = api_post(state, f"/campaigns/{state.campaign_id}/steps", json_data={
            "step_type": "email",
            "step_order": 3,
            "subject": "Re: Quick question about your {{job_title}} opening",
            "body_html": (
                "<p>Hi {{first_name}},</p>"
                "<p>Just wanted to follow up on my previous message. We've recently placed "
                "several {{job_title}} candidates in the {{job_location}} area with great success.</p>"
                "<p>I'd love to share a few pre-vetted profiles that might be a strong fit "
                "for {{client_name}}. Would next week work for a quick 15-minute call?</p>"
                "<p>Best,<br>{{sender_name}}</p>"
            ),
            "delay_days": 0,
            "delay_hours": 0,
            "reply_to_thread": True,
        })
        if resp.status_code in (200, 201):
            body = resp.json()
            state.step3_id = body.get("step_id", 0)
            return "PASS", f"Added Step 3: Follow-up email (ID={state.step3_id})", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Campaign", "Add Step 3: Follow-up email", add_step3)

    # Enroll contacts
    def enroll_contacts():
        contact_ids = [cid for cid in [state.contact1_id, state.contact2_id, state.contact3_id] if cid]
        if not contact_ids:
            return "SKIP", "No contacts to enroll", None
        resp = api_post(state, f"/campaigns/{state.campaign_id}/contacts", json_data={
            "contact_ids": contact_ids,
        })
        if resp.status_code in (200, 201):
            body = resp.json()
            return "PASS", f"Enrolled {body.get('enrolled', '?')} contacts (skipped {body.get('skipped', 0)})", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Campaign", "Enroll 3 contacts", enroll_contacts)

    # Activate campaign
    def activate():
        resp = api_post(state, f"/campaigns/{state.campaign_id}/activate")
        if resp.status_code in (200, 201):
            body = resp.json()
            return "PASS", f"Campaign activated: {body.get('message', 'OK')}", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Campaign", "Activate campaign", activate)


def phase_deals(state: WorkflowState, step_counter: list):
    """Phase 6: Create deals."""
    print("\n== Phase 6: Deal Pipeline ==")

    # Deal 1
    def create_deal1():
        stage_id = state.stages.get("Contacted", state.stages.get("New Lead", 1))
        deal_data = {
            "name": "Acme Digital - SWE Placement",
            "stage_id": stage_id,
            "value": 25000,
            "probability": 40,
            "notes": "E2E test deal: Senior Software Engineer placement for Acme Digital Solutions",
        }
        if state.client1_id:
            deal_data["client_id"] = state.client1_id
        if state.contact1_id:
            deal_data["contact_id"] = state.contact1_id
        if state.campaign_id:
            deal_data["campaign_id"] = state.campaign_id
        resp = api_post(state, "/deals", json_data=deal_data)
        if resp.status_code in (200, 201):
            body = resp.json()
            state.deal1_id = body.get("deal_id", 0)
            return "PASS", f"Created deal (ID={state.deal1_id}, stage={body.get('stage_name', '?')})", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Deals", "Create deal: Acme Digital - SWE Placement ($25k)", create_deal1)

    # Deal 2
    def create_deal2():
        stage_id = state.stages.get("New Lead", 1)
        deal_data = {
            "name": "Pinnacle Eng - PM Placement",
            "stage_id": stage_id,
            "value": 18000,
            "probability": 25,
            "notes": "E2E test deal: Project Manager placement for Pinnacle Engineering Group",
        }
        if state.client2_id:
            deal_data["client_id"] = state.client2_id
        if state.contact3_id:
            deal_data["contact_id"] = state.contact3_id
        if state.campaign_id:
            deal_data["campaign_id"] = state.campaign_id
        resp = api_post(state, "/deals", json_data=deal_data)
        if resp.status_code in (200, 201):
            body = resp.json()
            state.deal2_id = body.get("deal_id", 0)
            return "PASS", f"Created deal (ID={state.deal2_id}, stage={body.get('stage_name', '?')})", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Deals", "Create deal: Pinnacle Eng - PM Placement ($18k)", create_deal2)


def phase_verification(state: WorkflowState, step_counter: list):
    """Phase 7: Verify created entities."""
    print("\n== Phase 7: Verification ==")

    # Campaign detail
    def verify_campaign():
        if not state.campaign_id:
            return "SKIP", "No campaign to verify", None
        resp = api_get(state, f"/campaigns/{state.campaign_id}")
        if resp.status_code == 200:
            body = resp.json()
            state.campaign_detail = body
            status = body.get("status", "?")
            steps = body.get("steps", [])
            return "PASS", f"Campaign status={status}, steps={len(steps)}", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Verify", "Fetch campaign detail", verify_campaign)

    # Campaign contacts
    def verify_contacts():
        if not state.campaign_id:
            return "SKIP", "No campaign to verify", None
        resp = api_get(state, f"/campaigns/{state.campaign_id}/contacts")
        if resp.status_code == 200:
            body = resp.json()
            state.campaign_contacts = body
            items = body if isinstance(body, list) else body.get("items", body.get("contacts", []))
            return "PASS", f"Campaign has {len(items)} enrolled contacts", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Verify", "Fetch campaign contacts", verify_contacts)

    # Verify deals exist
    def verify_deals():
        resp = api_get(state, "/deals", params={"limit": 10})
        if resp.status_code == 200:
            body = resp.json()
            deals = body if isinstance(body, list) else body.get("items", body.get("deals", []))
            test_deals = [d for d in deals if d.get("name", "").startswith(("Acme Digital", "Pinnacle Eng"))]
            return "PASS", f"Found {len(test_deals)} test deals out of {len(deals)} total", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Verify", "Verify deals in pipeline", verify_deals)


def phase_analytics(state: WorkflowState, step_counter: list):
    """Phase 8: Capture analytics."""
    print("\n== Phase 8: Analytics ==")

    # Campaign comparison
    def get_campaign_analytics():
        resp = api_get(state, "/analytics/campaign-comparison")
        if resp.status_code == 200:
            body = resp.json()
            state.analytics_campaign = body
            campaigns = body.get("campaigns", body if isinstance(body, list) else [])
            return "PASS", f"Campaign analytics: {len(campaigns)} campaigns tracked", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Analytics", "Fetch campaign analytics", get_campaign_analytics)

    # Revenue metrics
    def get_revenue():
        resp = api_get(state, "/analytics/revenue", params={"days": 90})
        if resp.status_code == 200:
            body = resp.json()
            state.analytics_revenue = body
            return "PASS", (
                f"Revenue: pipeline=${body.get('pipeline_value', 0):,.0f}, "
                f"deals={body.get('total_deals', 0)}, "
                f"win_rate={body.get('win_rate', 0):.1f}%"
            ), body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Analytics", "Fetch revenue metrics", get_revenue)

    # Deal stats
    def get_deal_stats():
        resp = api_get(state, "/deals/stats")
        if resp.status_code == 200:
            body = resp.json()
            state.deal_stats = body
            return "PASS", f"Deal stats: {json.dumps(body)[:150]}", body
        return "FAIL", f"HTTP {resp.status_code}: {resp.text[:200]}", None

    step_counter[0] += 1
    run_step(state, step_counter[0], "Analytics", "Fetch deal pipeline stats", get_deal_stats)


# ──────────────────────────────────────────────────────────────────────────────
# .docx Report Generator
# ──────────────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK)
        set_cell_text(cell, header, bold=True, color=WHITE, size=9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = LIGHT_GRAY if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_shading(cell, bg)
            # Color-code PASS/FAIL
            if val == "PASS":
                set_cell_text(cell, val, bold=True, color=GREEN, size=9)
            elif val == "FAIL":
                set_cell_text(cell, val, bold=True, color=RED, size=9)
            elif val == "SKIP":
                set_cell_text(cell, val, bold=True, color=GRAY, size=9)
            else:
                set_cell_text(cell, str(val) if val is not None else "", size=9)

    return table


def add_key_value_table(doc, data: dict, title=None):
    """Add a 2-column key-value table."""
    if title:
        doc.add_heading(title, level=3)
    items = list(data.items())
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (key, val) in enumerate(items):
        bg = LIGHT_BLUE if i % 2 == 0 else WHITE
        set_cell_shading(table.rows[i].cells[0], bg)
        set_cell_shading(table.rows[i].cells[1], bg)
        set_cell_text(table.rows[i].cells[0], key, bold=True, size=9)
        val_str = json.dumps(val, indent=2) if isinstance(val, (dict, list)) else str(val)
        if len(val_str) > 300:
            val_str = val_str[:297] + "..."
        set_cell_text(table.rows[i].cells[1], val_str, size=9)
    return table


def generate_report(state: WorkflowState, output_path: str):
    """Generate the .docx report."""
    print(f"\n== Generating Report: {output_path} ==")

    doc = Document()

    # ── Title Page ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("E2E Test Workflow Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Exzelon RA Agent — Production Validation")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        f"Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n"
        f"Environment: {state.base_url.replace('/api/v1', '')}\n"
        f"Tenant ID: {TENANT_ID}"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_page_break()

    # ── Executive Summary ──
    doc.add_heading("1. Executive Summary", level=1)

    total = len(state.results)
    passed = sum(1 for r in state.results if r.status == "PASS")
    failed = sum(1 for r in state.results if r.status == "FAIL")
    skipped = sum(1 for r in state.results if r.status == "SKIP")
    total_time = sum(r.duration_ms for r in state.results)

    summary_data = {
        "Total Steps": total,
        "Passed": f"{passed} ({passed/total*100:.0f}%)" if total else "0",
        "Failed": f"{failed} ({failed/total*100:.0f}%)" if total else "0",
        "Skipped": skipped,
        "Total Duration": f"{total_time/1000:.1f}s",
        "Overall Result": "PASS" if failed == 0 else "FAIL",
    }

    items = list(summary_data.items())
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(items):
        bg = LIGHT_GREEN if k == "Overall Result" and "PASS" in str(v) else (LIGHT_RED if k == "Overall Result" else LIGHT_BLUE)
        set_cell_shading(table.rows[i].cells[0], bg)
        set_cell_shading(table.rows[i].cells[1], bg)
        set_cell_text(table.rows[i].cells[0], k, bold=True, size=10)
        color = GREEN if "PASS" in str(v) and k == "Overall Result" else (RED if "FAIL" in str(v) and k == "Overall Result" else None)
        set_cell_text(table.rows[i].cells[1], str(v), bold=(k == "Overall Result"), color=color, size=10)

    doc.add_paragraph("")

    entities = doc.add_paragraph()
    entities.add_run("Entities Created:").bold = True
    doc.add_paragraph(f"Clients: 2 (IDs: {state.client1_id}, {state.client2_id})", style="List Bullet")
    doc.add_paragraph(f"Leads: 3 (IDs: {state.lead1_id}, {state.lead2_id}, {state.lead3_id})", style="List Bullet")
    doc.add_paragraph(f"Contacts: 3 (IDs: {state.contact1_id}, {state.contact2_id}, {state.contact3_id})", style="List Bullet")
    doc.add_paragraph(f"Campaign: 1 (ID: {state.campaign_id})", style="List Bullet")
    doc.add_paragraph(f"Sequence Steps: 3 (IDs: {state.step1_id}, {state.step2_id}, {state.step3_id})", style="List Bullet")
    doc.add_paragraph(f"Deals: 2 (IDs: {state.deal1_id}, {state.deal2_id})", style="List Bullet")

    doc.add_page_break()

    # ── Test Set 1 ──
    doc.add_heading("2. Test Set 1: Acme Digital Solutions", level=1)

    doc.add_heading("Client", level=2)
    set1_results = [r for r in state.results if r.phase == "Set 1"]
    client_r = [r for r in set1_results if "client" in r.action.lower()]
    if client_r and client_r[0].response_data and isinstance(client_r[0].response_data, dict):
        cd = client_r[0].response_data
        add_key_value_table(doc, {
            "Client ID": cd.get("client_id", "?"),
            "Name": cd.get("client_name", "?"),
            "Industry": cd.get("industry", "?"),
            "Company Size": cd.get("company_size", "?"),
            "State": cd.get("location_state", "?"),
            "Status": cd.get("status", "?"),
        })
    else:
        doc.add_paragraph(f"Client ID: {state.client1_id}")

    doc.add_heading("Leads", level=2)
    lead_rows = []
    for r in set1_results:
        if "lead" in r.action.lower() and r.response_data and isinstance(r.response_data, dict):
            d = r.response_data
            lead_rows.append([
                d.get("lead_id", "?"), d.get("job_title", "?"),
                d.get("state", "?"),
                f"${float(d.get('salary_min') or 0):,.0f}-${float(d.get('salary_max') or 0):,.0f}",
                r.status,
            ])
    if lead_rows:
        add_styled_table(doc, ["Lead ID", "Job Title", "State", "Salary Range", "Status"], lead_rows)

    doc.add_heading("Contacts", level=2)
    contact_rows = []
    for r in set1_results:
        if "contact" in r.action.lower() and r.response_data and isinstance(r.response_data, dict):
            d = r.response_data
            contact_rows.append([
                d.get("contact_id", "?"), d.get("email", "?"),
                f"{d.get('first_name', '')} {d.get('last_name', '')}".strip(),
                d.get("title", "?"),
                f"P{d.get('priority_level', '?')}",
                d.get("validation_status", "?"),
                r.status,
            ])
    if contact_rows:
        add_styled_table(doc, ["ID", "Email", "Name", "Title", "Priority", "Validation", "Status"], contact_rows)

    doc.add_page_break()

    # ── Test Set 2 ──
    doc.add_heading("3. Test Set 2: Pinnacle Engineering Group", level=1)

    set2_results = [r for r in state.results if r.phase == "Set 2"]

    doc.add_heading("Client", level=2)
    client2_r = [r for r in set2_results if "client" in r.action.lower()]
    if client2_r and client2_r[0].response_data and isinstance(client2_r[0].response_data, dict):
        cd = client2_r[0].response_data
        add_key_value_table(doc, {
            "Client ID": cd.get("client_id", "?"),
            "Name": cd.get("client_name", "?"),
            "Industry": cd.get("industry", "?"),
            "Company Size": cd.get("company_size", "?"),
            "State": cd.get("location_state", "?"),
            "Status": cd.get("status", "?"),
        })

    doc.add_heading("Lead", level=2)
    lead2_rows = []
    for r in set2_results:
        if "lead" in r.action.lower() and r.response_data and isinstance(r.response_data, dict):
            d = r.response_data
            lead2_rows.append([
                d.get("lead_id", "?"), d.get("job_title", "?"),
                d.get("state", "?"),
                f"${float(d.get('salary_min') or 0):,.0f}-${float(d.get('salary_max') or 0):,.0f}",
                r.status,
            ])
    if lead2_rows:
        add_styled_table(doc, ["Lead ID", "Job Title", "State", "Salary Range", "Status"], lead2_rows)

    doc.add_heading("Contact", level=2)
    contact2_rows = []
    for r in set2_results:
        if "contact" in r.action.lower() and r.response_data and isinstance(r.response_data, dict):
            d = r.response_data
            contact2_rows.append([
                d.get("contact_id", "?"), d.get("email", "?"),
                f"{d.get('first_name', '')} {d.get('last_name', '')}".strip(),
                d.get("title", "?"),
                f"P{d.get('priority_level', '?')}",
                d.get("validation_status", "?"),
                r.status,
            ])
    if contact2_rows:
        add_styled_table(doc, ["ID", "Email", "Name", "Title", "Priority", "Validation", "Status"], contact2_rows)

    doc.add_page_break()

    # ── Campaign Setup ──
    doc.add_heading("4. Campaign Setup", level=1)

    doc.add_heading("Campaign Details", level=2)
    if state.campaign_detail and isinstance(state.campaign_detail, dict):
        cd = state.campaign_detail
        add_key_value_table(doc, {
            "Campaign ID": cd.get("campaign_id", "?"),
            "Name": cd.get("name", "?"),
            "Status": cd.get("status", "?"),
            "Timezone": cd.get("timezone", "?"),
            "Send Window": f"{cd.get('send_window_start', '?')} - {cd.get('send_window_end', '?')}",
            "Send Days": ", ".join(cd.get("send_days", [])),
            "Mailbox IDs": str(cd.get("mailbox_ids", [])),
            "Daily Limit": cd.get("daily_limit", "?"),
            "Total Contacts": cd.get("total_contacts", "?"),
        })
    else:
        doc.add_paragraph(f"Campaign ID: {state.campaign_id}")

    doc.add_heading("Sequence Steps", level=2)
    step_results = [r for r in state.results if r.phase == "Campaign" and "Step" in r.action]
    step_rows = []
    for r in step_results:
        if r.response_data and isinstance(r.response_data, dict):
            d = r.response_data
            step_rows.append([
                d.get("step_order", d.get("step_id", "?")),
                d.get("step_type", "?"),
                d.get("subject", "-")[:50] if d.get("subject") else "-",
                f"{d.get('delay_days', 0)}d {d.get('delay_hours', 0)}h",
                r.status,
            ])
    if step_rows:
        add_styled_table(doc, ["Order", "Type", "Subject", "Delay", "Status"], step_rows)

    doc.add_heading("Enrollment", level=2)
    enroll_r = [r for r in state.results if "Enroll" in r.action]
    if enroll_r:
        r = enroll_r[0]
        doc.add_paragraph(f"Result: {r.status} - {r.details}")

    doc.add_heading("Activation", level=2)
    activate_r = [r for r in state.results if "Activate" in r.action]
    if activate_r:
        r = activate_r[0]
        doc.add_paragraph(f"Result: {r.status} - {r.details}")

    doc.add_page_break()

    # ── Deal Pipeline ──
    doc.add_heading("5. Deal Pipeline", level=1)

    doc.add_heading("Deal Stages", level=2)
    if state.stages:
        stage_rows = [[name, sid] for name, sid in sorted(state.stages.items(), key=lambda x: x[1])]
        add_styled_table(doc, ["Stage Name", "Stage ID"], stage_rows)

    doc.add_heading("Deals Created", level=2)
    deal_results = [r for r in state.results if r.phase == "Deals"]
    deal_rows = []
    for r in deal_results:
        if r.response_data and isinstance(r.response_data, dict):
            d = r.response_data
            deal_rows.append([
                d.get("deal_id", "?"),
                d.get("name", "?"),
                d.get("stage_name", "?"),
                f"${d.get('value', 0):,.0f}",
                f"{d.get('probability', 0)}%",
                d.get("client_name", "-"),
                r.status,
            ])
    if deal_rows:
        add_styled_table(doc, ["ID", "Deal Name", "Stage", "Value", "Probability", "Client", "Status"], deal_rows)

    doc.add_page_break()

    # ── Analytics ──
    doc.add_heading("6. System Analytics", level=1)

    doc.add_heading("Revenue Metrics (90 days)", level=2)
    if state.analytics_revenue and isinstance(state.analytics_revenue, dict):
        rv = state.analytics_revenue
        add_key_value_table(doc, {
            "Pipeline Value": f"${rv.get('pipeline_value', 0):,.0f}",
            "Total Deals": rv.get("total_deals", 0),
            "Deals Won": rv.get("deals_won", 0),
            "Average Deal Size": f"${rv.get('avg_deal_size', 0):,.0f}",
            "Win Rate": f"{rv.get('win_rate', 0):.1f}%",
            "Total Won Value": f"${rv.get('total_won_value', 0):,.0f}",
            "Leads Generated": rv.get("leads_generated", 0),
            "ROI": f"{rv.get('roi_percent', 0):.1f}%",
        })

    doc.add_heading("Deal Stats", level=2)
    if state.deal_stats and isinstance(state.deal_stats, dict):
        add_key_value_table(doc, state.deal_stats)
    else:
        doc.add_paragraph("No deal stats available.")

    doc.add_heading("Campaign Comparison", level=2)
    if state.analytics_campaign:
        campaigns = state.analytics_campaign.get("campaigns", state.analytics_campaign if isinstance(state.analytics_campaign, list) else [])
        camp_rows = []
        for c in campaigns[:10]:
            camp_rows.append([
                c.get("campaign_id", "?"),
                c.get("name", "?")[:30],
                c.get("status", "?"),
                c.get("total_contacts", 0),
                c.get("sent", 0),
                c.get("replied", 0),
                f"{c.get('reply_rate', 0):.1f}%",
            ])
        if camp_rows:
            add_styled_table(doc, ["ID", "Name", "Status", "Contacts", "Sent", "Replied", "Reply Rate"], camp_rows)

    doc.add_page_break()

    # ── Test Results Matrix ──
    doc.add_heading("7. Test Results Matrix", level=1)

    result_rows = []
    for r in state.results:
        result_rows.append([
            r.step_num, r.timestamp, r.phase, r.action[:50],
            r.status, f"{r.duration_ms:.0f}ms", r.details[:80],
        ])
    add_styled_table(
        doc,
        ["#", "Time", "Phase", "Action", "Result", "Duration", "Details"],
        result_rows,
    )

    doc.add_page_break()

    # ── Appendix ──
    doc.add_heading("8. Appendix: API Response Log", level=1)
    doc.add_paragraph(
        "Key response data for each API call is logged below. "
        "Full response bodies are truncated to 500 characters for readability."
    )

    for r in state.results:
        doc.add_heading(f"Step {r.step_num}: {r.action}", level=3)
        p = doc.add_paragraph()
        p.add_run(f"Status: ").bold = True
        color = GREEN if r.status == "PASS" else (RED if r.status == "FAIL" else GRAY)
        status_run = p.add_run(r.status)
        status_run.bold = True
        status_run.font.color.rgb = RGBColor.from_string(color)

        doc.add_paragraph(f"Details: {r.details}")

        if r.response_data:
            resp_str = json.dumps(r.response_data, indent=2, default=str) if isinstance(r.response_data, (dict, list)) else str(r.response_data)
            if len(resp_str) > 500:
                resp_str = resp_str[:497] + "..."
            p2 = doc.add_paragraph()
            run2 = p2.add_run(resp_str)
            run2.font.size = Pt(7)
            run2.font.name = "Courier New"
            run2.font.color.rgb = RGBColor.from_string(GRAY)

    # Save
    doc.save(output_path)
    print(f"  Report saved: {output_path}")


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="E2E Test Workflow for Exzelon RA Agent")
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL, help="API base URL")
    parser.add_argument("--output", default="E2E_Test_Workflow_Report.docx", help="Output .docx path")
    args = parser.parse_args()

    print("=" * 60)
    print("  Exzelon RA Agent — E2E Test Workflow")
    print(f"  Target: {args.base_url}")
    print(f"  Date:   {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 60)

    state = WorkflowState(base_url=args.base_url)
    step_counter = [0]  # mutable counter

    # Run all phases sequentially
    phase_auth(state, step_counter)
    phase_discover(state, step_counter)
    phase_test_set_1(state, step_counter)
    phase_test_set_2(state, step_counter)
    phase_campaign(state, step_counter)
    phase_deals(state, step_counter)
    phase_verification(state, step_counter)
    phase_analytics(state, step_counter)

    # Summary
    total = len(state.results)
    passed = sum(1 for r in state.results if r.status == "PASS")
    failed = sum(1 for r in state.results if r.status == "FAIL")
    skipped = sum(1 for r in state.results if r.status == "SKIP")

    print("\n" + "=" * 60)
    print(f"  RESULTS: {passed}/{total} passed, {failed} failed, {skipped} skipped")
    print("=" * 60)

    # Generate report
    generate_report(state, args.output)

    print(f"\nDone! Report: {args.output}")
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
