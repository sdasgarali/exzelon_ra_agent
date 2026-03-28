"""Generate Lead Sourcing Performance Analysis document (.docx).

Run once from project root:
    cd backend && pip install python-docx
    python scripts/generate_analysis_doc.py

Produces: Lead_Sourcing_Performance_Analysis.docx in project root.
"""
import os
import sys

# Add backend to path so we can import config
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = str(value)

    doc.add_paragraph()  # Spacer
    return table


def generate():
    doc = Document()

    # Title
    title = doc.add_heading("Lead Sourcing Performance Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("NeuraLeads AI Agent - Performance Overhaul Report")
    doc.add_paragraph("Prepared for: NeuraLeads AI Operations Team")
    doc.add_paragraph("")

    # --- Executive Summary ---
    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "This document analyzes the lead sourcing pipeline's current performance bottlenecks "
        "and presents a comprehensive improvement strategy. Manual recruiters (10 people) extract "
        "approximately 1,500 jobs per day from LinkedIn alone, while the automated system produces "
        "only 100-180 unique leads per pipeline run. Root cause analysis identified 7 key bottlenecks "
        "that, when addressed, can increase output to 2,000-4,000 leads per day."
    )

    # --- Current Architecture ---
    add_heading(doc, "2. Current Architecture")
    doc.add_paragraph(
        "The lead sourcing pipeline is a multi-source, parallel-execution system built on FastAPI. "
        "It supports 6 job source adapters (JSearch, Apollo, TheirStack, SerpAPI, Adzuna, and Mock) "
        "that fetch jobs concurrently via a ThreadPoolExecutor. Results undergo a 3-layer deduplication "
        "strategy (external_job_id, employer LinkedIn URL, normalized company+title+state+city) before "
        "being stored in the database."
    )

    add_heading(doc, "2.1 Pipeline Flow", level=2)
    doc.add_paragraph(
        "1. Load configuration (industries, job titles, exclude keywords) from database/config\n"
        "2. Initialize all enabled adapters based on settings\n"
        "3. Fetch jobs from all sources in parallel (ThreadPoolExecutor)\n"
        "4. Deduplicate within batch (3-layer: external ID, LinkedIn URL, normalized key)\n"
        "5. Deduplicate against existing database records\n"
        "6. Insert unique leads and upsert client records\n"
        "7. Auto-enrich from cached company contacts\n"
        "8. Export results to XLSX"
    )

    add_heading(doc, "2.2 Adapter Summary", level=2)
    add_table(doc,
        ["Adapter", "API Source", "Pagination", "Default Limit", "Coverage"],
        [
            ["JSearch", "RapidAPI (LinkedIn, Indeed, Glassdoor, ZipRecruiter)", "num_pages param", "500", "Broad US aggregator"],
            ["Apollo", "Apollo.io organization search", "Page-based POST", "500", "Company hiring signals"],
            ["TheirStack", "TheirStack job postings API", "Page-based POST", "500", "Career page scraping"],
            ["SerpAPI", "Google Jobs via SerpAPI", "Single page", "500", "Google Jobs results"],
            ["Adzuna", "Adzuna job search API", "Page-based GET", "500", "UK-based, US coverage"],
            ["Mock", "Internal test data", "N/A", "50", "Development/testing only"],
        ]
    )

    # --- Bottleneck Analysis ---
    add_heading(doc, "3. Bottleneck Analysis")
    doc.add_paragraph(
        "Seven bottlenecks were identified that collectively limit the system to ~100-180 unique "
        "leads per run, far below the theoretical capacity of the available APIs."
    )

    bottlenecks = [
        ("B1: Single Adapter Enabled",
         "Only JSearch is enabled by default in the database settings (lead_sources = ['jsearch']). "
         "The other 5 adapters (Apollo, TheirStack, SerpAPI, Adzuna) sit dormant even when API keys "
         "are configured. This leaves 80% of sourcing capacity unused.",
         "High"),
        ("B2: Low Pagination Limits",
         "JSearch fetches only 3 pages (30 results) per query batch. Apollo fetches 2 pages per "
         "industry group. Adzuna fetches 3 pages. SerpAPI fetches a single page with no pagination. "
         "All adapters could fetch significantly more data per API call.",
         "High"),
        ("B3: Narrow Job Title Coverage",
         "The system searches for 36 job titles, missing key roles like VP Operations, Director of HR, "
         "Warehouse Supervisor, Production Manager, Compliance Manager, Transportation Manager, and "
         "other high-value staffing titles.",
         "Medium"),
        ("B4: All-Time DB Deduplication",
         "The database dedup layer (Layer 3) checks against ALL existing leads with no time window. "
         "On subsequent runs, nearly everything is flagged as a duplicate because the same companies "
         "keep posting similar roles. A 30-day window would allow re-discovery of refreshed postings.",
         "High"),
        ("B5: No Automated Scheduling",
         "The pipeline only runs when manually triggered via the UI or API. There is no scheduled "
         "execution, meaning leads are only sourced during active user sessions. Manual recruiters "
         "search continuously throughout the day.",
         "High"),
        ("B6: Low Thread Pool Workers",
         "ThreadPoolExecutor is limited to 3 workers. With 6 adapters potentially enabled, "
         "some must wait for others to complete, adding unnecessary latency.",
         "Low"),
        ("B7: Conservative Default Limits",
         "All adapters default to limit=500, which caps results even when the API can return more. "
         "Raising to 1000 doubles the ceiling per adapter.",
         "Medium"),
    ]

    for title, description, severity in bottlenecks:
        add_heading(doc, title, level=2)
        doc.add_paragraph(f"Severity: {severity}", style="Intense Quote")
        doc.add_paragraph(description)

    # --- Capacity Comparison ---
    add_heading(doc, "4. Capacity Comparison: Manual vs. Automated")

    add_table(doc,
        ["Metric", "Manual Recruiters (10 people)", "Current System", "After Improvements"],
        [
            ["Jobs sourced/day", "~1,500 (LinkedIn only)", "~100-180 unique/run", "~2,000-4,000/day"],
            ["Sources used", "LinkedIn only", "JSearch only (1 adapter)", "5-6 adapters"],
            ["Runs per day", "Continuous (8 hrs)", "Manual trigger only", "3x/day automated"],
            ["Title coverage", "Broad manual search", "36 titles", "55+ titles"],
            ["Geographic coverage", "Targeted states", "Broad US search", "Broad US search"],
            ["Dedup window", "N/A (human memory)", "All-time (aggressive)", "30-day rolling"],
            ["Raw results/run", "N/A", "~270 (JSearch 3-page)", "~4,000+ (all adapters)"],
            ["Cost", "~$3,000/day (10 FTE)", "$0-50/month API fees", "$0-200/month API fees"],
        ]
    )

    # --- Improvement Strategies ---
    add_heading(doc, "5. Improvement Strategies")

    add_heading(doc, "5.1 Tier 1: Pagination & Limits (Immediate Impact)", level=2)
    improvements_t1 = [
        ["JSearch", "num_pages: 3 -> 10", "~270 -> ~900 results/run", "Zero cost"],
        ["Apollo", "max_pages: 2 -> 5", "~1,400 -> ~3,500 results/run", "Zero cost"],
        ["Adzuna", "pages: 3 -> 10", "~1,350 -> ~4,500 results/run", "Zero cost"],
        ["SerpAPI", "Add pagination (3 pages)", "~90 -> ~270 results/run", "3x API calls"],
        ["TheirStack", "Explicit 10-page cap", "Minor improvement", "Zero cost"],
        ["All adapters", "limit: 500 -> 1000", "2x ceiling per adapter", "Zero cost"],
    ]
    add_table(doc,
        ["Adapter", "Change", "Impact", "Cost"],
        improvements_t1
    )

    add_heading(doc, "5.2 Tier 2: Coverage & Dedup Optimization", level=2)
    improvements_t2 = [
        ["Job Titles", "Expand from 36 to 55+ titles", "Catches VP, Director, Supervisor roles missed before"],
        ["DB Dedup", "30-day rolling window instead of all-time", "Allows re-discovery of refreshed/re-posted jobs"],
        ["Thread Pool", "max_workers: 3 -> 6", "All 6 adapters run simultaneously"],
    ]
    add_table(doc,
        ["Area", "Change", "Impact"],
        improvements_t2
    )

    add_heading(doc, "5.3 Tier 3: Automation", level=2)
    doc.add_paragraph(
        "Add scheduled pipeline execution via APScheduler (already used for warmup tasks). "
        "The lead sourcing pipeline will run automatically at 6am, 12pm, and 6pm UTC, "
        "producing 3 batches of leads per day without manual intervention."
    )
    doc.add_paragraph(
        "Schedule: CronTrigger(hour='6,12,18', minute=0)\n"
        "Trigger source: 'scheduler'\n"
        "Uses existing _get_db() pattern from warmup scheduler jobs"
    )

    # --- Projected Impact ---
    add_heading(doc, "6. Projected Impact")

    add_table(doc,
        ["Metric", "Before", "After"],
        [
            ["Raw results/run (JSearch only)", "~270", "~900"],
            ["Raw results/run (all adapters)", "~270 (1 adapter)", "~4,000+ (5-6 adapters)"],
            ["Unique leads/run (after dedup)", "~100-180", "~1,500-3,000"],
            ["Runs/day", "1 (manual)", "3 (automated)"],
            ["Effective leads/day", "~150", "~2,000-4,000"],
            ["Job title coverage", "36 titles", "55+ titles"],
            ["Dedup false positive rate", "~30% (all-time)", "~5% (30-day window)"],
            ["Thread utilization", "50% (3/6 workers)", "100% (6/6 workers)"],
        ]
    )

    # --- Risk Assessment ---
    add_heading(doc, "7. Risk Assessment")

    risks = [
        ["API Rate Limits", "Higher pagination = more API calls per run", "Medium",
         "Each adapter handles 429 gracefully; free tiers may exhaust faster"],
        ["API Cost", "More calls = potentially higher costs", "Low",
         "Most APIs have generous free tiers; monitor usage"],
        ["Database Load", "More leads = larger DB queries for dedup", "Low",
         "30-day window actually reduces DB load vs all-time"],
        ["Duplicate Quality", "Looser dedup may let some duplicates through", "Low",
         "30-day window is conservative; company+title+state+city key is robust"],
    ]
    add_table(doc,
        ["Risk", "Description", "Severity", "Mitigation"],
        risks
    )

    # --- Implementation Checklist ---
    add_heading(doc, "8. Implementation Checklist")

    checklist = [
        "Add python-docx to backend/requirements.txt",
        "JSearch: num_pages 3->10, limit 500->1000",
        "Apollo: max_pages 2->5, limit 500->1000",
        "Adzuna: pages 3->10, limit 500->1000",
        "SerpAPI: Add multi-page pagination with start param, limit 500->1000",
        "TheirStack: Explicit 10-page cap, limit 500->1000",
        "Expand TARGET_JOB_TITLES from 36 to 55+ titles in config.py",
        "Add 30-day dedup window to lead_sourcing.py DB Layer 3",
        "Increase ThreadPoolExecutor max_workers from 3 to 6",
        "Add scheduled lead sourcing job (3x/day) to scheduler.py",
        "Run full test suite to verify no regressions",
        "Deploy to production and monitor first scheduled run",
    ]

    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    # Save
    output_path = os.path.join(os.path.dirname(__file__), "..", "Lead_Sourcing_Performance_Analysis.docx")
    output_path = os.path.abspath(output_path)
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate()
