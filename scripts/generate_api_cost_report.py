"""Generate the API Cost & Usage Report (.docx).

Summarizes total external-API spend per provider (from the production
cost_entries table) and explains how job/company filters affect API credit
usage — including the new query-side negative-keyword optimization and the
newly-instrumented contact-discovery + AI cost tracking.

Data snapshot: production (ra.partnerwithus.tech), 2026-03-21 → 2026-05-23.

Run:     python scripts/generate_api_cost_report.py
Produces: API_Cost_Report.docx in the project root.
"""
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# -- Palette --
WHITE = "FFFFFF"
DARK = "1F2937"
GRAY = "6B7280"
HEADER_BG = "1E3A5F"
ALT_ROW = "F3F4F6"
GREEN = "166534"
RED = "991B1B"
BLUE = "1E40AF"
AMBER = "92400E"

# -- Production data snapshot (cost_entries, 2026-03-21 → 2026-05-23) --
SNAPSHOT_PERIOD = "2026-03-21 to 2026-05-23 (~2 months)"
PROVIDERS = [
    # provider, category, api_calls, results, cost_usd, per_request_rate
    ("SerpAPI",      "Lead sourcing", 2325, 16549, 108.60, "$0.05/req"),
    ("SearchAPI",    "Lead sourcing", 1228,  9173,  49.12, "$0.04/req"),
    ("Apollo",       "Lead sourcing", 3713, 161018, 37.13, "$0.01/req"),
    ("Adzuna",       "Lead sourcing", 1384,  3379,   3.83, "$0.004/req"),
    ("JSearch",      "Lead sourcing", 1968, 23551,   3.07, "$0.005/req"),
    ("TheirStack",   "Lead sourcing",   51,     0,   0.00, "$0.005/req"),
    ("NPI Registry", "Lead sourcing",    2,     0,   0.00, "Free"),
]
TOTAL_COST = 201.75
TOTAL_CALLS = 10671
TOTAL_RESULTS = 213670


def set_cell_shading(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))


def set_cell_text(cell, text, bold=False, color=None, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(size)


def add_table(doc, headers, rows, col_widths=None, total_row=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header, bold=True, color=WHITE, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row_data in enumerate(rows):
        is_total = total_row and r == len(rows) - 1
        bg = "E5EDF5" if is_total else (WHITE if r % 2 == 0 else ALT_ROW)
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            set_cell_shading(cell, bg)
            set_cell_text(cell, value, bold=is_total, size=9,
                          align=WD_ALIGN_PARAGRAPH.RIGHT if c >= 2 else None)
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def h(doc, text, level):
    return doc.add_heading(text, level=level)


def para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def generate():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    # ---- Title ----
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_heading("API Cost & Usage Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
    sub = para(doc, "Exzelon RA Agent — External API Spend per Provider & the Cost Impact of Job Filters",
               italic=True, color=GRAY, size=13)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = para(doc, f"Data source: production cost_entries  |  Period: {SNAPSHOT_PERIOD}",
                color=GRAY, size=10)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ---- 1. Executive Summary ----
    h(doc, "1. Executive Summary", 1)
    para(doc, f"Across the snapshot period, the platform recorded ${TOTAL_COST:,.2f} of tracked "
              f"external-API spend over {TOTAL_CALLS:,} API calls returning {TOTAL_RESULTS:,} "
              f"results. SerpAPI and SearchAPI alone account for ~78% of cost despite being a "
              f"minority of calls, because they carry the highest per-request rates.")
    bullet(doc, "Highest spend: SerpAPI ($108.60, 54%) and SearchAPI ($49.12, 24%).", "Top spenders — ")
    bullet(doc, "Cheapest at volume: JSearch and Adzuna (generous free tiers, low per-request rates).", "Best value — ")
    bullet(doc, "Until this release, only job-board sourcing was tracked. Contact-discovery, "
                "email-validation, and AI/LLM costs were NOT recorded — likely the largest "
                "untracked spend (15,989 contacts were enriched with $0 recorded).", "Blind spots (now fixed) — ")

    # ---- 2. Cost per Provider ----
    h(doc, "2. Total API Cost per Provider", 1)
    rows = [[p[0], p[1], f"{p[2]:,}", f"{p[3]:,}", f"${p[4]:,.2f}", p[5]] for p in PROVIDERS]
    rows.append(["TOTAL", "—", f"{TOTAL_CALLS:,}", f"{TOTAL_RESULTS:,}", f"${TOTAL_COST:,.2f}", "—"])
    add_table(doc,
              ["Provider", "Category", "API Calls", "Results", "Cost (USD)", "Rate"],
              rows, col_widths=[3.2, 3.2, 2.4, 2.4, 2.6, 2.6], total_row=True)

    para(doc, "Three caveats about these figures:", bold=True)
    bullet(doc, "Only the lead_sourcing category was populated. Contact discovery, email "
                "validation, and AI/LLM token costs were untracked (now instrumented — see §4).",
           "Scope — ")
    bullet(doc, "Costs are estimates from a hardcoded pricing table, not reconciled against "
                "provider invoices.", "Estimates — ")
    bullet(doc, "The free-tier deduction is applied per run rather than per month, so high-free-tier "
                "providers (JSearch, Adzuna) are understated; SerpAPI/SearchAPI are close to accurate.",
           "Free-tier skew — ")

    # ---- 3. Impact of Job Filters on API Credit Usage ----
    h(doc, "3. Do Job/Company Filters Reduce API Credits?", 1)
    para(doc, "Short answer: it depends on the API layer. The ~1,100-company exclusion blocklist and "
              "the exclusion keywords are applied AFTER the job-board API responses arrive, so by "
              "themselves they do NOT save job-board credits — but they DO save downstream "
              "contact-discovery credits, which are the most expensive layer.")

    h(doc, "Order of operations (lead sourcing → contact discovery)", 2)
    for step in [
        "1. Job-board APIs are called  →  credits charged HERE (SerpAPI, SearchAPI, JSearch, Adzuna, TheirStack, Apollo).",
        "2. Results normalized + de-duplicated in memory.",
        "3. Exclusion KEYWORDS applied (Python filter).",
        "4. ~1,100-company BLOCKLIST applied (Python filter).",
        "5. Surviving leads stored in the database.",
        "6. Contact-discovery APIs called only on stored leads  →  credits charged HERE (Apollo/Seamless/Hunter/Snov/PDL enrichment).",
    ]:
        doc.add_paragraph(step, style="List Number" if False else None)

    h(doc, "Effect by layer", 2)
    add_table(doc,
              ["API layer", "Filters reduce credits?", "Why"],
              [
                  ["Job boards (SerpAPI, SearchAPI, JSearch, Adzuna, TheirStack)",
                   "Not by themselves", "Filters run after the call; the credit is already spent."],
                  ["Contact discovery (Apollo/Seamless/Hunter/Snov/PDL)",
                   "Yes — significantly", "Excluded companies never become leads, so they never trigger 1–4 enrichment calls each."],
              ],
              col_widths=[6.2, 4.0, 6.0])

    para(doc, "Practical takeaway: keep the exclusion lists — they directly cut your most expensive "
              "(previously invisible) spend, contact enrichment. To cut job-board credits specifically, "
              "the levers are fewer queries (titles × locations × pages), a tighter posted-within window, "
              "lower page caps, and shifting volume away from SerpAPI/SearchAPI toward JSearch/Adzuna/Jooble/USAJOBS.",
          bold=True)

    # ---- 4. What changed in this release ----
    h(doc, "4. Improvements Shipped in This Release", 1)

    h(doc, "(a) Query-side negative-keyword filtering", 2)
    para(doc, "Exclusion keywords are now pushed INTO the upstream query where the source supports it, "
              "so excluded postings are filtered server-side instead of fetched-then-discarded. This "
              "reduces wasted result volume and, for paginated sources, can reduce the number of paid calls.")
    bullet(doc, "Google-style negative terms appended to the query (e.g. -intern -\"staffing agency\").", "JSearch & SerpAPI — ")
    bullet(doc, "Native what_exclude request parameter.", "Adzuna — ")
    bullet(doc, "The local filter still runs as a backstop; toggle via the push_exclusions_to_query setting (default on).",
           "Safety — ")

    h(doc, "(b) Contact-discovery cost tracking", 2)
    para(doc, "The contact-enrichment pipeline now records a CostEntry per provider "
              "(category = contact_discovery) using the per-adapter call/result counts it already "
              "collects — closing the biggest blind spot in the old report.")

    h(doc, "(c) AI/LLM cost tracking", 2)
    para(doc, "Every AI call now records token-based cost (category = ai) from the provider's reported "
              "usage, across OpenAI, Anthropic, Gemini, and Groq (Groq = free tier). The amount column "
              "was widened to 6 decimal places so sub-cent per-call costs are no longer truncated to $0.00. "
              "All new costs surface automatically in the existing /analytics/costs/per-source dashboard.")

    # ---- Footer ----
    doc.add_paragraph()
    f = para(doc, "Generated by scripts/generate_api_cost_report.py — pricing is configurable via the "
                  "Settings key provider_pricing (job/contact) and ai_model_pricing (AI).",
             italic=True, color=GRAY, size=9)
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out_path = os.path.join(os.path.dirname(__file__), "..", "API_Cost_Report.docx")
    out_path = os.path.abspath(out_path)
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    generate()
