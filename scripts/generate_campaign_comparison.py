"""Generate Campaign Deep-Dive: Exzelon RA Agent vs Instantly.ai (.docx).

Comprehensive campaign system comparison with:
- How Exzelon campaigns work (end-to-end flow)
- Feature-by-feature comparison against Instantly.ai
- Improvement areas with priority and effort estimates

Run: python scripts/generate_campaign_comparison.py
Produces: Campaign_Comparison_Instantly_vs_Exzelon.docx in project root.
"""
import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Color Constants ──────────────────────────────────────────────────
GREEN = "22C55E"
RED = "EF4444"
YELLOW = "F59E0B"
BLUE = "3B82F6"
GRAY = "6B7280"
PURPLE = "8B5CF6"
LIGHT_GREEN = "DCFCE7"
LIGHT_RED = "FEE2E2"
LIGHT_YELLOW = "FEF3C7"
LIGHT_BLUE = "DBEAFE"
LIGHT_PURPLE = "EDE9FE"
WHITE = "FFFFFF"
DARK = "1F2937"
HEADER_BG = "1E3A5F"


# ── Helper Functions ─────────────────────────────────────────────────

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_styled_table(doc, headers, rows, header_bg=HEADER_BG):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_bg)
        set_cell_text(cell, header, bold=True, color=WHITE, size=10)

    for row_idx, row_data in enumerate(rows):
        bg = WHITE if row_idx % 2 == 0 else "F3F4F6"
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            text = str(value)
            # Color-code comparison cells
            if col_idx >= 1:
                if text.startswith("[WIN]"):
                    set_cell_shading(cell, LIGHT_GREEN)
                    set_cell_text(cell, text.replace("[WIN]", "").strip(), size=9, color="166534")
                elif text.startswith("[LOSE]"):
                    set_cell_shading(cell, LIGHT_RED)
                    set_cell_text(cell, text.replace("[LOSE]", "").strip(), size=9, color="991B1B")
                elif text.startswith("[TIE]"):
                    set_cell_shading(cell, LIGHT_BLUE)
                    set_cell_text(cell, text.replace("[TIE]", "").strip(), size=9, color="1E40AF")
                elif text.startswith("[GAP]"):
                    set_cell_shading(cell, LIGHT_RED)
                    set_cell_text(cell, text.replace("[GAP]", "").strip(), size=9, color="991B1B")
                elif text.startswith("[PARTIAL]"):
                    set_cell_shading(cell, LIGHT_YELLOW)
                    set_cell_text(cell, text.replace("[PARTIAL]", "").strip(), size=9, color="92400E")
                else:
                    set_cell_shading(cell, bg)
                    set_cell_text(cell, text, size=9)
            else:
                set_cell_shading(cell, bg)
                set_cell_text(cell, text, size=9, bold=True)

    return table


def add_score_table(doc, category, exzelon_score, instantly_score, notes=""):
    """Add a single-row score comparison."""
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cells = table.rows[0].cells

    set_cell_text(cells[0], category, bold=True, size=10)
    set_cell_shading(cells[0], "F3F4F6")

    # Exzelon score
    ex_color = LIGHT_GREEN if exzelon_score >= instantly_score else LIGHT_RED
    set_cell_shading(cells[1], ex_color)
    set_cell_text(cells[1], f"Exzelon: {exzelon_score}/10", bold=True, size=10,
                  color="166534" if exzelon_score >= instantly_score else "991B1B")

    # Instantly score
    in_color = LIGHT_GREEN if instantly_score >= exzelon_score else LIGHT_RED
    set_cell_shading(cells[2], in_color)
    set_cell_text(cells[2], f"Instantly: {instantly_score}/10", bold=True, size=10,
                  color="166534" if instantly_score >= exzelon_score else "991B1B")

    set_cell_text(cells[3], notes, size=9)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(DARK)
    return h


def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


# ── Document Generation ──────────────────────────────────────────────

def generate_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("Campaign System Deep-Dive", level=0)
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor.from_string(DARK)

    subtitle = doc.add_heading("Exzelon RA Agent vs Instantly.ai", level=1)
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor.from_string(BLUE)

    add_para(doc, f"Generated: {datetime.now().strftime('%B %d, %Y')}", italic=True, size=11)
    add_para(doc, "Document Type: Technical Comparison & Improvement Analysis", size=11)
    add_para(doc, "Scope: Campaign creation, execution, safety, analytics, and deliverability", size=11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. How Exzelon Campaigns Work (End-to-End Flow)",
        "   2.1 Campaign Lifecycle",
        "   2.2 Sequence Engine Architecture",
        "   2.3 Safety & Compliance Layer",
        "   2.4 AI Integration",
        "3. Feature-by-Feature Comparison",
        "   3.1 Campaign Creation & Setup",
        "   3.2 Sequence Builder & Step Types",
        "   3.3 A/B Testing & Optimization",
        "   3.4 Personalization & Content",
        "   3.5 Send Scheduling & Timing",
        "   3.6 Deliverability & Safety",
        "   3.7 Analytics & Reporting",
        "   3.8 Contact Management & Enrollment",
        "   3.9 AI & Automation",
        "   3.10 Integrations & API",
        "   3.11 Multi-Channel",
        "   3.12 Team & Collaboration",
        "4. Scorecard Summary",
        "5. Improvement Areas (Prioritized)",
        "   5.1 Critical (Must-Have)",
        "   5.2 High Priority (Competitive Parity)",
        "   5.3 Medium Priority (Differentiation)",
        "   5.4 Nice-to-Have (Future)",
        "6. Implementation Roadmap",
    ]
    for item in toc_items:
        add_para(doc, item, size=10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Executive Summary", level=1)

    add_para(doc, (
        "This document provides a comprehensive deep-dive into Exzelon RA Agent's campaign system, "
        "comparing it feature-by-feature against Instantly.ai — the market-leading cold email platform. "
        "The analysis covers campaign creation, sequence execution, deliverability controls, AI integration, "
        "analytics, and identifies 23 improvement areas with effort estimates."
    ), size=10)

    doc.add_paragraph()

    # Summary scores
    add_para(doc, "Overall Campaign Capability Scores:", bold=True, size=11)
    summary_rows = [
        ["Platform", "Score", "Strengths", "Gaps"],
        ["Exzelon RA Agent", "7.5/10",
         "AI governance, safety controls, send gate, multi-tenant, preview mode",
         "No visual builder, no open tracking in conditions, limited multi-channel"],
        ["Instantly.ai", "8.5/10",
         "Visual builder, inbox placement testing, B2B database, polished UX",
         "No AI safety layer, no centralized send gate, no multi-tenant"],
    ]
    add_styled_table(doc, summary_rows[0], summary_rows[1:])

    doc.add_paragraph()
    add_para(doc, (
        "Key Finding: Exzelon's campaign engine is architecturally superior in safety and AI governance "
        "(send gate, campaign safety, AI orchestrator, content fingerprinting), but Instantly leads in "
        "user experience, visual campaign building, inbox placement testing, and multi-channel execution. "
        "Closing the gap requires 5 critical improvements and 8 high-priority enhancements."
    ), size=10, bold=True)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 2. HOW EXZELON CAMPAIGNS WORK
    # ════════════════════════════════════════════════════════════════
    add_heading(doc, "2. How Exzelon Campaigns Work", level=1)

    # 2.1 Campaign Lifecycle
    add_heading(doc, "2.1 Campaign Lifecycle", level=2)

    add_para(doc, (
        "Every campaign in Exzelon follows a strict state machine with 5 states and defined transitions:"
    ), size=10)

    lifecycle_rows = [
        ["Draft", "Initial state. Edit steps, assign mailboxes, configure settings.", "→ Active, → Archived"],
        ["Active", "Engine processes contacts every 2 minutes. Sends emails within send window.", "→ Paused, → Completed, → Archived"],
        ["Paused", "Manually paused or auto-paused (bounce threshold). No sends.", "→ Active, → Completed, → Archived"],
        ["Completed", "All contacts finished their sequences.", "→ Archived"],
        ["Archived", "Terminal state. Soft-deleted, hidden from active views.", "No transitions (terminal)"],
    ]
    add_styled_table(doc, ["Status", "Description", "Allowed Transitions"], lifecycle_rows)

    doc.add_paragraph()
    add_para(doc, "State Machine Enforcement:", bold=True, size=10)
    add_para(doc, (
        "All status changes go through validate_campaign_transition() in core/state_machine.py. "
        "Invalid transitions return HTTP 400. The state machine is enforced at the API layer — "
        "the frontend cannot bypass it. Auto-pause uses the same transition system."
    ), size=10)

    # 2.2 Sequence Engine Architecture
    add_heading(doc, "2.2 Sequence Engine Architecture", level=2)

    add_para(doc, "The campaign engine (campaign_engine.py) runs as a scheduler job every 2 minutes:", size=10)

    engine_rows = [
        ["1. Query Due Contacts", "SELECT campaign_contacts WHERE next_send_at <= NOW() AND status = active AND campaign is ACTIVE"],
        ["2. Send Window Check", "Per-campaign timezone + send_window_start/end + send_days (Mon-Fri default). Skips campaigns outside their window."],
        ["3. Safety Checks (per contact)", "Idempotency guard → Reply detection → Unified Send Gate (10 checks)"],
        ["4. Mailbox Selection", "Health-aware scoring: health*0.4 + quota*0.3 + warmup_age*0.15 + deliverability*0.15 + engagement*0.10"],
        ["5. Slow Ramp", "If enabled, effective_limit = slow_ramp_increment × (day + 1), capped at daily_limit"],
        ["6. Smart Throttle", "Jittered daily limit (85-95% of actual), random delay between sends (configurable MIN/MAX seconds)"],
        ["7. A/B Variant Resolution", "Weighted random assignment, stored in variant_assignments_json for consistency"],
        ["8. Spintax Processing", "Nested {option1|option2} resolution with contact_id + campaign_id mixed seed"],
        ["9. Jinja2 Templating", "Variables: contact_first_name, company_name, job_title, job_location, sender.name, etc."],
        ["10. Content Uniqueness", "Jaccard similarity check via content_fingerprint.py — logs warning if >80% similar to recent sends"],
        ["11. Email Send", "Via send_outreach_email() with SMTP. Tracks message_id for threading."],
        ["12. Post-Send", "Update mailbox counters, step stats, campaign stats, sync to unified inbox, trigger deal automation"],
        ["13. Advance Step", "Move contact to next step with delay. Mark COMPLETED if no more steps."],
    ]
    add_styled_table(doc, ["Step", "Description"], engine_rows)

    doc.add_paragraph()
    add_para(doc, "Step Types Supported:", bold=True, size=10)
    step_rows = [
        ["Email", "Send an email (with A/B variants, spintax, placeholders, signature)", "Yes — full implementation"],
        ["Wait", "Delay N days + M hours before next step", "Yes — advance past on next tick"],
        ["Condition", "Branch based on opened/clicked/replied/no_action within time window", "Yes — if/then branching"],
        ["SMS", "Send SMS via Twilio", "Model exists, not wired into engine"],
        ["LinkedIn", "LinkedIn outreach step", "Model exists, not wired into engine"],
        ["Call", "Phone call step", "Model exists, not wired into engine"],
    ]
    add_styled_table(doc, ["Type", "Description", "Implementation Status"], step_rows)

    # 2.3 Safety & Compliance Layer
    add_heading(doc, "2.3 Safety & Compliance Layer", level=2)

    add_para(doc, (
        "Exzelon has a unique 'Centralized Send Gate' — every email-sending path in the system "
        "MUST call unified_send_gate() before transmitting. This is a significant architectural advantage."
    ), size=10, bold=True)

    gate_rows = [
        ["1. Contact Status", "Block unsubscribed or inactive contacts", "Instant — no DB query"],
        ["2. Suppression List", "Check global suppression list (with expiry)", "1 DB query"],
        ["3. Email Validation", "Only send to contacts with 'Valid' email status", "1 DB query"],
        ["4. Contact+Lead Cooldown", "No duplicate sends to same contact for same lead within N days", "1 DB query"],
        ["5. Contact Cooldown", "No sends to same contact within N days (any lead/campaign)", "1 DB query"],
        ["6. Lead Contact Limit", "Max N contacts per lead (default 4)", "1 DB query"],
        ["7. Company Cap", "Max 5 contacts per company across all campaigns", "1 DB query"],
        ["8. Sequence Fatigue", "Block after 5 unanswered emails in 90 days", "2 DB queries"],
        ["9. Domain Throttle", "Per-recipient-domain daily caps (gmail: 30, general: 50)", "1 DB query"],
        ["10. AI Orchestrator", "AI sales agent policy + scoring + content evaluation", "AI call"],
    ]
    add_styled_table(doc, ["Check #", "What It Does", "Cost"], gate_rows)

    doc.add_paragraph()
    add_para(doc, "Additional Campaign Safety Controls (campaign_safety.py):", bold=True, size=10)
    safety_rows = [
        ["Idempotency Guard", "Prevents duplicate sends if scheduler runs twice (checks OutreachEvent for today)"],
        ["Company Contact Cap", "Limits contacts emailed per company across ALL campaigns (default 5)"],
        ["Smart Pause on Reply", "Auto-pauses contact sequence when inbox reply detected"],
        ["Sequence Fatigue", "Blocks contact after N unanswered emails in 90-day window"],
        ["Cross-Campaign Dedup", "Prevents sending from multiple campaigns to same contact simultaneously"],
    ]
    add_styled_table(doc, ["Control", "Description"], safety_rows)

    # 2.4 AI Integration
    add_heading(doc, "2.4 AI Integration", level=2)

    ai_rows = [
        ["AI Sales Agent Orchestrator", "orchestrate_send() gates every outbound email. Builds context, resolves policies, makes structured go/no-go decision with reason codes."],
        ["AI Sequence Generator", "Generates multi-step email sequences from goal/product/tone. 4 AI providers (Groq, OpenAI, Anthropic, Gemini) with template fallback."],
        ["AI Reply Intelligence", "2-tier classification: LLM first → keyword fallback. Categories: interested, objection, question, OOO, unsubscribe."],
        ["AI Reply Agent", "HITL + Autopilot auto-reply with 70%+ confidence gating. Generates drafts, auto-sends after delay."],
        ["AI Safety", "Prompt injection defense for inbound emails. Strips malicious patterns."],
        ["AI Audit Logger", "Logs every AI decision to automation_events with prompt hash, confidence, tokens, latency."],
        ["AI Content Schemas", "Pydantic structured output validation for all AI responses."],
        ["AI Resilience", "Retry with exponential backoff (3 attempts) + provider fallback chain."],
        ["Email Humanizer", "Post-generation pass: burstiness scoring, sentence variation, natural imperfections, conversational markers."],
        ["Content Fingerprint", "Jaccard similarity on 3-shingles + Shannon entropy to detect near-identical emails."],
    ]
    add_styled_table(doc, ["Component", "Description"], ai_rows)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 3. FEATURE-BY-FEATURE COMPARISON
    # ════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Feature-by-Feature Comparison", level=1)

    # 3.1 Campaign Creation & Setup
    add_heading(doc, "3.1 Campaign Creation & Setup", level=2)
    rows_31 = [
        ["Campaign wizard/builder", "[LOSE] Form-based creation (name, timezone, mailboxes, limits)", "[WIN] Visual step-by-step wizard with drag-and-drop"],
        ["Duplicate/clone campaign", "[TIE] Yes — POST /{id}/duplicate, clones steps + settings", "[TIE] Yes — one-click clone with all steps"],
        ["Template library", "[LOSE] No campaign templates (only email templates)", "[WIN] Pre-built campaign templates by industry/goal"],
        ["Campaign preview mode", "[WIN] Preview mode generates drafts instead of sending", "[LOSE] No equivalent — must use test emails"],
        ["Multi-tenant campaigns", "[WIN] Full tenant isolation, plan limits, RBAC", "[LOSE] Workspace-level only, no tenant isolation"],
        ["Campaign daily limit", "[TIE] Configurable per-campaign (default 30)", "[TIE] Configurable per-campaign"],
        ["Send days selection", "[TIE] JSON array of day abbreviations (mon-sun)", "[TIE] Day-of-week picker"],
        ["Mailbox assignment", "[TIE] JSON array of mailbox IDs per campaign", "[TIE] Select accounts to rotate"],
        ["Auto-enrollment rules", "[WIN] Rule-based auto-enrollment (validation, priority, state, score, source)", "[PARTIAL] List-based enrollment, no rule engine"],
        ["Enrollment preview", "[WIN] Preview matching contacts before enrolling", "[LOSE] No enrollment preview"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_31)

    # 3.2 Sequence Builder & Step Types
    add_heading(doc, "3.2 Sequence Builder & Step Types", level=2)
    rows_32 = [
        ["Visual sequence builder", "[LOSE] List-based step editor (add/edit/delete/reorder)", "[WIN] Visual drag-and-drop sequence builder"],
        ["Email steps", "[TIE] Full: subject, HTML body, text body, template link", "[TIE] Full: subject, body, variables"],
        ["Wait/delay steps", "[TIE] Configurable days + hours", "[TIE] Configurable delay between emails"],
        ["Condition steps (if/then)", "[WIN] Branch on opened/clicked/replied/no_action with time window", "[PARTIAL] Subsequences on Hypergrowth plan only, separate screen"],
        ["SMS steps", "[PARTIAL] Model defined (StepType.SMS), not wired to engine", "[WIN] Native SMS via Hyper CRM ($97/mo)"],
        ["LinkedIn steps", "[PARTIAL] Model defined (StepType.LINKEDIN), not wired to engine", "[LOSE] No native LinkedIn (requires third-party)"],
        ["Call steps", "[PARTIAL] Model defined (StepType.CALL), not wired to engine", "[WIN] Native calling via Hyper CRM ($97/mo)"],
        ["Step reordering", "[TIE] PUT /steps/reorder with ordered step_id list", "[TIE] Drag-and-drop reorder"],
        ["Reply-to-thread", "[WIN] Per-step toggle (reply_to_thread boolean)", "[TIE] Thread replies by default"],
        ["Template linking", "[WIN] Link step to email_template by template_id", "[PARTIAL] Copy from templates, no live linking"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_32)

    # 3.3 A/B Testing
    add_heading(doc, "3.3 A/B Testing & Optimization", level=2)
    rows_33 = [
        ["A/B variant support", "[TIE] Per-step variants_json with subject/body/weight", "[TIE] A/Z testing — unlimited variants per step"],
        ["Weighted distribution", "[WIN] Custom weight per variant (e.g., 70/30 split)", "[PARTIAL] Equal distribution by default"],
        ["Variant assignment", "[WIN] Deterministic per contact+step (stored in assignments JSON)", "[TIE] Random assignment"],
        ["Auto-optimize", "[LOSE] Chi-squared planned but not implemented", "[WIN] Auto-optimize switches off underperforming variants"],
        ["Per-step analytics", "[TIE] sent/opened/clicked/replied/bounced per step", "[TIE] Per-step performance metrics"],
        ["Winner selection", "[LOSE] Manual only", "[WIN] Automatic winner selection based on reply rates"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_33)

    # 3.4 Personalization & Content
    add_heading(doc, "3.4 Personalization & Content", level=2)
    rows_34 = [
        ["Variable placeholders", "[TIE] {{contact_first_name}}, {{company_name}}, {{job_title}}, etc.", "[TIE] {{firstName}}, {{companyName}}, custom variables"],
        ["Jinja2 templating", "[WIN] Full Jinja2 with nested objects (contact.*, lead.*, sender.*)", "[LOSE] Simple variable substitution only"],
        ["Spintax", "[WIN] Nested spintax with campaign-aware seeding for cross-campaign variation", "[TIE] Basic spintax support"],
        ["Email humanizer", "[WIN] AI anti-detection: burstiness, sentence variation, imperfections", "[LOSE] No equivalent anti-detection system"],
        ["AI sequence generation", "[TIE] 4-provider AI generation with template fallback", "[TIE] AI sequence writer from structured prompt"],
        ["Email signature", "[TIE] Per-mailbox signature injection via {{signature}}", "[TIE] Per-account email signature"],
        ["Content fingerprint", "[WIN] Jaccard similarity + Shannon entropy uniqueness check", "[LOSE] No content similarity detection"],
        ["Spam score check", "[WIN] 106 trigger words + 6 regex patterns + link/image ratio", "[PARTIAL] SpamAssassin integration (inbox placement test)"],
        ["Unsubscribe footer", "[TIE] Auto-generated per-email with tracking_id", "[TIE] Auto-inserted unsubscribe link"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_34)

    # 3.5 Send Scheduling & Timing
    add_heading(doc, "3.5 Send Scheduling & Timing", level=2)
    rows_35 = [
        ["Send window", "[TIE] Per-campaign start/end time (e.g., 09:00-17:00)", "[TIE] Per-campaign send window"],
        ["Timezone support", "[TIE] Per-campaign timezone setting", "[TIE] Per-campaign timezone"],
        ["Send-time optimization", "[WIN] US state→timezone mapping, optimal windows (Tue 9-11 AM best)", "[PARTIAL] Basic timezone-aware sending"],
        ["Day-of-week selection", "[TIE] Configurable send days (JSON array)", "[TIE] Day-of-week checkboxes"],
        ["Slow ramp", "[TIE] Configurable increment per day per account, tracked per campaign", "[TIE] Slow ramp warmup with gradual volume increase"],
        ["Smart throttle", "[WIN] Jittered daily limits (85-95%) + random inter-send delays", "[PARTIAL] Fixed limits per account"],
        ["Processing frequency", "[TIE] Every 2 minutes via APScheduler", "[TIE] Continuous processing"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_35)

    # 3.6 Deliverability & Safety
    add_heading(doc, "3.6 Deliverability & Safety", level=2)
    rows_36 = [
        ["Centralized send gate", "[WIN] 10-check unified_send_gate() on ALL send paths", "[LOSE] No centralized send gate"],
        ["Auto-pause on bounce", "[TIE] Hourly check, auto-pause when bounce rate exceeds threshold", "[TIE] Auto-pause on bounce threshold"],
        ["Domain throttle", "[WIN] Per-recipient-domain daily caps (gmail: 30, general: 50)", "[PARTIAL] Per-account limits only"],
        ["Company contact cap", "[WIN] Max 5 contacts per company across all campaigns", "[LOSE] No cross-campaign company cap"],
        ["Sequence fatigue", "[WIN] Block after 5 unanswered emails in 90 days", "[LOSE] No fatigue detection"],
        ["Cross-campaign dedup", "[WIN] Prevent simultaneous campaigns to same contact", "[PARTIAL] Lead-level dedup only"],
        ["Idempotency guard", "[WIN] Prevent duplicate sends on scheduler restart", "[LOSE] Not documented"],
        ["Bounce handler", "[WIN] SMTP 5xx auto-suppress + contact INACTIVE + mailbox bounce_count", "[TIE] Bounce tracking and auto-removal"],
        ["ESP feedback", "[WIN] Complaint rate tracking, auto-pause mailbox on 0.3% threshold", "[PARTIAL] Spam complaint monitoring"],
        ["Inbox placement test", "[LOSE] Not implemented", "[WIN] Automated tests across Gmail/Outlook/Yahoo"],
        ["Blacklist monitoring", "[TIE] IP/domain blacklist check in warmup engine", "[TIE] Blacklist monitoring with alerts"],
        ["DKIM signing", "[TIE] Optional DKIM for custom SMTP", "[TIE] DNS health checks including DKIM"],
        ["Read emulation", "[TIE] IMAP read emulation for warmup emails", "[TIE] Human-like opens and scrolling simulation"],
        ["SISR (IP rotation)", "[PARTIAL] ISP-aware round-robin (model exists)", "[WIN] Dedicated IPs on Light Speed plan"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_36)

    # 3.7 Analytics & Reporting
    add_heading(doc, "3.7 Analytics & Reporting", level=2)
    rows_37 = [
        ["Campaign dashboard", "[TIE] Overall stats + per-step analytics + funnel view", "[TIE] Dashboard with sent/opened/clicked/replied"],
        ["Per-step analytics", "[TIE] sent/opened/clicked/replied/bounced per step", "[TIE] Step-level performance"],
        ["Funnel view", "[WIN] Contacts-at-each-step funnel visualization", "[PARTIAL] Basic funnel"],
        ["Reply rate tracking", "[TIE] Per-campaign and per-step reply rates", "[TIE] Reply rate dashboard"],
        ["Open tracking", "[PARTIAL] Pixel tracking exists but limited in conditions", "[WIN] Full open tracking with pixel + history panel"],
        ["Click tracking", "[PARTIAL] Link redirect tracking exists", "[WIN] Full click tracking with link-level analytics"],
        ["A/B variant analytics", "[PARTIAL] Per-variant sent counts, no winner UI", "[WIN] Per-variant metrics with auto-winner"],
        ["Date range filtering", "[LOSE] No date range filter on analytics", "[WIN] Custom date range selector"],
        ["Campaign comparison", "[PARTIAL] Basic comparison via /analytics endpoint", "[WIN] Side-by-side campaign comparison"],
        ["Export/download", "[LOSE] No analytics export", "[WIN] CSV export of campaign activities"],
        ["Real-time activity feed", "[PARTIAL] Via automation_events table", "[WIN] Real-time activity feed with filters"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_37)

    # 3.8 Contact Management
    add_heading(doc, "3.8 Contact Management & Enrollment", level=2)
    rows_38 = [
        ["Manual enrollment", "[TIE] POST /contacts with contact_id list", "[TIE] Upload CSV or select from CRM"],
        ["Auto-enrollment", "[WIN] Rule-based engine (validation, priority, state, score, source, daily cap)", "[PARTIAL] List-based, no rule engine"],
        ["Enrollment dedup", "[TIE] Unique constraint on campaign_id + contact_id", "[TIE] Duplicate detection on upload"],
        ["Suppression checking", "[WIN] Global suppression list checked at enrollment + send gate", "[TIE] Blacklist/suppression list"],
        ["Contact progress tracking", "[TIE] current_step, next_send_at, status per enrollment", "[TIE] Contact status within campaign"],
        ["Remove contacts", "[TIE] DELETE /contacts with contact_id list", "[TIE] Remove from campaign"],
        ["Reply auto-pause", "[WIN] Automatic pause on reply + status change to REPLIED", "[TIE] Pause on reply"],
        ["Bounce auto-removal", "[TIE] Status change to BOUNCED", "[TIE] Auto-removal on bounce"],
        ["Contact enrichment", "[WIN] Enriched contact view with lead/company data", "[TIE] Basic contact details"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_38)

    # 3.9 AI & Automation
    add_heading(doc, "3.9 AI & Automation", level=2)
    rows_39 = [
        ["AI send decision", "[WIN] AI orchestrator evaluates every send (policy + scoring + content)", "[LOSE] No AI send gating"],
        ["AI reply classification", "[WIN] 2-tier LLM + keyword fallback with 5 categories", "[TIE] AI classification of replies"],
        ["AI auto-reply", "[WIN] HITL + Autopilot with 70%+ confidence gating", "[TIE] AI Reply Agents"],
        ["AI sequence generation", "[TIE] 4-provider AI with anti-AI-detection prompts", "[TIE] AI Sequence Writer"],
        ["AI lead scoring", "[WIN] Hiring signals + company size + industry + salary + web presence", "[PARTIAL] Basic lead scoring"],
        ["AI personalization planning", "[WIN] Per-step personalization (angle, tone, hooks, CTA type)", "[PARTIAL] AI-powered variable personalization"],
        ["AI safety (prompt injection)", "[WIN] Inbound email sanitization, pattern stripping", "[LOSE] No AI safety layer"],
        ["AI audit logging", "[WIN] Every AI decision logged with confidence, tokens, latency", "[LOSE] No AI audit trail"],
        ["AI resilience", "[WIN] Retry + backoff + provider fallback chain", "[PARTIAL] Single-provider"],
        ["Webhook events", "[TIE] 8 event types with HMAC-SHA256 signing", "[TIE] Webhook delivery for campaign events"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_39)

    # 3.10 Integrations & API
    add_heading(doc, "3.10 Integrations & API", level=2)
    rows_310 = [
        ["REST API", "[TIE] Full CRUD API for campaigns, steps, contacts, analytics", "[TIE] API V2 with comprehensive endpoints"],
        ["Webhooks", "[TIE] HMAC-SHA256 signed, 8 event types, exponential backoff retry", "[TIE] Real-time webhooks for email events"],
        ["CRM integration", "[PARTIAL] HubSpot + Salesforce sync (bidirectional)", "[WIN] HubSpot + Salesforce + Pipedrive + native CRM"],
        ["Zapier/Make", "[LOSE] No native Zapier/Make integration", "[WIN] Zapier (8000+ apps) + Make (3000+ apps)"],
        ["Email verification", "[TIE] 7 providers (NeverBounce, ZeroBounce, Hunter, etc.)", "[TIE] Built-in verification (0.25 credits/address)"],
        ["B2B lead database", "[LOSE] No native lead database (uses external APIs)", "[WIN] 450M+ contacts with filters"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_310)

    # 3.11 Multi-Channel
    add_heading(doc, "3.11 Multi-Channel Support", level=2)
    rows_311 = [
        ["Email", "[TIE] Full SMTP sending with provider rotation", "[TIE] Full email sending with inbox rotation"],
        ["SMS", "[PARTIAL] Twilio integration exists, not in campaign engine", "[WIN] Native SMS in Hyper CRM tier"],
        ["LinkedIn", "[PARTIAL] Step type defined, no execution", "[LOSE] No native LinkedIn (third-party required)"],
        ["Phone/Calling", "[PARTIAL] Step type defined, Twilio integration exists", "[WIN] Native calling in Hyper CRM tier"],
        ["Multi-channel orchestration", "[LOSE] Single-channel (email) campaign engine", "[PARTIAL] SMS + Call in CRM, not in sequence builder"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_311)

    # 3.12 Team & Collaboration
    add_heading(doc, "3.12 Team & Collaboration", level=2)
    rows_312 = [
        ["Multi-user access", "[TIE] 4 roles: super_admin, admin, operator, viewer", "[TIE] Team access with role-based permissions"],
        ["Multi-tenant", "[WIN] Full tenant isolation with plan limits", "[LOSE] Workspace-level, no tenant isolation"],
        ["Campaign ownership", "[TIE] created_by field tracks campaign creator", "[TIE] Creator tracking"],
        ["Permission control", "[WIN] 3-layer permission system (defaults + role + user override)", "[PARTIAL] Basic role-based access"],
        ["Audit trail", "[WIN] Automation events log all actions", "[PARTIAL] Basic activity log"],
        ["Agency/white-label", "[LOSE] No white-label support", "[WIN] Multi-client setup, agency features"],
    ]
    add_styled_table(doc, ["Feature", "Exzelon RA Agent", "Instantly.ai"], rows_312)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 4. SCORECARD SUMMARY
    # ════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Scorecard Summary", level=1)

    scorecard_rows = [
        ["Campaign Creation & Setup", "7", "9", "Instantly has visual wizard + template library"],
        ["Sequence Builder & Steps", "7", "8", "Instantly has drag-and-drop + native multi-channel"],
        ["A/B Testing & Optimization", "7", "9", "Instantly has auto-optimize + winner selection"],
        ["Personalization & Content", "9", "7", "Exzelon has Jinja2 + humanizer + fingerprint"],
        ["Send Scheduling & Timing", "8", "7", "Exzelon has send-time optimizer + smart throttle"],
        ["Deliverability & Safety", "9", "8", "Exzelon has send gate + 5 safety controls"],
        ["Analytics & Reporting", "6", "9", "Instantly has date ranges + export + comparison"],
        ["Contact Management", "8", "7", "Exzelon has auto-enrollment + enrichment"],
        ["AI & Automation", "10", "7", "Exzelon has orchestrator + safety + audit + resilience"],
        ["Integrations & API", "7", "9", "Instantly has Zapier/Make + B2B database"],
        ["Multi-Channel", "4", "7", "Instantly has native SMS + calling"],
        ["Team & Collaboration", "8", "7", "Exzelon has multi-tenant + 3-layer permissions"],
    ]

    add_styled_table(doc, ["Category", "Exzelon /10", "Instantly /10", "Notes"], scorecard_rows)

    doc.add_paragraph()

    # Calculate totals
    ex_total = sum(int(r[1]) for r in scorecard_rows)
    in_total = sum(int(r[2]) for r in scorecard_rows)
    ex_pct = round(ex_total / (len(scorecard_rows) * 10) * 100)
    in_pct = round(in_total / (len(scorecard_rows) * 10) * 100)

    add_para(doc, f"Exzelon Total: {ex_total}/{len(scorecard_rows)*10} ({ex_pct}%)", bold=True, size=12)
    add_para(doc, f"Instantly Total: {in_total}/{len(scorecard_rows)*10} ({in_pct}%)", bold=True, size=12)
    add_para(doc, f"Gap: {in_total - ex_total} points", bold=True, size=12)

    doc.add_paragraph()
    add_para(doc, "Category Breakdown:", bold=True, size=11)
    wins = sum(1 for r in scorecard_rows if int(r[1]) > int(r[2]))
    losses = sum(1 for r in scorecard_rows if int(r[1]) < int(r[2]))
    ties = sum(1 for r in scorecard_rows if int(r[1]) == int(r[2]))
    add_para(doc, f"Exzelon leads in {wins} categories, Instantly leads in {losses}, tied in {ties}", size=10)
    add_para(doc, (
        f"Exzelon's strengths: AI & Automation (10), Deliverability & Safety (9), "
        f"Personalization & Content (9). Instantly's strengths: Analytics (9), "
        f"Campaign Creation (9), A/B Testing (9), Integrations (9)."
    ), size=10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 5. IMPROVEMENT AREAS
    # ════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Improvement Areas (Prioritized)", level=1)

    add_para(doc, (
        "The following improvements are organized by priority. Critical items address competitive gaps "
        "that directly impact user experience and sales demos. High-priority items achieve competitive "
        "parity with Instantly.ai. Medium items create differentiation."
    ), size=10)

    # 5.1 Critical
    add_heading(doc, "5.1 Critical (Must-Have for Competitive Demos)", level=2)
    critical_rows = [
        ["C1", "Visual Sequence Builder",
         "Replace list-based step editor with drag-and-drop visual builder (React Flow or similar). "
         "Show email/wait/condition nodes with connections. Preview email content inline.",
         "40h", "Campaign Creation, Sequence Builder"],
        ["C2", "A/B Auto-Optimize",
         "Implement chi-squared statistical test on reply rates per variant. Auto-disable losing "
         "variants after minimum sample (50 sends). Show confidence percentage in UI.",
         "16h", "A/B Testing"],
        ["C3", "Campaign Analytics Dashboard Overhaul",
         "Add date range filtering, per-variant breakdown, CSV export, real-time activity feed, "
         "and campaign comparison view. Use Recharts for open/reply/bounce trend charts.",
         "32h", "Analytics & Reporting"],
        ["C4", "Inbox Placement Testing",
         "Build or integrate inbox placement testing (Gmail/Outlook/Yahoo seed accounts). Show "
         "inbox vs spam ratio per provider. Alert on placement drops.",
         "24h", "Deliverability"],
        ["C5", "Open & Click Tracking in Campaign Engine",
         "Wire pixel-based open tracking and link redirect tracking into campaign stats and condition "
         "steps. Currently 'opened' condition falls back to reply detection.",
         "12h", "Analytics, Sequence Builder"],
    ]
    add_styled_table(doc,
        ["ID", "Improvement", "Description", "Effort", "Affected Categories"],
        critical_rows, header_bg="991B1B")

    # 5.2 High Priority
    add_heading(doc, "5.2 High Priority (Competitive Parity)", level=2)
    high_rows = [
        ["H1", "Campaign Template Library",
         "Pre-built campaign templates by industry/goal (meeting booking, referral, "
         "re-engagement, event, break-up). One-click import into new campaign.",
         "16h", "Campaign Creation"],
        ["H2", "Zapier/Make Integration",
         "Build Zapier app or Make module for triggers (new reply, bounce, campaign completed) "
         "and actions (enroll contact, create campaign). Enables 8000+ app connections.",
         "24h", "Integrations"],
        ["H3", "Wire SMS Steps into Campaign Engine",
         "Connect StepType.SMS to Twilio send in campaign_engine.py. Handle SMS opt-out, "
         "delivery receipts, and character limits. Add SMS to send gate.",
         "20h", "Multi-Channel"],
        ["H4", "Wire Call Steps into Campaign Engine",
         "Connect StepType.CALL to Twilio click-to-call or scheduled call. Log call outcome "
         "as outreach event. Auto-advance sequence based on result.",
         "16h", "Multi-Channel"],
        ["H5", "Campaign Comparison View",
         "Side-by-side comparison of 2-4 campaigns: reply rates, bounce rates, sends per day, "
         "best performing step, top variant. Export as PDF/CSV.",
         "12h", "Analytics"],
        ["H6", "Date Range Analytics",
         "Add from_date/to_date query params to analytics endpoint. Frontend date picker with "
         "presets (7d, 30d, 90d, custom). Time-series charts for trends.",
         "8h", "Analytics"],
        ["H7", "Analytics CSV Export",
         "Export campaign contacts, step performance, and activity log as CSV. "
         "Stream large exports to avoid timeout.",
         "8h", "Analytics"],
        ["H8", "Activity Feed with Real-Time Updates",
         "Dedicated activity tab showing sent/opened/clicked/replied/bounced events. "
         "WebSocket or polling for real-time updates. Filterable by event type.",
         "16h", "Analytics"],
    ]
    add_styled_table(doc,
        ["ID", "Improvement", "Description", "Effort", "Affected Categories"],
        high_rows, header_bg="92400E")

    # 5.3 Medium Priority
    add_heading(doc, "5.3 Medium Priority (Differentiation)", level=2)
    medium_rows = [
        ["M1", "LinkedIn Step Execution",
         "Integrate LinkedIn automation (connection request, message, profile view) via "
         "browser extension or API partner. Handle LinkedIn rate limits.",
         "40h", "Multi-Channel"],
        ["M2", "B2B Lead Database Integration",
         "Integrate with Apollo/Instantly credits/similar for in-app lead discovery. "
         "Search by role, seniority, company size, industry directly from campaign screen.",
         "24h", "Integrations"],
        ["M3", "Agency/White-Label Support",
         "Sub-accounts per client, custom branding, client-level billing, "
         "shared campaign templates across clients.",
         "60h", "Team & Collaboration"],
        ["M4", "Advanced Subsequences UI",
         "Visual branching UI for condition steps: show yes/no paths visually. "
         "Allow nested conditions (opened AND clicked). Add condition on link click.",
         "24h", "Sequence Builder"],
        ["M5", "Smart Send Time per Contact",
         "Use engagement history to calculate per-contact optimal send time. "
         "Adjust next_send_at to align with their historical open/reply pattern.",
         "16h", "Scheduling"],
        ["M6", "Campaign Health Score",
         "Composite score combining bounce rate, reply rate, open rate, complaint rate. "
         "Traffic-light indicator (green/yellow/red) on campaign list.",
         "8h", "Analytics"],
        ["M7", "Predictive Reply Scoring",
         "ML model predicting reply probability per contact based on historical data. "
         "Surface high-probability contacts for priority sending.",
         "32h", "AI & Automation"],
    ]
    add_styled_table(doc,
        ["ID", "Improvement", "Description", "Effort", "Affected Categories"],
        medium_rows, header_bg="1E40AF")

    # 5.4 Nice-to-Have
    add_heading(doc, "5.4 Nice-to-Have (Future Enhancements)", level=2)
    nice_rows = [
        ["N1", "Campaign Scheduling (Future Send)",
         "Schedule campaign activation for a future date/time instead of immediate activation.",
         "4h", "Campaign Creation"],
        ["N2", "Sending Speed Control Slider",
         "UI slider to control sending speed (conservative/balanced/aggressive) that adjusts "
         "inter-send delay, daily jitter range, and batch size.",
         "8h", "Scheduling"],
        ["N3", "Email Thread Preview",
         "Preview the full email thread a contact will see (step 1 → 2 → 3) with rendered "
         "variables, spintax resolved, and signature injected.",
         "12h", "Campaign Creation"],
    ]
    add_styled_table(doc,
        ["ID", "Improvement", "Description", "Effort", "Affected Categories"],
        nice_rows, header_bg="6B7280")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 6. IMPLEMENTATION ROADMAP
    # ════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Implementation Roadmap", level=1)

    add_para(doc, (
        "Recommended implementation order based on impact-to-effort ratio and competitive urgency:"
    ), size=10)

    roadmap_rows = [
        ["Phase 1", "2-3 weeks", "C5, C2, H6, H7, M6",
         "Open/click tracking + A/B auto-optimize + date range analytics + CSV export + health score",
         "76h", "Fills biggest analytics and A/B gaps with moderate effort"],
        ["Phase 2", "3-4 weeks", "C3, C1, H1",
         "Analytics dashboard overhaul + visual sequence builder + campaign templates",
         "88h", "Transforms UX — the most visible competitive gap"],
        ["Phase 3", "2-3 weeks", "H3, H4, H8, H5",
         "SMS steps + call steps + activity feed + campaign comparison",
         "64h", "Multi-channel execution + real-time monitoring"],
        ["Phase 4", "3-4 weeks", "C4, H2, M4",
         "Inbox placement testing + Zapier/Make + advanced subsequences",
         "72h", "Deliverability assurance + integration ecosystem"],
        ["Phase 5", "6-8 weeks", "M1, M2, M3, M5, M7",
         "LinkedIn steps + B2B database + agency support + smart timing + predictive scoring",
         "172h", "Long-term differentiation and market expansion"],
    ]
    add_styled_table(doc,
        ["Phase", "Timeline", "Items", "Description", "Effort", "Rationale"],
        roadmap_rows)

    doc.add_paragraph()

    # Total effort
    total_effort = sum(int(r[0].replace("h", "")) for r in [
        *[("{}h".format(r[3].replace("h", "")),) for r in critical_rows],
        *[("{}h".format(r[3].replace("h", "")),) for r in high_rows],
        *[("{}h".format(r[3].replace("h", "")),) for r in medium_rows],
        *[("{}h".format(r[3].replace("h", "")),) for r in nice_rows],
    ])

    add_para(doc, "Total Estimated Effort:", bold=True, size=11)
    effort_summary = [
        ["Critical (C1-C5)", "124h", "5 items"],
        ["High Priority (H1-H8)", "120h", "8 items"],
        ["Medium Priority (M1-M7)", "204h", "7 items"],
        ["Nice-to-Have (N1-N3)", "24h", "3 items"],
        ["TOTAL", "472h", "23 items"],
    ]
    add_styled_table(doc, ["Priority", "Effort", "Count"], effort_summary)

    doc.add_paragraph()
    add_para(doc, (
        "With Phase 1 + Phase 2 completed (~164h, ~6 weeks), Exzelon's campaign system would "
        "achieve approximate parity with Instantly.ai on UX and analytics while maintaining its "
        "lead in AI governance, safety controls, and multi-tenant architecture."
    ), size=10, bold=True)

    # ── Save ─────────────────────────────────────────────────────────
    output_path = os.path.join(
        os.path.dirname(__file__), "..",
        "Campaign_Comparison_Instantly_vs_Exzelon.docx"
    )
    doc.save(output_path)
    print(f"Document saved to: {os.path.abspath(output_path)}")
    print(f"Exzelon Campaign Score: {ex_total}/{len(scorecard_rows)*10} ({ex_pct}%)")
    print(f"Instantly Campaign Score: {in_total}/{len(scorecard_rows)*10} ({in_pct}%)")
    return output_path


if __name__ == "__main__":
    generate_document()
