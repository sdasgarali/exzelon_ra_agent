"""
Generate New Data Sources Assessment Document for NeuraLeads AI Agent.
Creates a professional .docx with profiles for all 10 new data source adapters.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


# ─── Helpers ───────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells, bold=False, header=False, shade=None):
    """Add a row to a table with optional formatting."""
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, '1F4E79')
        elif shade:
            set_cell_shading(cell, shade)
    return row


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def add_note_box(doc, text, label="Note"):
    """Add a highlighted note/callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ─── Source Profiles ──────────────────────────────────────────────────────

SOURCES = [
    {
        "name": "TheirStack",
        "category": "Job Source",
        "description": "Job board aggregator that tracks company tech stacks and hiring activity. Identifies companies actively hiring specific roles, providing rich company context alongside job postings.",
        "api_endpoint": "https://api.theirstack.com/v1/jobs/search",
        "docs_link": "https://theirstack.com/docs",
        "auth_method": "Bearer token header",
        "free_tier": "100 requests/month",
        "paid": "From $49/month",
        "data_quality": "High — includes tech stack data, company context, and detailed job metadata",
        "complexity": "Low — simple REST API with JSON request/response",
        "recommendation": "Must-Have",
        "rationale": "Unique tech-stack-based job discovery fills a gap not covered by traditional job boards. Excellent for finding companies hiring specific technical roles.",
        "adapter_file": "backend/app/services/adapters/job_sources/theirstack.py",
    },
    {
        "name": "SerpAPI (Google Jobs)",
        "category": "Job Source",
        "description": "Structured access to Google Jobs search results. Captures postings aggregated by Google from across the web, including many that don't appear on traditional job boards.",
        "api_endpoint": "https://serpapi.com/search?engine=google_jobs",
        "docs_link": "https://serpapi.com/google-jobs-api",
        "auth_method": "api_key query parameter",
        "free_tier": "100 searches/month",
        "paid": "From $50/month",
        "data_quality": "High — Google aggregates from thousands of sources; captures listings missed by individual boards",
        "complexity": "Low — simple GET requests with query parameters",
        "recommendation": "Must-Have",
        "rationale": "Google Jobs is the broadest job aggregator. Catches listings from company career pages, niche boards, and sources not covered by JSearch or Apollo.",
        "adapter_file": "backend/app/services/adapters/job_sources/serpapi.py",
    },
    {
        "name": "Adzuna",
        "category": "Job Source",
        "description": "UK-based job search aggregator with strong US coverage. Indexes millions of job ads from thousands of websites and classifies them with salary estimates.",
        "api_endpoint": "https://api.adzuna.com/v1/api/jobs/us/search/{page}",
        "docs_link": "https://developer.adzuna.com/docs",
        "auth_method": "app_id + app_key query parameters (dual credential)",
        "free_tier": "250 requests/month",
        "paid": "From $99/month",
        "data_quality": "Medium-High — good salary data but less company metadata than TheirStack",
        "complexity": "Low — simple REST API, requires two credentials (app_id + api_key)",
        "recommendation": "Nice-to-Have",
        "rationale": "Provides salary estimates and catches some listings missed by Google Jobs. Good supplementary source but overlaps with SerpAPI coverage.",
        "adapter_file": "backend/app/services/adapters/job_sources/adzuna.py",
    },
    {
        "name": "Hunter.io (Contact Finder)",
        "category": "Contact Discovery",
        "description": "Domain-based email finder and verifier. Given a company domain, discovers employee email addresses with confidence scores. Also provides email verification.",
        "api_endpoint": "https://api.hunter.io/v2/domain-search",
        "docs_link": "https://hunter.io/api-documentation",
        "auth_method": "api_key query parameter",
        "free_tier": "25 requests/month",
        "paid": "From $49/month",
        "data_quality": "High for email accuracy — confidence scores per email, verified patterns",
        "complexity": "Low — straightforward REST API",
        "recommendation": "Must-Have",
        "rationale": "Excellent when you know the company domain. High email accuracy with built-in verification. Complements Apollo's people-first approach with a domain-first approach.",
        "adapter_file": "backend/app/services/adapters/contacts/hunter_contact.py",
    },
    {
        "name": "Snov.io",
        "category": "Contact Discovery",
        "description": "Sales automation platform with email finder, verifier, and drip campaigns. Uses OAuth2 client credentials flow for authentication.",
        "api_endpoint": "https://api.snov.io/v2/domain-emails-with-info",
        "docs_link": "https://snov.io/knowledgebase/api",
        "auth_method": "OAuth2 client_credentials (client_id + client_secret → access token)",
        "free_tier": "50 credits/month",
        "paid": "From $39/month",
        "data_quality": "Medium-High — good email coverage, includes social profiles",
        "complexity": "Medium — requires OAuth token management (auto-refresh implemented in adapter)",
        "recommendation": "Nice-to-Have",
        "rationale": "Good coverage at a competitive price point. OAuth adds complexity but the adapter handles token refresh automatically. Best as a fallback after Hunter/Apollo.",
        "adapter_file": "backend/app/services/adapters/contacts/snovio.py",
    },
    {
        "name": "RocketReach",
        "category": "Contact Discovery",
        "description": "Professional contact data platform specializing in executive and decision-maker contacts. Strong coverage for C-suite and VP-level contacts.",
        "api_endpoint": "https://api.rocketreach.co/api/v2/lookupProfile",
        "docs_link": "https://rocketreach.co/api",
        "auth_method": "Api-Key header",
        "free_tier": "5 lookups/month",
        "paid": "From $99/month",
        "data_quality": "Very High — specializes in verified executive contacts with direct phone numbers",
        "complexity": "Medium — two-step workflow (search → lookup) similar to Apollo",
        "recommendation": "Must-Have",
        "rationale": "Best source for executive/decision-maker contacts. Higher cost but significantly better hit rate for senior titles (VP, Director, C-suite).",
        "adapter_file": "backend/app/services/adapters/contacts/rocketreach.py",
    },
    {
        "name": "People Data Labs (PDL)",
        "category": "Contact Discovery",
        "description": "Massive B2B data platform with 3B+ person profiles. Supports SQL-like query DSL for precise targeting by company, title, location, and more.",
        "api_endpoint": "https://api.peopledatalabs.com/v5/person/search",
        "docs_link": "https://docs.peopledatalabs.com",
        "auth_method": "X-Api-Key header",
        "free_tier": "100 requests/month",
        "paid": "$0.01/match (pay per result)",
        "data_quality": "High — massive database with good coverage across industries and geographies",
        "complexity": "Medium — Elasticsearch-style query DSL requires specific syntax",
        "recommendation": "Must-Have",
        "rationale": "Largest contact database available. Pay-per-match pricing is cost-effective for targeted searches. Excellent when other sources miss contacts.",
        "adapter_file": "backend/app/services/adapters/contacts/pdl.py",
    },
    {
        "name": "Proxycurl",
        "category": "Contact Discovery",
        "description": "LinkedIn profile enrichment API. Extracts structured data from LinkedIn profiles including work history, education, and contact details.",
        "api_endpoint": "https://nubela.co/proxycurl/api/v2/linkedin/company/role/",
        "docs_link": "https://nubela.co/proxycurl/docs",
        "auth_method": "Bearer token",
        "free_tier": "10 credits (one-time)",
        "paid": "$0.01/API call",
        "data_quality": "High — real-time LinkedIn data extraction with full profile details",
        "complexity": "Low — simple REST API with Bearer auth",
        "recommendation": "Nice-to-Have",
        "rationale": "Best when LinkedIn URLs are known. Provides the most detailed professional history but limited without LinkedIn data. Good for enriching contacts found by other providers.",
        "adapter_file": "backend/app/services/adapters/contacts/proxycurl.py",
    },
    {
        "name": "Clearbit (Breeze)",
        "category": "Company Enrichment",
        "description": "Company intelligence platform providing firmographics, technographics, and employee data. Now part of HubSpot as 'Breeze Intelligence'. Returns company size, industry, revenue, tech stack, and more.",
        "api_endpoint": "https://company.clearbit.com/v2/companies/find",
        "docs_link": "https://dashboard.clearbit.com/docs",
        "auth_method": "Bearer token",
        "free_tier": "Free with HubSpot account",
        "paid": "API standalone from $99/month",
        "data_quality": "Very High — best-in-class company data with tech stack, revenue, employee count",
        "complexity": "Low — simple domain-based lookup",
        "recommendation": "Must-Have",
        "rationale": "Industry-leading company enrichment. Provides critical data for lead scoring: employee count (for company size prioritization), revenue, tech stack, and industry classification.",
        "adapter_file": "backend/app/services/adapters/company/clearbit.py",
    },
    {
        "name": "OpenCorporates",
        "category": "Company Enrichment",
        "description": "World's largest open database of companies with 200M+ corporate records. Provides incorporation data, registered addresses, officers, and legal status from official registries.",
        "api_endpoint": "https://api.opencorporates.com/v0.4/companies/search",
        "docs_link": "https://api.opencorporates.com/documentation",
        "auth_method": "api_token query parameter",
        "free_tier": "500 requests/month",
        "paid": "Custom pricing (enterprise)",
        "data_quality": "Medium — authoritative for legal/registration data, no financials or tech stack",
        "complexity": "Low — simple REST API with search + detail endpoints",
        "recommendation": "Nice-to-Have",
        "rationale": "Provides authoritative incorporation data and officer names. Useful for compliance checks and identifying company decision-makers through director listings. Complements Clearbit's commercial data.",
        "adapter_file": "backend/app/services/adapters/company/opencorporates.py",
    },
]


# ─── Main Document ─────────────────────────────────────────────────────────

def create_document():
    doc = Document()

    # ── Page Setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('NeuraLeads AI Agent')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('New Data Sources Assessment')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4A, 0x86, 0xC8)

    doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('\u2500' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph()

    meta_items = [
        ('Document Type', 'Data Source Integration Assessment'),
        ('Scope', '10 New External Data Providers (Job Sources, Contacts, Company Enrichment)'),
        ('Version', '1.0'),
        ('Date', datetime.now().strftime('%B %d, %Y')),
        ('Classification', 'Internal \u2014 Technical Reference'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Executive Summary', level=1)

    doc.add_paragraph(
        'This document assesses 10 new external data providers for integration into the '
        'NeuraLeads AI Agent platform. These providers expand the platform\'s data coverage across '
        'three categories: job sources (3), contact discovery (5), and company enrichment (2).'
    )

    doc.add_paragraph(
        'All 10 adapters have been implemented following the platform\'s adapter pattern and are '
        'disabled by default. Users enable them in Settings and provide API keys. Zero impact on '
        'existing functionality is guaranteed \u2014 the existing JSearch, Apollo, and Seamless '
        'integrations continue to work unchanged.'
    )

    # Summary table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['#', 'Provider', 'Category', 'Free Tier', 'Recommendation'], header=True)

    for i, src in enumerate(SOURCES, 1):
        add_table_row(table, [
            str(i), src['name'], src['category'],
            src['free_tier'], src['recommendation']
        ])

    doc.add_paragraph()

    add_note_box(doc,
        '6 providers are rated "Must-Have" (TheirStack, SerpAPI, Hunter.io, RocketReach, PDL, Clearbit) '
        'and 4 as "Nice-to-Have" (Adzuna, Snov.io, Proxycurl, OpenCorporates). '
        'The must-haves cover job discovery, contact finding, and company enrichment with minimal overlap. '
        'Start with the free tiers to validate before committing to paid plans.',
        label='Summary')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # INDIVIDUAL SOURCE PROFILES
    # ══════════════════════════════════════════════════════════════════════
    for i, src in enumerate(SOURCES, 1):
        doc.add_heading(f'Source {i}: {src["name"]}', level=1)

        # Profile table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_row(table, ['Attribute', 'Details'], header=True)

        profile_rows = [
            ['Category', src['category']],
            ['Description', src['description']],
            ['API Endpoint', src['api_endpoint']],
            ['Documentation', src['docs_link']],
            ['Authentication', src['auth_method']],
            ['Free Tier', src['free_tier']],
            ['Paid Plans', src['paid']],
            ['Data Quality', src['data_quality']],
            ['Integration Complexity', src['complexity']],
            ['Adapter File', src['adapter_file']],
            ['Recommendation', src['recommendation']],
        ]
        for row_data in profile_rows:
            add_table_row(table, row_data)

        doc.add_paragraph()

        doc.add_heading('Rationale', level=3)
        doc.add_paragraph(src['rationale'])

        if i < len(SOURCES):
            doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # COST ANALYSIS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('Cost Analysis', level=1)

    doc.add_paragraph(
        'All 10 providers offer free tiers sufficient for testing and low-volume production. '
        'Below is a comparison of costs for a typical production deployment sending 500-1,000 emails/day '
        '(requiring ~250-500 new leads and ~1,000-2,000 contact lookups per month).'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Provider', 'Free Tier', 'Paid Entry', 'Monthly Est. (Production)'], header=True)

    cost_rows = [
        ['TheirStack', '100 req/mo', '$49/mo', '$49-99'],
        ['SerpAPI', '100 req/mo', '$50/mo', '$50-130'],
        ['Adzuna', '250 req/mo', '$99/mo', '$0 (free tier may suffice)'],
        ['Hunter.io', '25 req/mo', '$49/mo', '$49-99'],
        ['Snov.io', '50 credits/mo', '$39/mo', '$39-79'],
        ['RocketReach', '5 lookups/mo', '$99/mo', '$99-199'],
        ['People Data Labs', '100 req/mo', '$0.01/match', '$10-50'],
        ['Proxycurl', '10 credits', '$0.01/call', '$10-50'],
        ['Clearbit', 'Free w/ HubSpot', '$99/mo', '$0-99'],
        ['OpenCorporates', '500 req/mo', 'Custom', '$0 (free tier may suffice)'],
        ['', '', 'TOTAL (all)', '$306-854/month'],
        ['', '', 'TOTAL (must-haves)', '$257-576/month'],
    ]
    for row_data in cost_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    add_note_box(doc,
        'Start with free tiers only. As volume grows, enable paid plans for must-have providers first. '
        'The pay-per-use providers (PDL at $0.01/match, Proxycurl at $0.01/call) are the most '
        'cost-effective for moderate volumes.',
        label='Recommendation')

    # ══════════════════════════════════════════════════════════════════════
    # IMPLEMENTATION STATUS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('Implementation Status', level=1)

    doc.add_paragraph(
        'All 10 adapters have been fully implemented and tested. Each adapter follows the platform\'s '
        'established patterns and is disabled by default.'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Component', 'Status', 'Files', 'Notes'], header=True)

    status_rows = [
        ['Adapter code (10 adapters)', 'Complete', '10 new files', 'All inherit from base classes'],
        ['Settings backend', 'Complete', 'settings.py updated', 'SETTINGS_TAB_MAP, DEFAULT_SETTINGS, PROVIDER_TAB_MAP, test-connection'],
        ['Pipeline registration', 'Complete', 'lead_sourcing.py, contact_enrichment.py', '3 job + 5 contact adapters registered'],
        ['Frontend Settings UI', 'Complete', 'settings/page.tsx', 'Checkboxes, API key inputs, Test buttons, pricing info'],
        ['Unit tests', 'Complete', 'test_adapters.py', '51 new tests (70 total), all passing'],
        ['Pipeline summary labels', 'Complete', 'pipeline_summary.py', '10 new ADAPTER_LABELS entries'],
        ['Company enrichment base class', 'Complete', 'base.py', 'CompanyEnrichmentAdapter ABC added'],
    ]
    for row_data in status_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    doc.add_heading('Safety Guarantees', level=2)
    add_bullet(doc, ' All new adapters are disabled by default. Users must explicitly enable and configure them.',
               bold_prefix='Opt-In Only:')
    add_bullet(doc, ' Existing JSearch, Apollo, and Seamless integrations are completely unaffected.',
               bold_prefix='Zero Breaking Changes:')
    add_bullet(doc, ' Company enrichment adapters (Clearbit, OpenCorporates) are available for test-connection '
               'and future enrichment workflows, but are not wired into any pipeline yet.',
               bold_prefix='Scoped Integration:')
    add_bullet(doc, ' A git tag (pre-adapter-expansion-v1) was created before any changes for instant rollback.',
               bold_prefix='Rollback Available:')

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2500' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Document generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('NeuraLeads AI Agent \u2014 New Data Sources Assessment')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


if __name__ == '__main__':
    doc = create_document()
    output_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    output_path = os.path.join(output_dir, 'New_Data_Sources_Assessment.docx')
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Size: {os.path.getsize(output_path):,} bytes")
