"""Generate Email Deliverability & Anti-AI Detection Gap Analysis document (.docx).

Identifies 12 gaps in the current email deliverability and anti-AI-detection posture
of the Exzelon RA Agent, grouped by priority (CRITICAL/HIGH/MEDIUM/LOW) with
technical remediation plans, effort estimates, and affected files.

Run: python scripts/generate_deliverability_gaps_doc.py
Produces: Email_Deliverability_Gaps_Analysis.docx in project root.
"""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# -- Color Constants --
CRITICAL_BG = "FEE2E2"
CRITICAL_FG = "991B1B"
HIGH_BG = "FEF3C7"
HIGH_FG = "92400E"
MEDIUM_BG = "DBEAFE"
MEDIUM_FG = "1E40AF"
LOW_BG = "DCFCE7"
LOW_FG = "166534"

WHITE = "FFFFFF"
DARK = "1F2937"
GRAY = "6B7280"
HEADER_BG = "1E3A5F"
ALT_ROW = "F3F4F6"
GREEN = "22C55E"
RED = "EF4444"
BLUE = "3B82F6"


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a table with colored header row and alternating row backgrounds."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header, bold=True, color=WHITE, size=10)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        bg = WHITE if row_idx % 2 == 0 else ALT_ROW
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            set_cell_text(cell, str(value), size=9)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)

    doc.add_paragraph()
    return table


def add_priority_badge(paragraph, priority):
    """Add a colored priority label as a bold run."""
    color_map = {
        "CRITICAL": CRITICAL_FG,
        "HIGH": HIGH_FG,
        "MEDIUM": MEDIUM_FG,
        "LOW": LOW_FG,
    }
    run = paragraph.add_run(f"  [{priority}]")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color_map.get(priority, GRAY))
    run.font.size = Pt(11)


def add_gap_section(doc, gap_num, title, priority, current_state, impact, fix, files, effort):
    """Add a complete gap section with all subsections."""
    # Gap heading
    heading = doc.add_heading(f"Gap {gap_num}: {title}", level=2)
    # Add priority badge after heading
    p = doc.add_paragraph()
    run = p.add_run("Priority: ")
    run.bold = True
    run.font.size = Pt(11)
    bg_map = {
        "CRITICAL": (CRITICAL_BG, CRITICAL_FG),
        "HIGH": (HIGH_BG, HIGH_FG),
        "MEDIUM": (MEDIUM_BG, MEDIUM_FG),
        "LOW": (LOW_BG, LOW_FG),
    }
    bg, fg = bg_map.get(priority, (ALT_ROW, DARK))
    run = p.add_run(f" {priority} ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(fg)
    run.font.size = Pt(11)

    # Current State
    doc.add_heading("Current State", level=3)
    doc.add_paragraph(current_state)

    # Impact
    doc.add_heading("Impact", level=3)
    p = doc.add_paragraph()
    run = p.add_run(impact)
    run.italic = True

    # Suggested Fix
    doc.add_heading("Suggested Fix", level=3)
    doc.add_paragraph(fix)

    # Files Affected
    doc.add_heading("Files Affected", level=3)
    for f in files:
        doc.add_paragraph(f, style="List Bullet")

    # Effort Estimate
    p = doc.add_paragraph()
    run = p.add_run("Effort Estimate: ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(effort)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    run.bold = True

    # Separator
    doc.add_paragraph("_" * 80)


def generate():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ======================================================================
    # TITLE PAGE
    # ======================================================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading("Email Deliverability & Anti-AI Detection", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(26)

    title2 = doc.add_heading("Gap Analysis", level=0)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title2.runs:
        run.font.size = Pt(26)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Exzelon RA Agent \u2014 Comprehensive Audit & Remediation Plan")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("April 2026")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_page_break()

    # ======================================================================
    # TABLE OF CONTENTS
    # ======================================================================
    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        "1. Executive Summary",
        "2. Gap 1: Real-Time Bounce-Back Parsing (CRITICAL)",
        "3. Gap 2: ESP Feedback Loop (FBL) Integration (CRITICAL)",
        "4. Gap 3: Engagement-Based Reputation Tracking (CRITICAL)",
        "5. Gap 4: No Anti-AI-Detection Measures (HIGH)",
        "6. Gap 5: AI Content Prompts Lack Anti-Detection Specifics (HIGH)",
        "7. Gap 6: IP Rotation Not Implemented (MEDIUM)",
        "8. Gap 7: No DKIM Signing In-App (MEDIUM)",
        "9. Gap 8: Domain Reputation Module Empty (MEDIUM)",
        "10. Gap 9: Tracking Pixel Fragile (MEDIUM)",
        "11. Gap 10: No Send-Time Optimization (LOW)",
        "12. Gap 11: No Email Client Rendering Preview (LOW)",
        "13. Gap 12: Spintax is Deterministic (LOW)",
        "14. Summary Table",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ======================================================================
    # EXECUTIVE SUMMARY
    # ======================================================================
    doc.add_heading("1. Executive Summary", level=1)

    # Overall score
    p = doc.add_paragraph()
    run = p.add_run("Overall Deliverability & Anti-AI-Detection Score: ")
    run.bold = True
    run.font.size = Pt(13)
    run = p.add_run("6.5 / 10")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(RED)

    doc.add_paragraph()

    doc.add_paragraph(
        "This audit evaluates the Exzelon RA Agent's email deliverability infrastructure and "
        "its resilience against AI-content detection systems deployed by major ESPs (Gmail, "
        "Outlook, Yahoo). While the platform has strong foundations \u2014 DNS health checking, "
        "warmup engine, spam word detection, and domain throttling \u2014 there are 12 identified "
        "gaps that collectively limit deliverability ceiling and increase the risk of inbox "
        "placement degradation at scale."
    )

    doc.add_paragraph()

    doc.add_paragraph(
        "The 12 gaps are grouped into four priority tiers based on their impact on sender "
        "reputation, inbox placement, and detection risk:"
    )

    doc.add_paragraph()

    # Priority breakdown
    priority_summary = [
        ("CRITICAL (3 gaps)", "Gaps 1-3", "Directly destroy sender reputation if unaddressed. "
         "Hard bounces, invisible spam complaints, and static health scoring create a "
         "compounding reputation death spiral."),
        ("HIGH (2 gaps)", "Gaps 4-5", "AI-generated email detection is an emerging threat. "
         "ESPs increasingly fingerprint AI-written content. Without humanization and "
         "anti-detection prompt engineering, deliverability will degrade over the next 6-12 months."),
        ("MEDIUM (4 gaps)", "Gaps 6-9", "Infrastructure limitations that cap deliverability "
         "at scale. IP rotation, DKIM signing, ISP-specific warmup, and engagement signal "
         "triangulation are table-stakes for high-volume senders."),
        ("LOW (3 gaps)", "Gaps 10-12", "Quality-of-life improvements that improve open rates "
         "and reduce detection surface. Important but not urgent."),
    ]

    for label, gaps, description in priority_summary:
        p = doc.add_paragraph()
        run = p.add_run(f"{label} \u2014 {gaps}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(description)

    doc.add_paragraph()

    # Effort summary
    p = doc.add_paragraph()
    run = p.add_run("Total Remediation Effort: ")
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run("79 hours")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_paragraph()

    # Effort breakdown table
    add_styled_table(doc,
        ["Priority", "Gap Count", "Effort (hours)", "Target Score Impact"],
        [
            ["CRITICAL", "3", "24h", "+1.5 points (6.5 \u2192 8.0)"],
            ["HIGH", "2", "16h", "+0.5 points (8.0 \u2192 8.5)"],
            ["MEDIUM", "4", "26h", "+1.0 points (8.5 \u2192 9.5)"],
            ["LOW", "3", "13h", "+0.5 points (9.5 \u2192 10.0)"],
        ],
        col_widths=[3, 3, 3, 8]
    )

    doc.add_page_break()

    # ======================================================================
    # GAP SECTIONS (1-12)
    # ======================================================================

    # -- GAP 1 --
    doc.add_heading("2. Gap 1: Real-Time Bounce-Back Parsing", level=1)
    add_gap_section(doc,
        gap_num=1,
        title="Real-Time Bounce-Back Parsing",
        priority="CRITICAL",
        current_state=(
            "Bounce status exists in the database but is only populated by pre-send email "
            "validation (NeverBounce, ZeroBounce, etc.). Actual SMTP 5xx bounce responses "
            "returned during the send operation are logged as generic errors in the application "
            "logs but do not automatically populate the suppression list, update the contact's "
            "validation status, or adjust the sender mailbox's health score."
        ),
        impact=(
            "Continuing to send to hard-bounced addresses is the single fastest way to destroy "
            "sender reputation. One hard bounce from a well-known ESP triggers IP reputation "
            "flags. Repeated 5xx responses to the same address cause the sending IP to be "
            "blocklisted. Without automated suppression, the system will keep attempting delivery "
            "to known-bad addresses on subsequent campaign steps, compounding the damage."
        ),
        fix=(
            "Parse SMTP response codes in send_outreach_email() within the outreach pipeline. "
            "On 5xx (permanent failure): auto-add the recipient to the suppression list, mark "
            "the contact's email_validation_status as 'invalid', increment the mailbox's "
            "bounce_count, and create an OutreachEvent with bounced status and the bounce reason. "
            "On 4xx (temporary failure): implement a retry queue with exponential backoff "
            "(3 attempts over 24 hours), then suppress after the third failure. Add a "
            "bounce_rate metric to the mailbox health score calculation."
        ),
        files=[
            "services/pipelines/outreach.py \u2014 SMTP response code parsing in send function",
            "db/models/suppression.py \u2014 auto-add on hard bounce",
            "db/models/sender_mailbox.py \u2014 bounce_count tracking, health score adjustment",
            "services/mailbox_selector.py \u2014 incorporate bounce rate into selection scoring",
        ],
        effort="8 hours"
    )

    doc.add_page_break()

    # -- GAP 2 --
    doc.add_heading("3. Gap 2: ESP Feedback Loop (FBL) Integration", level=1)
    add_gap_section(doc,
        gap_num=2,
        title="ESP Feedback Loop (FBL) Integration",
        priority="CRITICAL",
        current_state=(
            "If a recipient marks an email as spam in Gmail, Outlook, or Yahoo, the platform "
            "has zero visibility into this action. There is no complaint tracking mechanism, "
            "no FBL header processing, and no complaint rate monitoring. The system continues "
            "sending from the same mailbox to the same domain, accumulating invisible complaints "
            "that silently degrade sender reputation."
        ),
        impact=(
            "Gmail's postmaster guidelines state that a complaint rate above 0.3% triggers "
            "bulk folder placement for all mail from that sender. Outlook and Yahoo have similar "
            "thresholds. Without FBL integration, there is no way to detect when a mailbox is "
            "approaching or exceeding these thresholds. By the time deliverability drops are "
            "noticed via reduced open rates, the reputation damage is already severe and takes "
            "weeks to recover."
        ),
        fix=(
            "Create services/esp_feedback.py with complaint rate tracking per mailbox. Add a "
            "complaint_count column on SenderMailbox with daily reset tracking. Implement "
            "auto-pause logic: when complaint_rate exceeds a configurable threshold (default "
            "0.25%), automatically pause the mailbox and flag it for review. Integrate complaint "
            "rate into the mailbox health scoring formula in mailbox_selector.py. For Gmail "
            "specifically, monitor Google Postmaster Tools API if credentials are configured. "
            "For other ESPs, implement ARF (Abuse Reporting Format) email parsing for mailboxes "
            "registered with ISP FBL programs."
        ),
        files=[
            "services/esp_feedback.py \u2014 NEW: complaint rate tracking + auto-pause",
            "db/models/sender_mailbox.py \u2014 complaint_count, complaint_rate columns",
            "services/mailbox_selector.py \u2014 penalize high complaint rate in health score",
            "services/warmup/scheduler.py \u2014 daily complaint count reset job",
        ],
        effort="10 hours"
    )

    doc.add_page_break()

    # -- GAP 3 --
    doc.add_heading("4. Gap 3: Engagement-Based Reputation Tracking", level=1)
    add_gap_section(doc,
        gap_num=3,
        title="Engagement-Based Reputation Tracking",
        priority="CRITICAL",
        current_state=(
            "Mailbox health scoring in mailbox_selector.py uses a static formula: "
            "health_score * 0.4 + quota_remaining * 0.3 + warmup_age * 0.15 + "
            "deliverability * 0.15. This formula does not incorporate real-time engagement "
            "data (open rates, reply rates, click rates) from actual campaign sends. A mailbox "
            "can score highly despite having terrible open rates because the scoring relies "
            "solely on configuration-level metrics."
        ),
        impact=(
            "ESPs like Gmail use engagement signals as the primary input for sender reputation. "
            "A mailbox with perfect DNS records and zero bounces but a 2% open rate is treated "
            "as a low-value sender. The current system cannot detect this scenario and will "
            "continue routing high-value sends through underperforming mailboxes, dragging down "
            "overall campaign performance. This creates a hidden deliverability ceiling that "
            "cannot be diagnosed from the existing metrics."
        ),
        fix=(
            "Track per-mailbox engagement metrics over a trailing 7-day window: open_rate_7d, "
            "reply_rate_7d, click_rate_7d. Query OutreachEvent grouped by mailbox_id for the "
            "last 7 days. Incorporate engagement into mailbox_selector.py as a new scoring "
            "factor (suggested weight: 0.25, reducing warmup_age and deliverability weights). "
            "Auto-throttle mailboxes with open_rate_7d < 5% by reducing their daily send limit "
            "to 50%. Add a 'reputation_warning' flag visible in the mailbox management UI."
        ),
        files=[
            "services/mailbox_selector.py \u2014 engagement factor in health score formula",
            "db/models/sender_mailbox.py \u2014 cached engagement rate columns",
            "services/warmup/scheduler.py \u2014 daily engagement rate recalculation job",
            "api/endpoints/mailboxes.py \u2014 expose engagement metrics in API response",
        ],
        effort="6 hours"
    )

    doc.add_page_break()

    # -- GAP 4 --
    doc.add_heading("5. Gap 4: No Anti-AI-Detection Measures", level=1)
    add_gap_section(doc,
        gap_num=4,
        title="No Anti-AI-Detection Measures",
        priority="HIGH",
        current_state=(
            "AI-generated email content relies entirely on LLM \"naturalness\" plus spintax "
            "variation for text diversity. There is no explicit post-generation humanization "
            "pass. The current pipeline is: AI generates content \u2192 spintax applied \u2192 "
            "variables substituted \u2192 sent. No analysis of the output's \"AI fingerprint\" "
            "(uniform sentence lengths, consistent formality, absence of natural imperfections, "
            "predictable structure)."
        ),
        impact=(
            "Gmail and Outlook are increasingly deploying AI classifiers to detect and "
            "deprioritize machine-generated emails. Perfectly polished, zero-typo, "
            "consistent-structure emails are statistical outliers in human communication and "
            "trigger these classifiers. As detection models improve over the next 6-12 months, "
            "AI-generated cold emails without humanization will see progressively lower inbox "
            "placement rates. This is an emerging but accelerating threat."
        ),
        fix=(
            "Create services/email_humanizer.py as a post-generation humanization pass. "
            "The module should: (1) Measure burstiness \u2014 analyze sentence length variance "
            "and inject variation if too uniform. (2) Add natural imperfections \u2014 occasional "
            "em-dash instead of comma, sentence fragments, conversational asides. "
            "(3) Vary paragraph length \u2014 mix 1-sentence and 3-sentence paragraphs. "
            "(4) Randomize formality \u2014 alternate between formal and casual register within "
            "the email. (5) Score the output on an 'AI detection likelihood' scale before "
            "sending. Wire the humanizer into campaign_engine.py after AI content generation "
            "and before spintax application. Also integrate into email_preview_service.py so "
            "preview drafts show humanized output."
        ),
        files=[
            "services/email_humanizer.py \u2014 NEW: post-generation humanization engine",
            "services/campaign_engine.py \u2014 wire humanizer after AI content generation",
            "services/pipelines/outreach.py \u2014 humanize pipeline-generated emails",
            "services/email_preview_service.py \u2014 humanize preview drafts",
        ],
        effort="12 hours"
    )

    doc.add_page_break()

    # -- GAP 5 --
    doc.add_heading("6. Gap 5: AI Content Prompts Lack Anti-Detection Specifics", level=1)
    add_gap_section(doc,
        gap_num=5,
        title="AI Content Prompts Lack Anti-Detection Specifics",
        priority="HIGH",
        current_state=(
            "Current AI content generation prompts include generic instructions like \"write "
            "in a natural and human-sounding tone\" but do not include specific anti-detection "
            "strategies. The prompt_registry.py templates focus on personalization, AIDA "
            "structure, and business relevance but not on evading AI content classifiers."
        ),
        impact=(
            "LLMs produce detectable patterns when given generic instructions: uniform sentence "
            "length distributions, consistent formality level throughout, absence of colloquial "
            "contractions, predictable paragraph structure (intro \u2192 value prop \u2192 CTA), "
            "and over-polished grammar. These patterns are increasingly recognizable by ESP "
            "classifiers trained on millions of AI-generated marketing emails."
        ),
        fix=(
            "Enhance prompt engineering across all email generation prompts with specific "
            "anti-detection directives: \"Write as a busy sales rep typing quickly between "
            "meetings.\" \"Vary sentence length dramatically \u2014 mix 4-word sentences with "
            "20-word compound sentences.\" \"Include one natural imperfection per email (a "
            "slight run-on, a casual aside in parentheses, or a self-correction).\" \"Never "
            "start with 'I hope this email finds you well' or any templated opener.\" \"Use "
            "first-person anecdotes where relevant.\" \"Occasionally use contractions (don't, "
            "can't, we're).\" Add these as a shared anti-detection prompt fragment that is "
            "injected into all outreach-related prompts."
        ),
        files=[
            "services/ai_sales_agent/prompt_registry.py \u2014 add anti-detection prompt fragment",
            "adapters/ai_content.py \u2014 inject anti-detection directives in content generation",
            "services/email_preview_service.py \u2014 use updated prompts for AI rewrite",
        ],
        effort="4 hours"
    )

    doc.add_page_break()

    # -- GAP 6 --
    doc.add_heading("7. Gap 6: IP Rotation Not Implemented", level=1)
    add_gap_section(doc,
        gap_num=6,
        title="IP Rotation Not Implemented",
        priority="MEDIUM",
        current_state=(
            "The ip_rotation.py service module exists with a basic structure, and a "
            "dedicated_ip column exists on the SenderMailbox model. However, the column "
            "is never populated during mailbox setup, and the IP rotation logic is never "
            "invoked during send operations. All outbound emails route through whatever IP "
            "the SMTP provider assigns, with no application-level control."
        ),
        impact=(
            "High-volume senders (100+ emails/day across multiple mailboxes) benefit from "
            "distributing sends across multiple IP addresses to avoid single-IP reputation "
            "concentration. If one IP gets flagged, all mail from that IP is affected. Without "
            "rotation, a single spam complaint or blocklisting event can impact all outbound "
            "mail simultaneously."
        ),
        fix=(
            "Complete the ip_rotation.py implementation: (1) On mailbox creation/update, "
            "resolve the SMTP host to its IP address and populate dedicated_ip. (2) Add an "
            "ip_pool table or in-memory registry tracking all unique IPs across mailboxes. "
            "(3) Integrate into mailbox selection for high-volume campaigns: when daily send "
            "volume exceeds a threshold (configurable, default 50), prefer distributing across "
            "multiple IPs via round-robin within the health-weighted selection. (4) Track "
            "per-IP reputation metrics (bounce rate, complaint rate) to deprioritize "
            "underperforming IPs."
        ),
        files=[
            "services/ip_rotation.py \u2014 complete implementation",
            "services/mailbox_selector.py \u2014 integrate IP diversity into selection",
            "db/models/sender_mailbox.py \u2014 populate dedicated_ip on setup",
        ],
        effort="8 hours"
    )

    doc.add_page_break()

    # -- GAP 7 --
    doc.add_heading("8. Gap 7: No DKIM Signing In-App", level=1)
    add_gap_section(doc,
        gap_num=7,
        title="No DKIM Signing In-App",
        priority="MEDIUM",
        current_state=(
            "The DNS health checker in the warmup module reads and validates SPF, DKIM, and "
            "DMARC records for mailbox domains. However, the application itself does not "
            "perform DKIM signing of outbound emails. It relies 100% on the SMTP provider "
            "(Gmail, O365, or third-party SMTP relay) to handle DKIM signing during the SMTP "
            "handshake."
        ),
        impact=(
            "For mailboxes using Gmail or Office 365, this is not a problem \u2014 these "
            "providers sign all outbound mail automatically. However, for mailboxes using "
            "custom SMTP relay services (e.g., Amazon SES, Mailgun, or a self-hosted SMTP "
            "server), emails may go out unsigned. Unsigned emails from DMARC-enforcing domains "
            "will be rejected or sent to spam by receiving ESPs."
        ),
        fix=(
            "Create services/dkim_signer.py with optional DKIM signing for mailboxes "
            "configured with custom SMTP. Add dkim_private_key (encrypted) and dkim_selector "
            "columns on SenderMailbox. When both are populated, sign the outbound email using "
            "the dkim library before SMTP submission. Make this opt-in per mailbox (most "
            "users on Gmail/O365 will not need it). Add a UI indicator showing DKIM signing "
            "status per mailbox."
        ),
        files=[
            "services/dkim_signer.py \u2014 NEW: DKIM signing module using dkim library",
            "services/pipelines/outreach.py \u2014 sign outbound emails if configured",
            "db/models/sender_mailbox.py \u2014 dkim_private_key, dkim_selector columns",
            "api/endpoints/mailboxes.py \u2014 DKIM configuration endpoints",
        ],
        effort="6 hours"
    )

    doc.add_page_break()

    # -- GAP 8 --
    doc.add_heading("9. Gap 8: Domain Reputation Module Empty", level=1)
    add_gap_section(doc,
        gap_num=8,
        title="Domain Reputation Module Empty",
        priority="MEDIUM",
        current_state=(
            "The warmup/domain_reputation.py file exists but is a stub with placeholder "
            "functions. There are no ISP-specific warmup strategies. The warmup engine treats "
            "all recipient domains identically \u2014 a warmup email to a Gmail address follows "
            "the same volume ramp and timing as one to an Outlook or Yahoo address."
        ),
        impact=(
            "Gmail, Outlook, and Yahoo each use fundamentally different reputation algorithms. "
            "Gmail heavily weights engagement signals (opens, replies, not-spam actions). "
            "Outlook prioritizes authentication and IP reputation. Yahoo focuses on volume "
            "consistency and complaint rates. A one-size-fits-all warmup strategy is suboptimal "
            "for each and may actively harm reputation on ESPs where the wrong signals are "
            "being sent."
        ),
        fix=(
            "Implement ISP-specific warmup profiles in domain_reputation.py: "
            "(1) Gmail profile \u2014 engagement-focused: start with known-engaged recipients, "
            "prioritize reply generation, slower volume ramp but higher engagement targets. "
            "(2) Outlook profile \u2014 authentication-focused: verify DKIM/SPF/DMARC compliance "
            "first, steady volume ramp, monitor junk folder placement. "
            "(3) Yahoo profile \u2014 volume-consistency-focused: very gradual ramp, consistent "
            "daily volumes, avoid spikes. "
            "Classify warmup peers by their email domain and apply the appropriate profile."
        ),
        files=[
            "services/warmup/domain_reputation.py \u2014 complete ISP-specific profiles",
            "services/warmup/scheduler.py \u2014 use ISP profiles in peer selection",
        ],
        effort="8 hours"
    )

    doc.add_page_break()

    # -- GAP 9 --
    doc.add_heading("10. Gap 9: Tracking Pixel Fragile", level=1)
    add_gap_section(doc,
        gap_num=9,
        title="Tracking Pixel Fragile",
        priority="MEDIUM",
        current_state=(
            "Open tracking relies on a 1x1 transparent GIF pixel loaded from the tracking "
            "endpoint (/t/{id}/px.gif). This is the sole mechanism for detecting email opens. "
            "Reply detection and click tracking exist separately but are not combined into a "
            "unified engagement signal."
        ),
        impact=(
            "Gmail, Outlook, and Apple Mail increasingly block or proxy-cache images by "
            "default. Apple Mail's Mail Privacy Protection (MPP) pre-fetches all images, "
            "generating false opens. Gmail's image proxy caches pixels, so repeat opens are "
            "not counted. Outlook blocks images by default for unknown senders. Realistic "
            "open rate underreporting is 40-60%, making open rate an unreliable engagement "
            "metric for mailbox reputation decisions."
        ),
        fix=(
            "Instead of treating pixel loads as the definitive engagement signal, create a "
            "composite engagement_detected system: (1) Add an engagement_detected boolean "
            "on OutreachEvent that is set to True when ANY of: pixel loaded, link clicked, "
            "or reply received. (2) Use reply detection (already implemented in inbox_syncer) "
            "as the primary and most reliable engagement signal. (3) Weight engagement signals: "
            "reply = 1.0, click = 0.8, pixel = 0.5 (pixel is the least reliable). "
            "(4) Feed composite engagement into the per-mailbox engagement rate calculation "
            "from Gap 3."
        ),
        files=[
            "db/models/outreach_event.py \u2014 add engagement_detected column",
            "services/inbox_syncer.py \u2014 set engagement_detected on reply",
            "services/warmup/tracking.py \u2014 set engagement_detected on pixel + click",
            "services/mailbox_selector.py \u2014 use composite engagement in scoring",
        ],
        effort="4 hours"
    )

    doc.add_page_break()

    # -- GAP 10 --
    doc.add_heading("11. Gap 10: No Send-Time Optimization", level=1)
    add_gap_section(doc,
        gap_num=10,
        title="No Send-Time Optimization",
        priority="LOW",
        current_state=(
            "Campaign send windows are configurable per campaign (e.g., 9:00 AM - 5:00 PM) "
            "and timezone-aware using the contact's US state. However, there is no per-recipient "
            "send-time optimization based on historical engagement data. All contacts in the "
            "same timezone receive emails at whatever time the campaign processor happens to "
            "reach them within the window."
        ),
        impact=(
            "Emails sent at suboptimal times get buried under newer emails in the recipient's "
            "inbox, reducing open rates by 10-30%. For high-volume campaigns, this is a "
            "significant missed optimization. Competitors like Instantly.ai use machine learning "
            "to predict optimal send times per recipient."
        ),
        fix=(
            "Create services/send_time_optimizer.py: (1) Track the hour-of-day for each "
            "open/reply event per contact. (2) Build a per-contact preferred_send_hour based "
            "on historical engagement peaks. (3) For contacts with no history, fall back to "
            "industry-standard optimal times (Tuesday-Thursday, 9-11 AM recipient local time). "
            "(4) Integrate into campaign_engine.py by adjusting next_send_at to align with "
            "the contact's optimal hour within the campaign's send window."
        ),
        files=[
            "services/send_time_optimizer.py \u2014 NEW: per-contact optimal send time",
            "services/campaign_engine.py \u2014 use optimizer for next_send_at scheduling",
        ],
        effort="6 hours"
    )

    doc.add_page_break()

    # -- GAP 11 --
    doc.add_heading("12. Gap 11: No Email Client Rendering Preview", level=1)
    add_gap_section(doc,
        gap_num=11,
        title="No Email Client Rendering Preview",
        priority="LOW",
        current_state=(
            "The email preview system (email_preview_service.py) shows the email's HTML as-is "
            "in the preview modal. There are no warnings about HTML patterns that render "
            "differently across email clients (Gmail, Outlook desktop, Outlook web, Apple Mail, "
            "mobile clients). Users cannot see how their email will look in different clients "
            "before sending."
        ),
        impact=(
            "Emails that look broken in certain clients (especially Outlook desktop, which has "
            "notoriously poor CSS support) reduce engagement and increase spam reports. "
            "Recipients who see a malformed email are more likely to mark it as spam. This is "
            "a secondary deliverability concern but contributes to overall reputation."
        ),
        fix=(
            "Create services/rendering_checker.py that analyzes HTML email content and flags "
            "client-specific rendering issues: (1) CSS properties unsupported in Outlook "
            "desktop (flexbox, grid, max-width, background images). (2) Dark mode compatibility "
            "issues (forced color inversion, transparent backgrounds). (3) Image blocking "
            "warnings (alt text missing, image-heavy layouts). (4) Mobile responsiveness issues "
            "(fixed-width tables, small font sizes, unscalable images). "
            "Integrate into email_preview_service.py as a 'rendering_warnings' field in the "
            "preview response."
        ),
        files=[
            "services/rendering_checker.py \u2014 NEW: client-specific rendering analysis",
            "services/email_preview_service.py \u2014 add rendering_warnings to preview response",
            "api/endpoints/email_preview.py \u2014 expose warnings in preview endpoint",
        ],
        effort="6 hours"
    )

    doc.add_page_break()

    # -- GAP 12 --
    doc.add_heading("13. Gap 12: Spintax is Deterministic", level=1)
    add_gap_section(doc,
        gap_num=12,
        title="Spintax is Deterministic",
        priority="LOW",
        current_state=(
            "The spintax engine seeds its random number generator with the contact_id, ensuring "
            "the same contact always receives the same variation of spintax content. This was "
            "designed for consistency (same contact sees the same email if previewed multiple "
            "times), but it means that across multiple campaigns, the same contact receives "
            "identical spintax selections."
        ),
        impact=(
            "ESPs can detect when the same sender repeatedly sends near-identical content to "
            "the same recipient across different campaigns. This is a spam signal \u2014 "
            "legitimate senders vary their messaging. With deterministic spintax, a contact "
            "enrolled in Campaign A and Campaign B will receive the same word choices in both, "
            "making the emails appear templated to the ESP's content fingerprinting system."
        ),
        fix=(
            "Add campaign_id to the spintax seed so the same contact gets different variations "
            "across different campaigns. Change the seed from hash(contact_id) to "
            "hash(contact_id + campaign_id). This preserves within-campaign consistency "
            "(preview matches send) while ensuring cross-campaign variation. This is a "
            "minimal-effort, high-impact change."
        ),
        files=[
            "services/spintax.py \u2014 add campaign_id to random seed",
            "services/campaign_engine.py \u2014 pass campaign_id to spintax resolver",
        ],
        effort="1 hour"
    )

    doc.add_page_break()

    # ======================================================================
    # SUMMARY TABLE
    # ======================================================================
    doc.add_heading("14. Summary Table", level=1)

    doc.add_paragraph(
        "The following table provides a consolidated view of all 12 gaps with their "
        "priority, effort estimate, and current status."
    )

    doc.add_paragraph()

    # Create the summary table with priority-colored rows
    summary_headers = ["Gap #", "Title", "Priority", "Effort", "Status"]
    summary_rows = [
        (1, "Real-Time Bounce-Back Parsing", "CRITICAL", "8h", "Not Started"),
        (2, "ESP Feedback Loop (FBL) Integration", "CRITICAL", "10h", "Not Started"),
        (3, "Engagement-Based Reputation Tracking", "CRITICAL", "6h", "Not Started"),
        (4, "No Anti-AI-Detection Measures", "HIGH", "12h", "Not Started"),
        (5, "AI Content Prompts Lack Anti-Detection Specifics", "HIGH", "4h", "Not Started"),
        (6, "IP Rotation Not Implemented", "MEDIUM", "8h", "Not Started"),
        (7, "No DKIM Signing In-App", "MEDIUM", "6h", "Not Started"),
        (8, "Domain Reputation Module Empty", "MEDIUM", "8h", "Not Started"),
        (9, "Tracking Pixel Fragile", "MEDIUM", "4h", "Not Started"),
        (10, "No Send-Time Optimization", "LOW", "6h", "Not Started"),
        (11, "No Email Client Rendering Preview", "LOW", "6h", "Not Started"),
        (12, "Spintax is Deterministic", "LOW", "1h", "Not Started"),
    ]

    table = doc.add_table(rows=1 + len(summary_rows), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(summary_headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header, bold=True, color=WHITE, size=10)

    # Data rows with priority-colored backgrounds
    priority_bg_map = {
        "CRITICAL": CRITICAL_BG,
        "HIGH": HIGH_BG,
        "MEDIUM": MEDIUM_BG,
        "LOW": LOW_BG,
    }
    priority_fg_map = {
        "CRITICAL": CRITICAL_FG,
        "HIGH": HIGH_FG,
        "MEDIUM": MEDIUM_FG,
        "LOW": LOW_FG,
    }

    for row_idx, (gap_num, title, priority, effort, status) in enumerate(summary_rows):
        row = table.rows[row_idx + 1]

        # Gap number
        cell = row.cells[0]
        set_cell_shading(cell, WHITE)
        set_cell_text(cell, str(gap_num), bold=True, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

        # Title
        cell = row.cells[1]
        set_cell_shading(cell, WHITE)
        set_cell_text(cell, title, size=9)

        # Priority (colored)
        cell = row.cells[2]
        bg = priority_bg_map.get(priority, WHITE)
        fg = priority_fg_map.get(priority, DARK)
        set_cell_shading(cell, bg)
        set_cell_text(cell, priority, bold=True, color=fg, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

        # Effort
        cell = row.cells[3]
        set_cell_shading(cell, WHITE)
        set_cell_text(cell, effort, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Status
        cell = row.cells[4]
        set_cell_shading(cell, ALT_ROW)
        set_cell_text(cell, status, italic=True, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # Set column widths
    col_widths = [1.5, 8, 2.5, 2, 3]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph()

    # Totals
    total_effort = 8 + 10 + 6 + 12 + 4 + 8 + 6 + 8 + 4 + 6 + 6 + 1
    p = doc.add_paragraph()
    run = p.add_run(f"Total Effort: {total_effort} hours")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_paragraph()

    # Recommended implementation order
    doc.add_heading("Recommended Implementation Order", level=2)
    doc.add_paragraph(
        "Address gaps in priority order. Within the same priority tier, tackle the "
        "lowest-effort items first to build momentum and close quick wins."
    )

    order_items = [
        "Phase A (CRITICAL, 24h): Gap 3 (6h) \u2192 Gap 1 (8h) \u2192 Gap 2 (10h)",
        "Phase B (HIGH, 16h): Gap 5 (4h) \u2192 Gap 4 (12h)",
        "Phase C (MEDIUM, 26h): Gap 9 (4h) \u2192 Gap 7 (6h) \u2192 Gap 6 (8h) \u2192 Gap 8 (8h)",
        "Phase D (LOW, 13h): Gap 12 (1h) \u2192 Gap 10 (6h) \u2192 Gap 11 (6h)",
    ]
    for item in order_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        "Completing Phases A and B (40 hours) addresses all reputation-critical and "
        "AI-detection gaps, raising the score from 6.5/10 to approximately 8.5/10."
    )
    run.italic = True

    # ---- Save ----
    output_path = os.path.join(
        os.path.dirname(__file__), "..", "Email_Deliverability_Gaps_Analysis.docx"
    )
    output_path = os.path.abspath(output_path)
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate()
