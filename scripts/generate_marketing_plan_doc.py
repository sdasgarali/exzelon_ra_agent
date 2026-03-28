"""
Generate Commercial Marketing Website Plan as .docx
Usage: python scripts/generate_marketing_plan_doc.py
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
    return h

def add_table(doc, headers, rows, highlight_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if highlight_col is not None and c == highlight_col:
                        run.bold = True
    return table

def main():
    doc = Document()

    # Title
    title = doc.add_heading('NeuraLeads AI — Commercial Marketing Website Plan', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x0E, 0x27)
    doc.add_paragraph('Comprehensive plan for the high-conversion, SEO-optimized commercial marketing website.')
    doc.add_paragraph('')

    # ─── Overview ───
    add_heading(doc, '1. Overview')
    doc.add_paragraph(
        'Build a high-conversion, SEO-optimized commercial marketing website for NeuraLeads AI Agent within '
        'the existing Next.js 14 frontend. The site features animated landing pages, pricing tiers, feature '
        'showcases, and competitor comparisons — all designed to beat Instantly.ai, Smartlead, and Lemlist.'
    )

    # ─── Competitive Landscape ───
    add_heading(doc, '2. Competitive Landscape')
    add_table(doc,
        ['Competitor', 'Pricing', 'Model', 'Key Hook'],
        [
            ['Instantly.ai', '$30–$77.6/mo', 'Flat-fee, unlimited seats', '"Unlimited email accounts"'],
            ['Smartlead', '$39–$94/mo', 'Flat-fee', '"Unlimited warmup"'],
            ['Lemlist', '$69–$99/user/mo', 'Per-seat', '"Multichannel sequences"'],
            ['Woodpecker', '$29–$56/mo', 'Per-slot', '"Cold email that converts"'],
            ['Apollo.io', '$49–$119/user/mo', 'Per-seat', '"B2B database + outreach"'],
            ['Saleshandy', '$25–$219/mo', 'Flat-fee', '"Unlimited email accounts"'],
            ['Reply.io', '$49–$89/user/mo', 'Per-seat', '"AI SDR"'],
        ]
    )
    doc.add_paragraph('')

    # ─── NeuraLeads Advantages ───
    add_heading(doc, '3. NeuraLeads Advantages')
    advantages = [
        'Feature Score: 135/150 (90%) vs Instantly 131/150 (87%) — beats #1 competitor',
        'Self-hosted: Full data ownership, no per-seat fees',
        'Cost: $15–200/mo total vs $141–652/mo for competitors (70–90% savings)',
        '7 contact discovery providers (vs 1–2 for competitors)',
        '7 email validation providers (vs 1 built-in)',
        '4 AI engines (Groq, OpenAI, Anthropic, Gemini)',
        '10 job source APIs for lead sourcing',
        'Full pipeline automation from lead sourcing → contact enrichment → validation → outreach',
    ]
    for a in advantages:
        doc.add_paragraph(a, style='List Bullet')

    # ─── Architecture ───
    add_heading(doc, '4. Architecture')

    add_heading(doc, '4.1 Route Structure', level=2)
    routes = [
        ['(marketing)/layout.tsx', 'Marketing layout (navbar + footer)'],
        ['(marketing)/page.tsx', 'Landing/home page (root /)'],
        ['(marketing)/pricing/page.tsx', 'Pricing tiers page'],
        ['(marketing)/features/page.tsx', 'Feature deep-dive page'],
        ['(marketing)/compare/page.tsx', 'Competitor comparison page'],
    ]
    add_table(doc, ['Route File', 'Purpose'], routes)
    doc.add_paragraph('')

    add_heading(doc, '4.2 Component Structure', level=2)
    components = [
        ['Navbar.tsx', 'Sticky nav with logo, links, Login/Get Started CTA'],
        ['Footer.tsx', '4-column links, social, legal'],
        ['Hero.tsx', 'Animated hero with headline, subhead, CTA, stats'],
        ['FeatureGrid.tsx', '3-col animated feature cards (6 features)'],
        ['FeatureShowcase.tsx', 'Alternating image+text sections'],
        ['PricingCards.tsx', '3-tier pricing with annual toggle'],
        ['ComparisonTable.tsx', 'Feature-by-feature vs competitors'],
        ['StatsCounter.tsx', 'Animated counting numbers'],
        ['CTABanner.tsx', 'Full-width call-to-action section'],
        ['ROICalculator.tsx', 'Interactive savings calculator'],
        ['FAQAccordion.tsx', 'Expandable FAQ section'],
        ['ScrollReveal.tsx', 'Framer Motion scroll-triggered wrapper'],
    ]
    add_table(doc, ['Component', 'Description'], components)
    doc.add_paragraph('')

    add_heading(doc, '4.3 Dependencies', level=2)
    doc.add_paragraph('framer-motion — animations (scroll reveals, counters, hover effects)', style='List Bullet')
    doc.add_paragraph('All other dependencies already available (Tailwind, Lucide, Next.js 14)', style='List Bullet')

    # ─── Page Designs ───
    add_heading(doc, '5. Page Designs')

    add_heading(doc, '5.1 Landing Page (/) — 8 Sections', level=2)
    sections = [
        ['1. Hero', 'Gradient background, headline, subheadline, 2 CTAs, animated stats row'],
        ['2. Integration Bar', 'Auto-scrolling marquee of integration logos'],
        ['3. Feature Grid', '6 animated cards: Lead Intelligence, Campaigns, Inbox, Warmup, CRM, Analytics'],
        ['4. Feature Showcase', '3 alternating sections with mock screenshots and bullet points'],
        ['5. Comparison Snapshot', 'Mini comparison table (4 competitors, 8 categories)'],
        ['6. ROI Calculator', 'Interactive: team size, emails/day, current cost → projected savings'],
        ['7. Testimonials', '3 testimonial cards with star ratings'],
        ['8. Final CTA', '"Ready to Outperform Your Competition?" with large CTA button'],
    ]
    add_table(doc, ['Section', 'Description'], sections)
    doc.add_paragraph('')

    add_heading(doc, '5.2 Pricing Page (/pricing)', level=2)
    pricing_features = [
        ['Email accounts', '5', '25', 'Unlimited'],
        ['Emails/day', '500', '2,500', '10,000+'],
        ['Lead sources', '3', '7', 'All 10'],
        ['Contact providers', '2', '5', 'All 7'],
        ['AI engines', '1 (Groq)', '2', 'All 4'],
        ['Active campaigns', '3', '15', 'Unlimited'],
        ['A/B testing', '—', 'Yes', 'Yes'],
        ['Unified Inbox', '—', 'Yes', 'Yes'],
        ['CRM Deals', '—', 'Yes', 'Yes'],
        ['Analytics', 'Basic', 'Advanced', 'Full + API'],
        ['Warmup mailboxes', '5', '25', 'Unlimited'],
        ['Self-hosted option', '—', '—', 'Yes'],
        ['White-label', '—', '—', 'Yes'],
    ]
    add_table(doc, ['Feature', 'Starter ($49/mo)', 'Professional ($99/mo) ★', 'Enterprise ($199/mo)'], pricing_features, highlight_col=2)
    doc.add_paragraph('')

    add_heading(doc, '5.3 Features Page (/features)', level=2)
    doc.add_paragraph('Deep-dive into 10 major features with alternating layout, icons, and bullet points.')
    feature_list = [
        'Lead Sourcing Pipeline', 'Contact Discovery', 'Email Validation Engine',
        'Campaign Engine', 'Warmup Engine', 'Unified Inbox',
        'AI Content Generation', 'CRM Deal Pipeline', 'Analytics & Reporting',
        'Integrations & Webhooks',
    ]
    for f in feature_list:
        doc.add_paragraph(f, style='List Bullet')

    add_heading(doc, '5.4 Comparison Page (/compare)', level=2)
    doc.add_paragraph('Full feature-by-feature comparison: NeuraLeads vs Instantly vs Smartlead vs Lemlist.')
    doc.add_paragraph('8 categories with expandable rows, score badges, and total score summary.')

    # ─── SEO Strategy ───
    add_heading(doc, '6. SEO Strategy')

    add_heading(doc, '6.1 Technical SEO', level=2)
    seo_items = [
        'Next.js metadata exports on every page (title, description, openGraph, twitter)',
        'JSON-LD structured data (SoftwareApplication, AggregateOffer)',
        'Semantic HTML with proper heading hierarchy',
        'metadataBase set for Open Graph URL resolution',
        'Robots: index and follow all marketing pages',
    ]
    for s in seo_items:
        doc.add_paragraph(s, style='List Bullet')

    add_heading(doc, '6.2 Target Keywords', level=2)
    keywords = [
        'cold email software', 'email outreach platform', 'instantly alternative',
        'sales automation tool', 'B2B lead generation', 'self-hosted email outreach',
        'best cold email tool for small teams', 'cheaper than instantly.ai',
    ]
    for k in keywords:
        doc.add_paragraph(k, style='List Bullet')

    add_heading(doc, '6.3 Performance', level=2)
    doc.add_paragraph('Lazy load below-fold sections with Framer Motion scroll triggers', style='List Bullet')
    doc.add_paragraph('useReducedMotion respect for accessibility', style='List Bullet')
    doc.add_paragraph('CSS-only gradient backgrounds (no heavy images in hero)', style='List Bullet')

    # ─── CRO Elements ───
    add_heading(doc, '7. CRO (Conversion Rate Optimization) Elements')
    cro = [
        'Sticky nav with "Get Started" always visible',
        'Hero CTA above fold (primary + secondary)',
        'Animated stats counters (social proof via numbers)',
        'Comparison table (competitive positioning)',
        'ROI calculator (personalized value proposition)',
        'Pricing page "Most Popular" nudge',
        'FAQ sections addressing objections',
        'Final CTA on every page',
        'Annual billing toggle (urgency + savings)',
        '"No credit card required" friction reducer',
    ]
    for c in cro:
        doc.add_paragraph(c, style='List Bullet')

    # ─── Files Created ───
    add_heading(doc, '8. Files Created/Modified')

    add_heading(doc, '8.1 New Files (17)', level=2)
    new_files = [
        'frontend/src/app/(marketing)/layout.tsx',
        'frontend/src/app/(marketing)/page.tsx',
        'frontend/src/app/(marketing)/pricing/page.tsx',
        'frontend/src/app/(marketing)/features/page.tsx',
        'frontend/src/app/(marketing)/compare/page.tsx',
        'frontend/src/components/marketing/Navbar.tsx',
        'frontend/src/components/marketing/Footer.tsx',
        'frontend/src/components/marketing/Hero.tsx',
        'frontend/src/components/marketing/FeatureGrid.tsx',
        'frontend/src/components/marketing/FeatureShowcase.tsx',
        'frontend/src/components/marketing/PricingCards.tsx',
        'frontend/src/components/marketing/ComparisonTable.tsx',
        'frontend/src/components/marketing/StatsCounter.tsx',
        'frontend/src/components/marketing/CTABanner.tsx',
        'frontend/src/components/marketing/ROICalculator.tsx',
        'frontend/src/components/marketing/FAQAccordion.tsx',
        'frontend/src/components/marketing/ScrollReveal.tsx',
    ]
    for f in new_files:
        doc.add_paragraph(f, style='List Bullet')

    add_heading(doc, '8.2 Modified Files (4)', level=2)
    modified = [
        'frontend/src/app/page.tsx — Removed (replaced by marketing route group)',
        'frontend/src/app/layout.tsx — Updated metadata with SEO fields',
        'frontend/package.json — Added framer-motion dependency',
        'frontend/tailwind.config.js — Added navy colors, marquee & glow animations',
        'frontend/src/app/globals.css — Added marketing CSS classes',
    ]
    for m in modified:
        doc.add_paragraph(m, style='List Bullet')

    # ─── Implementation Notes ───
    add_heading(doc, '9. Implementation Notes')
    notes = [
        'No backend changes required — marketing pages are fully static (SSG-compatible)',
        'No database changes needed',
        'Existing auth/dashboard routes are completely untouched',
        'Marketing pages use the (marketing) route group for layout isolation',
        'All "Get Started" and "Login" CTAs point to /login (existing auth page)',
        'Framer Motion animations respect useReducedMotion for accessibility',
        'Build size: ~5KB for landing page, ~3KB for pricing, ~1KB for features/compare',
    ]
    for n in notes:
        doc.add_paragraph(n, style='List Bullet')

    # Save
    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                               'Marketing_Website_Plan.docx')
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    main()
