"""Generate Microsoft 365 OAuth Setup Guide as .docx"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def add_heading_with_number(doc, text, level=1):
    """Add a heading."""
    h = doc.add_heading(text, level=level)
    return h

def add_info_box(doc, text, bold_prefix=""):
    """Add a highlighted info paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 153)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0, 51, 153)
    return p

def create_guide():
    doc = Document()

    # Title
    title = doc.add_heading("Microsoft 365 OAuth2 Setup Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("NeuraLeads Email Integration")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # Overview
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "This guide walks you through connecting your Microsoft 365 mailboxes to NeuraLeads "
        "using OAuth2 authentication. Unlike password-based authentication, OAuth2 uses secure "
        "tokens that automatically refresh — meaning password changes will never break your "
        "email connection."
    )

    # Benefits box
    doc.add_heading("Why OAuth2?", level=2)
    benefits = [
        ("Password-Independent", "Changing your M365 password won't disconnect your mailbox"),
        ("More Secure", "No passwords stored — uses short-lived tokens with automatic refresh"),
        ("M365 Compliant", "Works even when Microsoft blocks Basic Authentication (legacy passwords)"),
        ("One-Click Setup", "Connect with a single click — no manual SMTP/IMAP configuration needed"),
    ]
    for title_text, desc in benefits:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(desc)

    # Prerequisites
    doc.add_heading("Prerequisites", level=1)
    prereqs = [
        "Microsoft 365 Business account (not free Outlook.com/Hotmail)",
        "Admin access to Microsoft Entra (Azure AD) admin center",
        "NeuraLeads platform access with admin privileges",
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style="List Bullet")

    # Part A: Azure AD App Registration
    doc.add_heading("Part A: Azure AD App Registration (One-Time Admin Setup)", level=1)
    doc.add_paragraph(
        "This section is for the IT administrator. It only needs to be done once — "
        "after setup, all users can connect their mailboxes with a single click."
    )

    # Step 1
    doc.add_heading("Step 1: Create the App Registration", level=2)
    steps = [
        "Go to https://portal.azure.com/#view/Microsoft_AAD_RegisteredApps/ApplicationsListBlade",
        "Or in the Entra admin center: search for \"App registrations\" in the top search bar",
        "Click \"+ New registration\"",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_paragraph("Fill in the registration form:")
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [
        ("Field", "Value"),
        ("Name", "NeuraLeads Email (or your preferred name)"),
        ("Supported account types", "Accounts in this organizational directory only (Single tenant)"),
        ("Redirect URI (Web)", "https://ra.partnerwithus.tech/dashboard/mailboxes"),
    ]
    for i, (field, value) in enumerate(headers):
        row = table.rows[i]
        row.cells[0].text = field
        row.cells[1].text = value
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph("Click \"Register\" to create the application.")

    # Step 2
    doc.add_heading("Step 2: Copy Application IDs", level=2)
    doc.add_paragraph("After registration, you'll see the Overview page. Copy these two values:")
    id_items = [
        ("Application (client) ID", "MS365_OAUTH_CLIENT_ID"),
        ("Directory (tenant) ID", "MS365_OAUTH_TENANT_ID"),
    ]
    for label, env_var in id_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}")
        run.bold = True
        p.add_run(f" → used as {env_var} in NeuraLeads configuration")

    # Step 3
    doc.add_heading("Step 3: Create a Client Secret", level=2)
    steps = [
        "In your app registration, go to \"Certificates & secrets\" in the left sidebar",
        "Click \"+ New client secret\"",
        "Description: \"NeuraLeads production\"",
        "Expiry: Choose 24 months (recommended)",
        "Click \"Add\"",
        "IMPORTANT: Immediately copy the \"Value\" column — it is shown only once!",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {step}")
        if "IMPORTANT" in step:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(204, 0, 0)

    add_info_box(doc, " This secret value is your MS365_OAUTH_CLIENT_SECRET. Store it securely.", "⚠ ")

    # Step 4
    doc.add_heading("Step 4: Add API Permissions", level=2)
    doc.add_paragraph("Your app needs permission to send emails and read inbox via SMTP/IMAP:")

    steps = [
        "Go to \"API permissions\" in the left sidebar",
        "Click \"+ Add a permission\"",
        "Select \"APIs my organization uses\" tab",
        "Search for and select \"Office 365 Exchange Online\"",
        "Choose \"Delegated permissions\"",
        "Check these two permissions:",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    # Permissions table
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    perm_headers = [("API", "Permission", "Type"),
                    ("Office 365 Exchange Online", "SMTP.Send", "Delegated"),
                    ("Office 365 Exchange Online", "IMAP.AccessAsUser.All", "Delegated")]
    for i, (api, perm, ptype) in enumerate(perm_headers):
        row = table.rows[i]
        row.cells[0].text = api
        row.cells[1].text = perm
        row.cells[2].text = ptype
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph("7. Click \"Add permissions\" to save")
    doc.add_paragraph("8. Click \"Grant admin consent for [Your Organization]\" (green checkmark button)")
    doc.add_paragraph("9. Verify all permissions show a green checkmark under \"Status\"")

    add_info_box(doc, " Microsoft Graph > offline_access and User.Read are usually added automatically. "
                 "If not, add them manually under Microsoft Graph > Delegated permissions.", "ℹ ")

    # Step 5
    doc.add_heading("Step 5: Enable Authenticated SMTP", level=2)
    doc.add_paragraph("Each mailbox must have SMTP AUTH enabled in Microsoft 365:")
    steps = [
        "Go to Microsoft 365 Admin Center (https://admin.microsoft.com)",
        "Navigate to Users > Active users",
        "Select the user/mailbox",
        "Click the \"Mail\" tab",
        "Under \"Email apps\", click \"Manage email apps\"",
        "Ensure \"Authenticated SMTP\" is checked/enabled",
        "Click \"Save changes\"",
        "Repeat for each mailbox you want to connect",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    add_info_box(doc, " If you don't see this option, your M365 plan may not support it, "
                 "or it may need to be enabled at the organization level first.", "⚠ ")

    # Part B: NeuraLeads Configuration
    doc.add_heading("Part B: NeuraLeads Configuration (One-Time Admin Setup)", level=1)
    doc.add_paragraph(
        "Add the following values to your NeuraLeads environment configuration (.env file):"
    )

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    config_rows = [
        ("Variable", "Value"),
        ("MS365_OAUTH_CLIENT_ID", "Your Application (client) ID from Step 2"),
        ("MS365_OAUTH_CLIENT_SECRET", "Your client secret Value from Step 3"),
        ("MS365_OAUTH_TENANT_ID", "Your Directory (tenant) ID from Step 2"),
        ("MS365_OAUTH_REDIRECT_URI", "https://ra.partnerwithus.tech/dashboard/mailboxes"),
    ]
    for i, (var, val) in enumerate(config_rows):
        row = table.rows[i]
        row.cells[0].text = var
        row.cells[1].text = val
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph("After adding these values, restart the NeuraLeads backend service.")

    # Part C: Connecting Mailboxes
    doc.add_heading("Part C: Connecting Mailboxes (Per User)", level=1)
    doc.add_paragraph(
        "Once the admin setup is complete, any user can connect their mailbox with OAuth2:"
    )

    steps = [
        "Log in to NeuraLeads and navigate to Dashboard > Mailboxes",
        "For an existing mailbox: click the mailbox row to open its details, then click \"Connect with Microsoft 365\"",
        "For a new mailbox: click \"+ Add Mailbox\", enter the email address, and select \"Connect with OAuth2\"",
        "You will be redirected to Microsoft's login page",
        "Sign in with the mailbox email and password (this is a one-time authentication)",
        "Review the permissions and click \"Accept\"",
        "You will be redirected back to NeuraLeads with a success confirmation",
        "The mailbox will now show \"OAuth Connected\" badge — you're done!",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    add_info_box(doc, " After OAuth connection, you can change your Microsoft 365 password freely — "
                 "the connection uses tokens that refresh automatically and independently of your password.", "✅ ")

    # Troubleshooting
    doc.add_heading("Troubleshooting", level=1)

    issues = [
        ("\"OAuth not configured\" error",
         "The MS365_OAUTH_CLIENT_ID is not set in the environment. "
         "Contact your NeuraLeads administrator to complete Part B."),
        ("\"AADSTS50011: Reply URL does not match\"",
         "The redirect URI in Azure AD doesn't match the NeuraLeads URL. "
         "Go to App Registration > Authentication and verify the redirect URI is exactly: "
         "https://ra.partnerwithus.tech/dashboard/mailboxes"),
        ("\"AADSTS7000215: Invalid client secret\"",
         "The client secret has expired or is incorrect. Create a new secret in Azure AD "
         "(Step 3) and update MS365_OAUTH_CLIENT_SECRET."),
        ("\"Insufficient privileges\" after consent",
         "Admin consent was not granted. Go to API permissions and click "
         "\"Grant admin consent for [Your Organization]\"."),
        ("Connection works but IMAP shows \"Basic Auth blocked\"",
         "IMAP over OAuth2 requires the IMAP.AccessAsUser.All permission. "
         "Verify it's added and admin-consented in API permissions."),
        ("Mailbox shows \"OAuth Connected\" but emails aren't sending",
         "Check that Authenticated SMTP is enabled for that specific user in M365 Admin Center "
         "(Part A, Step 5)."),
    ]

    for issue, solution in issues:
        p = doc.add_paragraph()
        run = p.add_run(f"Problem: {issue}")
        run.bold = True
        doc.add_paragraph(f"Solution: {solution}")
        doc.add_paragraph()  # spacing

    # FAQ
    doc.add_heading("FAQ", level=1)

    faqs = [
        ("How long does the OAuth connection last?",
         "Indefinitely, as long as the refresh token remains valid. Microsoft refresh tokens "
         "typically last 90 days but are automatically renewed each time they're used. Since "
         "NeuraLeads sends emails regularly, the token stays fresh."),
        ("What happens if I change my M365 password?",
         "Nothing! OAuth2 tokens are independent of your password. Your mailbox will continue "
         "to send and receive emails normally."),
        ("Do I need to re-authenticate periodically?",
         "Only if the refresh token expires (rare — usually only if the mailbox is inactive "
         "for 90+ days) or if the Azure AD admin revokes consent."),
        ("Can I switch from password to OAuth2 for an existing mailbox?",
         "Yes. Open the mailbox details and click \"Connect with Microsoft 365\". "
         "The authentication method will switch to OAuth2 automatically."),
        ("Is OAuth2 more secure than passwords?",
         "Yes. OAuth2 uses short-lived access tokens (1 hour) with automatic refresh. "
         "No password is stored in NeuraLeads — only encrypted tokens that can be revoked "
         "from the Azure AD admin center at any time."),
    ]

    for question, answer in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f"Q: {question}")
        run.bold = True
        doc.add_paragraph(f"A: {answer}")
        doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph("—")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Generated by NeuraLeads • For support, contact your system administrator")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # Save
    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                                "Microsoft_365_OAuth_Setup_Guide.docx")
    doc.save(output_path)
    print(f"Guide saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_guide()
