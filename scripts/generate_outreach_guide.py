"""Generate styled .docx guide: NeuraLeads Outreach & Email System — Complete Guide."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ──────────────────────────────────────────────────────────────────
# Color palette
# ──────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0F, 0x17, 0x2A)
DARK_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
PRIMARY   = RGBColor(0x4F, 0x46, 0xE5)   # Indigo-600
ACCENT    = RGBColor(0x7C, 0x3A, 0xED)   # Violet-600
GREEN     = RGBColor(0x05, 0x96, 0x69)
ORANGE    = RGBColor(0xEA, 0x58, 0x0C)
RED       = RGBColor(0xDC, 0x26, 0x26)
CYAN      = RGBColor(0x06, 0x91, 0xB7)
GRAY      = RGBColor(0x64, 0x74, 0x8B)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = "F0F4FF"
GREEN_BG  = "ECFDF5"
ORANGE_BG = "FFF7ED"
RED_BG    = "FEF2F2"
CYAN_BG   = "ECFEFF"
GRAY_BG   = "F1F5F9"
PURPLE_BG = "F5F3FF"

doc = Document()

# ──────────────────────────────────────────────────────────────────
# Styles
# ──────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = NAVY if level == 1 else (DARK_BLUE if level == 2 else PRIMARY)
    h.font.size = Pt(22 - level * 4)
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(6)

# ──────────────────────────────────────────────────────────────────
# Helper functions
# ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(headers, rows, col_widths=None, header_color="2D3A8C", alt_color=LIGHT_BG):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = WHITE
                r.font.size = Pt(10)
        set_cell_bg(cell, header_color)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
            if ri % 2 == 1:
                set_cell_bg(cell, alt_color)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl

def add_colored_box(text, bg_hex, text_color, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.color.rgb = text_color
        r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.color.rgb = text_color
    r.font.size = Pt(10)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_hex}" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(0.5)
    pf.right_indent = Cm(0.5)

def add_bullet(text, bold_prefix=None, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.2 + indent * 0.6)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)

def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)

def add_para(text, bold=False, color=None, size=10.5, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold:
        r.font.bold = True
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align


# ══════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════

# ── Title Page ────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
add_para("NeuraLeads", bold=True, color=PRIMARY, size=32, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Outreach & Email System", bold=True, color=DARK_BLUE, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Complete Guide", bold=True, color=ACCENT, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("How emails are composed, sent, and automated — explained for everyone.",
         color=GRAY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("April 2026  |  Version 1.0", color=GRAY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ── Table of Contents (manual) ───────────────────────────────────
doc.add_heading("Table of Contents", level=1)
toc = [
    "1. The Big Picture — Three Ways to Send Emails",
    "2. Method A: Outreach Pipeline (One-Shot Bulk Send)",
    "3. Method B: Campaigns (Multi-Step Automated Sequences)",
    "4. Method C: Mailmerge Export (Manual Send Outside System)",
    "5. Email Templates vs AI-Drafted Emails",
    "6. When Master Automation Is ON",
    "7. Manual vs Automated — Side-by-Side Comparison",
    "8. The Complete Email Journey (Step-by-Step)",
    "9. Safety Rules & Rate Limits",
    "10. Quick Reference — Where to Go for What",
    "11. Glossary of Terms",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_BLUE

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 1: Big Picture
# ══════════════════════════════════════════════════════════════════
doc.add_heading("1. The Big Picture — Three Ways to Send Emails", level=1)

add_para(
    "NeuraLeads gives you three distinct ways to reach out to prospects. Each serves "
    "a different purpose, and understanding when to use which is key to getting results.",
    size=11
)

add_styled_table(
    ["Method", "What It Does", "Best For", "Where to Find It"],
    [
        ["A) Outreach Pipeline\n(One-Shot Send)",
         "Sends one email to every eligible contact in a single batch run.",
         "Quick blast to new validated contacts. No follow-ups needed.",
         "Dashboard → Pipelines → Run Outreach\nOR\nDashboard → Outreach → Send Emails"],
        ["B) Campaigns\n(Multi-Step Sequences)",
         "Sends a series of timed emails (e.g., Day 1: Intro, Day 4: Follow-up, Day 10: Final).",
         "Proper outreach with follow-ups, A/B testing, and conditional branching.",
         "Dashboard → Campaigns → Create Campaign"],
        ["C) Mailmerge Export\n(CSV Download)",
         "Exports contacts to a CSV file. You send the emails yourself using Word or another tool.",
         "When you want full manual control, or your SMTP isn't configured.",
         "Dashboard → Pipelines → Run Outreach (Mailmerge mode)\nOR\nDashboard → Outreach (Mailmerge tab)"],
    ],
    col_widths=[3.5, 5.5, 4, 5],
    header_color="312E81"
)

doc.add_paragraph()
add_colored_box(
    " Campaigns (Method B) is the recommended approach for serious outreach. "
    "It handles follow-ups, tracks replies, and stops automatically when someone responds.",
    PURPLE_BG, ACCENT, bold_prefix="★ Recommendation:"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 2: Outreach Pipeline
# ══════════════════════════════════════════════════════════════════
doc.add_heading("2. Method A: Outreach Pipeline (One-Shot Bulk Send)", level=1)

doc.add_heading("What is it?", level=2)
add_para(
    "The Outreach Pipeline is Stage 4 of the four-stage data pipeline "
    "(Source Leads → Enrich Contacts → Validate Emails → Send Outreach). "
    "It sends a single email to every eligible contact in one batch. No follow-ups. "
    "Think of it as a one-time broadcast."
)

doc.add_heading("How to run it manually", level=2)

add_para("Option 1: From the Pipelines page", bold=True, color=PRIMARY)
add_numbered("Go to Dashboard → Pipelines.")
add_numbered('In the "Outreach" section, choose Send mode (not Mailmerge).')
add_numbered("Optionally toggle Dry Run ON to preview without actually sending.")
add_numbered('Click "Run Outreach". The system runs as a background job.')
add_numbered("Monitor progress on the Pipelines page (job status, sent count, errors).")

doc.add_paragraph()
add_para("Option 2: From the Outreach page", bold=True, color=PRIMARY)
add_numbered("Go to Dashboard → Outreach.")
add_numbered('Click "Send Emails" button (top-right).')
add_numbered("Set the limit (default: 30 — the daily cap per mailbox).")
add_numbered("Toggle Dry Run if you want a preview first.")
add_numbered("View sent/bounced/replied stats on the same page.")

doc.add_heading("What happens behind the scenes", level=2)
add_numbered("System finds all contacts with validation_status = \"Valid\" and not archived.")
add_numbered("For each contact, checks eligibility (not unsubscribed, not on suppression list, cooldown period respected, daily limit not exceeded).")
add_numbered("Picks the healthiest available mailbox (uses a scoring formula based on health, quota remaining, warmup age, and deliverability).")
add_numbered("Composes the email: uses your active Email Template if one exists → falls back to AI draft → falls back to a basic hardcoded template.")
add_numbered("Applies spintax variations (e.g., {Hi|Hello|Hey} becomes one random choice per contact).")
add_numbered("Fills in placeholders: {{contact_first_name}}, {{company_name}}, {{job_title}}, etc.")
add_numbered("Adds unsubscribe footer with a unique tracking link.")
add_numbered("Sends via SMTP with a random 45–180 second delay between emails (to look human).")
add_numbered("Records an OutreachEvent for tracking (sent/bounced/replied status).")

doc.add_heading("Limitations", level=2)
add_bullet("Only sends one email per contact — no automatic follow-ups.", bold_prefix="No follow-ups: ")
add_bullet("Once sent, you'd have to manually re-run for another round.", bold_prefix="Manual re-run: ")
add_bullet("30 emails per mailbox per day (configurable).", bold_prefix="Daily cap: ")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 3: Campaigns
# ══════════════════════════════════════════════════════════════════
doc.add_heading("3. Method B: Campaigns (Multi-Step Automated Sequences)", level=1)

add_colored_box(
    " This is the most powerful and recommended way to do outreach in NeuraLeads.",
    GREEN_BG, GREEN, bold_prefix="✓ RECOMMENDED:"
)

doc.add_heading("What is it?", level=2)
add_para(
    "A Campaign is a multi-step email sequence that sends a series of emails to enrolled "
    "contacts over days or weeks. Each step can be an email, a wait period, or a conditional "
    "check (e.g., \"did they reply?\"). The system handles everything automatically — sending, "
    "timing, follow-ups, and stopping when someone responds."
)

doc.add_heading("Example Campaign Sequence", level=2)
add_styled_table(
    ["Step", "Type", "Delay", "What Happens"],
    [
        ["1", "Email", "Immediately", "Introduction email (\"Hi Sarah, I noticed your company is hiring...\")"],
        ["2", "Wait", "3 days", "System waits 3 days before next step"],
        ["3", "Condition", "—", "Check: Did they reply? YES → stop. NO → continue."],
        ["4", "Email", "Immediately", "Follow-up email (\"Just wanted to circle back on my last email...\")"],
        ["5", "Wait", "5 days", "Wait 5 more days"],
        ["6", "Email", "Immediately", "Final email (\"This is my last follow-up...\")"],
    ],
    col_widths=[1.5, 2, 2.5, 12],
    header_color="7C3AED"
)

doc.add_heading("How to create a Campaign manually", level=2)
add_numbered("Go to Dashboard → Campaigns → click \"New Campaign\".")
add_numbered("Fill in the campaign name, select mailboxes to send from, set send window (e.g., 9 AM – 5 PM weekdays).")
add_numbered("Add sequence steps: Email steps (write subject + body or use AI to generate), Wait steps (set delay in days), Condition steps (branch based on reply/open).")
add_numbered("For each Email step, you can create A/B variants (e.g., two different subject lines — the system will randomly assign contacts and track which performs better).")
add_numbered("Go to the Contacts tab → click \"Enroll Contacts\" → select contacts from your validated list.")
add_numbered("Alternatively, set up Auto-Enrollment Rules (e.g., automatically enroll contacts with validation_status=valid and priority=P1).")
add_numbered("Click \"Activate\" to start the campaign.")
add_numbered("The system's scheduler (runs every 2 minutes) picks up due contacts and sends their current step's email.")

doc.add_heading("How sending works inside a Campaign", level=2)
add_bullet("Every 2 minutes, the Campaign Engine checks: which contacts are due for their next email?", bold_prefix="Scheduler: ")
add_bullet("Checks the send window — only sends during the hours you configured (respects contact timezone).", bold_prefix="Send window: ")
add_bullet("Picks the best mailbox using health-aware scoring (same logic as the pipeline).", bold_prefix="Mailbox selection: ")
add_bullet("If the step has A/B variants, picks one based on weighted randomization (seeded by contact ID for consistency).", bold_prefix="A/B testing: ")
add_bullet("Applies spintax, fills placeholders, adds unsubscribe footer.", bold_prefix="Content: ")
add_bullet("Sends via SMTP with 45–180 second random delays between contacts.", bold_prefix="Sending: ")
add_bullet("After sending, advances the contact to the next step and calculates when the next email is due.", bold_prefix="Advancement: ")
add_bullet("When a contact replies → the system automatically marks them as REPLIED and stops sending them further emails in that campaign.", bold_prefix="Reply handling: ")
add_bullet("When a contact bounces → marked as BOUNCED and excluded from further sends.", bold_prefix="Bounce handling: ")

doc.add_heading("Key Campaign Features", level=2)
add_styled_table(
    ["Feature", "Description"],
    [
        ["A/B Testing", "Create 2+ subject/body variants per step. System auto-assigns and tracks open/reply rates. Auto-optimizes winner using chi-squared test."],
        ["Conditional Branching", "\"If replied → go to step X, else → go to step Y.\" Supports reply and open conditions."],
        ["Auto-Enrollment", "Define rules (priority, validation status, state, job title, source) and the system automatically enrolls matching new contacts daily."],
        ["Send Window", "Set hours + days of week. System respects contact's timezone (US state-based)."],
        ["Spintax", "Write {Hello|Hi|Hey} and each contact gets a random variation — makes emails unique."],
        ["Deal Integration", "When an email is sent, it auto-logs activity on associated CRM deals and can auto-advance deal stages."],
        ["Duplicate Protection", "A contact cannot be enrolled twice in the same campaign."],
        ["Suppression List", "Contacts on the suppression list are automatically skipped."],
    ],
    col_widths=[3.5, 14.5],
    header_color="7C3AED"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 4: Mailmerge
# ══════════════════════════════════════════════════════════════════
doc.add_heading("4. Method C: Mailmerge Export (Manual Send Outside System)", level=1)

doc.add_heading("What is it?", level=2)
add_para(
    "Mailmerge mode exports your validated contacts into a CSV file that you can download. "
    "You then use Microsoft Word, Google Sheets, or any mail-merge tool to send the emails yourself. "
    "No emails are sent by NeuraLeads — it only prepares the data."
)

doc.add_heading("When to use it", level=2)
add_bullet("Your SMTP credentials aren't configured in NeuraLeads yet.")
add_bullet("You want to review and manually approve every email before it goes out.")
add_bullet("You use an external email tool (Outlook, Mailchimp, etc.) for sending.")
add_bullet("You want a CSV backup of your validated contacts for other purposes.")

doc.add_heading("How to run it", level=2)
add_numbered("Go to Dashboard → Pipelines → Outreach section.")
add_numbered("Select \"Mailmerge\" mode.")
add_numbered("Click Run. The system generates a CSV file in the data/exports/ folder.")
add_numbered("Download the CSV from the Outreach page.")

doc.add_heading("What's in the CSV", level=2)
add_para(
    "Contact first name, last name, email, job title, company name, state, phone, "
    "priority level, validation status — everything you need for a mail merge."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 5: Templates vs AI
# ══════════════════════════════════════════════════════════════════
doc.add_heading("5. Email Templates vs AI-Drafted Emails", level=1)

add_para(
    "This is one of the most common questions: \"When does NeuraLeads use my template, "
    "and when does it write the email with AI?\" Here's the clear answer:"
)

doc.add_heading("The Priority Order (Highest → Lowest)", level=2)

add_styled_table(
    ["Priority", "Source", "When It's Used", "Where to Configure"],
    [
        ["1 (Highest)", "Campaign Step Content",
         "If the email is sent by a Campaign, the subject + body you wrote in the Sequence Step is ALWAYS used. This is the most common case.",
         "Dashboard → Campaigns → Edit → Sequence Steps"],
        ["2", "Active Email Template",
         "If NOT a campaign (i.e., Outreach Pipeline send), the system looks for an Email Template with status = Active.",
         "Dashboard → Email Templates → Create/Activate"],
        ["3", "AI-Drafted Email",
         "If no active template exists, and the AI Outreach Drafting feature is enabled, the system asks AI (Groq/OpenAI/Anthropic/Gemini) to write a personalized email.",
         "Dashboard → Settings → AI provider + API key. Toggle: ai_outreach_drafting = true"],
        ["4 (Lowest)", "Hardcoded Fallback",
         "If everything above fails (no template, AI disabled or errored), a basic boilerplate email is used.",
         "Built into the system. Not configurable."],
    ],
    col_widths=[2, 3, 6, 5],
    header_color="0E7490"
)

doc.add_paragraph()
doc.add_heading("How Email Templates work", level=2)
add_bullet("Go to Dashboard → Email Templates to create templates.", bold_prefix="Create: ")
add_bullet("Use placeholders like {{contact_first_name}}, {{company_name}}, {{job_title}}, {{job_location}}, {{sender_first_name}}, {{signature}}, {{unsubscribe_link}}.", bold_prefix="Placeholders: ")
add_bullet("Mark a template as \"Active\" — only active templates are used by the Pipeline.", bold_prefix="Activate: ")
add_bullet("You can have multiple templates but only the active one is picked for Pipeline sends.", bold_prefix="One active: ")
add_bullet("Templates are NOT used by Campaigns — campaigns use their own step content.", bold_prefix="Not for Campaigns: ")

doc.add_heading("How AI-Drafted Emails work", level=2)
add_bullet("Enable the feature: Settings → set ai_outreach_drafting = true.", bold_prefix="Enable: ")
add_bullet("Configure an AI provider: Groq (default), OpenAI, Anthropic, or Google Gemini + API key.", bold_prefix="Provider: ")
add_bullet("When triggered, the AI receives the contact's name, title, company, industry, and job context. It generates a personalized subject + body.", bold_prefix="Personalization: ")
add_bullet("The AI also researches the company (cached) if the context is thin — fetching industry, size, and description.", bold_prefix="Research: ")
add_bullet("AI drafting is only used by the Outreach Pipeline (Method A), never by Campaigns.", bold_prefix="Pipeline only: ")

doc.add_heading("How Campaign Email Content works", level=2)
add_bullet("When you create a Campaign, you write the subject and body for each Email step directly.", bold_prefix="You write it: ")
add_bullet("You can also click \"Generate with AI\" to have the AI Sequence Generator create an entire multi-step sequence for you (goal, product, tone, number of steps).", bold_prefix="AI Sequence Generator: ")
add_bullet("Once generated (or written), the content is saved in the SequenceStep. The Campaign Engine uses that exact content — it does NOT re-generate with AI at send time.", bold_prefix="Saved content: ")
add_bullet("The system still processes spintax and fills in placeholders ({{contact_first_name}}, etc.) at send time.", bold_prefix="Dynamic parts: ")

doc.add_paragraph()
add_colored_box(
    " Email Templates are for Pipeline sends. Campaign step content is for Campaign sends. "
    "AI generation happens either at Pipeline send time (if no template) or at Campaign creation time "
    "(AI Sequence Generator). The Campaign Engine never calls AI during actual sending.",
    CYAN_BG, CYAN, bold_prefix="Key Takeaway: "
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 6: Master Automation
# ══════════════════════════════════════════════════════════════════
doc.add_heading("6. When Master Automation Is ON", level=1)

add_para(
    "NeuraLeads has a Master Automation toggle (Settings → Automation → Master Enabled). "
    "When ON, the system runs scheduled jobs automatically without you clicking anything."
)

doc.add_heading("What runs automatically", level=2)

add_styled_table(
    ["Job", "Schedule", "What It Does"],
    [
        ["Lead Sourcing", "Every 4 hours\n(0, 4, 8, 12, 16, 20 UTC)", "Scrapes job boards (Indeed, LinkedIn, Glassdoor, etc.) for new leads. Deduplicates against existing data."],
        ["Auto-Chain\n(Enrich → Validate → Enroll)", "After each sourcing run", "Automatically runs contact enrichment on new leads, then validates their emails, then enrolls eligible contacts into active campaigns."],
        ["Campaign Processor", "Every 2 minutes", "Checks all active campaigns. Sends due emails. Handles waits, conditions, and step advancement."],
        ["Auto-Enrollment", "Every 30 minutes", "Checks active campaigns with enrollment rules. Enrolls new matching contacts automatically."],
        ["Reply Check", "Every 15 minutes\n(8 AM – 7 PM UTC)", "Checks mailbox inboxes for replies. Updates OutreachEvents and marks campaign contacts as REPLIED."],
        ["Inbox Sync", "Every 10 minutes", "Syncs replies into the Unified Inbox with AI sentiment analysis."],
        ["Daily Count Reset", "Midnight UTC", "Resets each mailbox's emails_sent_today counter back to 0."],
        ["Warmup Engine", "Every 30 min + daily", "Sends/replies to warmup emails between mailboxes. DNS checks. Blacklist monitoring."],
        ["Lead Scoring", "Daily at 2 AM", "Recalculates lead priority scores."],
        ["CRM Sync", "Daily at 1 AM", "Syncs contacts/deals with HubSpot or Salesforce (if configured)."],
        ["Backup", "Daily at 2 AM", "Database backup + old backup cleanup."],
    ],
    col_widths=[3, 3.5, 11.5],
    header_color="0F766E"
)

doc.add_heading("The Automated Workflow (End to End)", level=2)
add_para("When Master Automation is ON, here's what happens with zero human intervention:", bold=True)

doc.add_paragraph()
add_numbered("Every 4 hours, the system sources new leads from job boards.", bold_prefix="Source: ")
add_numbered("Immediately after sourcing, it enriches contacts for those leads (finds emails and names via Apollo, Seamless, etc.).", bold_prefix="Enrich: ")
add_numbered("Then validates the discovered emails (NeverBounce, ZeroBounce, etc.).", bold_prefix="Validate: ")
add_numbered("Then enrolls valid contacts into active campaigns that have auto-enrollment rules matching them.", bold_prefix="Enroll: ")
add_numbered("Every 2 minutes, the campaign engine sends due emails from active campaigns.", bold_prefix="Send: ")
add_numbered("Every 15 minutes, the reply checker looks for responses.", bold_prefix="Reply: ")
add_numbered("Replied contacts are automatically removed from the campaign (no more follow-ups). Their messages appear in the Inbox.", bold_prefix="Stop: ")
add_numbered("Meanwhile, the warmup engine keeps your mailbox reputation healthy.", bold_prefix="Warmup: ")

doc.add_paragraph()
add_colored_box(
    " Even with automation ON, you still have full control. You can pause/resume campaigns, "
    "disable individual jobs (e.g., keep sourcing manual but automate campaigns), or turn off "
    "the master toggle entirely. Every job checks the master toggle + its own toggle before running.",
    GREEN_BG, GREEN, bold_prefix="Control: "
)

doc.add_heading("How to toggle automation", level=2)
add_bullet("Go to Dashboard → Settings → Automation tab.", bold_prefix="Master toggle: ")
add_bullet("automation_master_enabled = true/false (controls all jobs).", bold_prefix="Setting: ")
add_bullet("Each job also has its own toggle: automation_{job_name}_enabled.", bold_prefix="Per-job: ")
add_bullet("Example: Keep campaign_processor enabled but lead_sourcing_run disabled = campaigns send on schedule, but no new leads are sourced automatically.", bold_prefix="Mix & match: ")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 7: Side-by-Side Comparison
# ══════════════════════════════════════════════════════════════════
doc.add_heading("7. Manual vs Automated — Side-by-Side Comparison", level=1)

add_styled_table(
    ["Aspect", "Manual Approach", "Automated (Master ON)"],
    [
        ["Lead Sourcing", "Click \"Run Lead Sourcing\" on Pipelines page", "Runs automatically every 4 hours"],
        ["Contact Enrichment", "Select leads → click \"Enrich Contacts\"", "Auto-chains after each sourcing run"],
        ["Email Validation", "Select contacts → click \"Validate Emails\"", "Auto-chains after enrichment"],
        ["Campaign Creation", "You create campaign + write steps manually", "Same — campaigns are always created manually"],
        ["Contact Enrollment", "Select contacts → \"Enroll\" in campaign", "Auto-enrollment based on rules you define"],
        ["Email Sending (Campaign)", "Activate campaign → scheduler sends automatically", "Same — campaign engine runs every 2 min either way"],
        ["Email Sending (Pipeline)", "Click \"Send Emails\" on Outreach page", "Not automated — Pipeline sends are always manual"],
        ["Reply Detection", "Click \"Check Replies\" on Outreach page", "Automatic every 15 minutes"],
        ["Template/AI Usage", "You control via Templates page + Settings", "Same — no difference"],
    ],
    col_widths=[3.5, 7, 7.5],
    header_color="1E3A5F"
)

doc.add_paragraph()
add_colored_box(
    " Campaign email sending is ALWAYS automated once you activate the campaign "
    "(the scheduler runs every 2 minutes regardless of the master toggle). "
    "The master toggle controls pipeline stages (sourcing, enrichment, validation) "
    "and supporting jobs (reply check, inbox sync, warmup, etc.).",
    ORANGE_BG, ORANGE, bold_prefix="Important Note: "
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 8: Complete Email Journey
# ══════════════════════════════════════════════════════════════════
doc.add_heading("8. The Complete Email Journey (Step-by-Step)", level=1)

add_para("Here's the full lifecycle of an outreach email, from job posting to closed deal:", bold=True, size=11)

steps = [
    ("Job Board", "A company posts a job on Indeed/LinkedIn/Glassdoor.",
     "Lead Sourcing pipeline scrapes it and creates a LeadDetails record."),
    ("Company Identified", "The company is saved as a ClientInfo record (deduplicated by name + LinkedIn URL).",
     "Industry, size, website, and headquarters are enriched."),
    ("Contact Discovered", "Contact Enrichment finds the hiring manager (e.g., via Apollo or Seamless AI).",
     "Creates ContactDetails with email, name, title, and priority (P1–P5)."),
    ("Email Validated", "Email Validation checks the contact's email (NeverBounce, ZeroBounce, etc.).",
     "Status set to: Valid, Invalid, Catch-All, or Unknown. Only Valid contacts get emails."),
    ("Enrolled in Campaign", "Contact is added to an active Campaign (manually or via auto-enrollment rules).",
     "CampaignContact record created with status=ACTIVE, pointing to step 1."),
    ("Email Composed", "Campaign Engine loads the step's subject + body. Applies spintax, fills placeholders, "
     "adds signature and unsubscribe footer.",
     "If using Pipeline instead: checks Template → AI → fallback."),
    ("Mailbox Selected", "System picks the healthiest mailbox (score = health × 0.4 + quota × 0.3 + warmup age × 0.15 + deliverability × 0.15).",
     "Only mailboxes with warmup status COLD_READY or ACTIVE, connection=successful, and quota remaining."),
    ("Email Sent", "SMTP connection established. Email sent with Message-ID header. Random 45–180s delay to next send.",
     "OutreachEvent created (status=SENT). Mailbox counters updated. Contact.last_outreach_date set."),
    ("Tracking", "Open pixel and link tracking embedded. Unsubscribe link with HMAC-signed token.",
     "If contact clicks unsubscribe → outreach_status = UNSUBSCRIBED, all future sends blocked."),
    ("Reply Detected", "Reply checker scans mailbox IMAP every 15 min. Matches reply to original Message-ID.",
     "CampaignContact.status = REPLIED. No more follow-ups. Message appears in Unified Inbox."),
    ("AI Categorized", "Inbox sync runs AI sentiment analysis: interested, not_interested, OOO, question, referral, do_not_contact.",
     "Category stored on InboxMessage. Interested replies flagged for follow-up."),
    ("Deal Created", "Interested prospect → create Deal in CRM pipeline (New Lead → Contacted → Qualified → Proposal → Won).",
     "Deal auto-logged with email activity. Can auto-advance stages."),
]

for i, (title, desc, detail) in enumerate(steps, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"Step {i}: {title}")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = PRIMARY
    add_bullet(desc)
    add_bullet(detail, indent=1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 9: Safety Rules
# ══════════════════════════════════════════════════════════════════
doc.add_heading("9. Safety Rules & Rate Limits", level=1)

add_para(
    "NeuraLeads has multiple layers of protection to keep your email reputation safe "
    "and comply with anti-spam regulations."
)

add_styled_table(
    ["Rule", "Value", "Purpose"],
    [
        ["Daily Send Limit", "30 emails / mailbox / day", "Prevents spam flags. Each mailbox tracks its own count, reset at midnight UTC."],
        ["Cooldown Period", "10 days between emails to same contact", "Prevents harassment. No contact receives more than 1 email per 10 days."],
        ["Send Delay", "45–180 seconds between emails", "Randomized delay makes sending pattern look human, not automated."],
        ["Hourly Cap", "daily_limit ÷ 8 per hour", "Spreads sends throughout the day instead of blasting all at once."],
        ["Contact Cap per Lead", "4 contacts per company per job", "Prevents over-contacting a single company."],
        ["Send Window", "Configurable (default 9 AM – 5 PM)", "Only sends during business hours in the contact's timezone."],
        ["Send Days", "Mon–Fri (configurable)", "No weekend emails by default."],
        ["Suppression List", "Permanent or expiring blocks", "Contacts on suppression list are always skipped."],
        ["Unsubscribe", "One-click in every email", "Contact marked as UNSUBSCRIBED — permanently blocked from future sends."],
        ["Valid Emails Only", "validation_status must be \"Valid\"", "Invalid, catch-all, unknown, and pending emails are never sent to."],
        ["Warmup Status", "Mailbox must be COLD_READY or ACTIVE", "New mailboxes go through warmup before sending real outreach."],
        ["Connection Check", "Mailbox connection_status = successful", "Failed SMTP connections exclude the mailbox from selection."],
        ["Salary Filter", "$30,000+ only", "Filters out low-salary leads to focus on quality targets."],
        ["Industry Filter", "22 non-IT target industries", "IT roles and staffing agencies excluded."],
    ],
    col_widths=[3, 4, 11],
    header_color="991B1B"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 10: Quick Reference
# ══════════════════════════════════════════════════════════════════
doc.add_heading("10. Quick Reference — Where to Go for What", level=1)

add_styled_table(
    ["I Want To...", "Go Here", "Notes"],
    [
        ["Send a one-time email blast", "Dashboard → Outreach → Send Emails\nor Pipelines → Run Outreach (Send mode)", "Uses active Email Template → AI → fallback"],
        ["Create a multi-step follow-up sequence", "Dashboard → Campaigns → New Campaign", "Recommended for serious outreach"],
        ["Write my own email content", "Dashboard → Campaigns → Edit → Steps\nor Email Templates", "Campaigns: per-step. Pipeline: via Templates."],
        ["Let AI write my emails", "Dashboard → Campaigns → \"Generate with AI\"\nor Settings → ai_outreach_drafting=true", "Campaigns: at design time. Pipeline: at send time."],
        ["Export contacts to CSV", "Dashboard → Pipelines → Run Outreach (Mailmerge mode)", "No emails sent — just CSV download"],
        ["See what was sent", "Dashboard → Outreach → Event list", "Filter by status, channel, date"],
        ["See replies", "Dashboard → Inbox", "AI-categorized: interested, OOO, etc."],
        ["Check email deliverability", "Dashboard → Mailboxes → Health score\nDashboard → Warmup Engine", "SPF/DKIM/DMARC checks, blacklist monitoring"],
        ["Turn on full automation", "Dashboard → Settings → Automation → Master Enabled = true", "Sources leads → enriches → validates → enrolls → sends"],
        ["Stop all automation", "Dashboard → Settings → Automation → Master Enabled = false", "Campaigns still send (disable per-campaign)"],
        ["Pause a specific campaign", "Dashboard → Campaigns → Select → Pause", "Contacts stay enrolled, just paused"],
    ],
    col_widths=[4.5, 5.5, 8],
    header_color="312E81"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# Section 11: Glossary
# ══════════════════════════════════════════════════════════════════
doc.add_heading("11. Glossary of Terms", level=1)

terms = [
    ("Outreach Pipeline", "Stage 4 of the data pipeline. Sends a single email to eligible contacts in bulk."),
    ("Campaign", "A multi-step email sequence with timed delays, conditions, and follow-ups."),
    ("Sequence Step", "One step in a campaign: Email (send a message), Wait (pause N days), or Condition (branch based on reply/open)."),
    ("Campaign Contact", "A contact enrolled in a specific campaign, tracking their current step and next send time."),
    ("OutreachEvent", "A record of a single email action (sent, opened, clicked, replied, bounced)."),
    ("Email Template", "A reusable email layout with placeholders ({{contact_first_name}}, etc.). Used by the Outreach Pipeline."),
    ("Spintax", "Text variation syntax: {Hello|Hi|Hey} randomly picks one option per contact."),
    ("A/B Variant", "Two or more versions of an email step in a campaign. System assigns contacts randomly and tracks which performs better."),
    ("Mailbox Health Score", "A 0–100 score calculated from: email health (40%), quota remaining (30%), warmup age (15%), deliverability (15%)."),
    ("Warmup", "Process of gradually increasing send volume on a new mailbox to build reputation. Peer-to-peer warmup emails between your own mailboxes."),
    ("Suppression List", "A list of email addresses that should never receive outreach (permanent or with expiry date)."),
    ("Cooldown", "The minimum number of days (default 10) before the same contact can receive another email."),
    ("Auto-Enrollment", "Rules on a campaign that automatically enroll new contacts matching criteria (priority, validation, state, etc.)."),
    ("Auto-Chain", "When automation is ON, after lead sourcing completes, the system automatically runs enrichment → validation → enrollment."),
    ("Master Automation Toggle", "A global ON/OFF switch in Settings that controls all scheduled automation jobs."),
    ("AI Sequence Generator", "A tool that uses AI to create an entire multi-step campaign sequence (subjects, bodies, delays) from a goal description."),
    ("Unified Inbox", "Central hub for all email replies across all mailboxes and campaigns, with AI sentiment analysis."),
    ("Deal", "A sales opportunity tracked through pipeline stages (New Lead → Contacted → Qualified → Proposal → Won/Lost)."),
]

for term, definition in terms:
    p = doc.add_paragraph()
    r = p.add_run(f"{term}: ")
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = PRIMARY
    r = p.add_run(definition)
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


# ── Save ──────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        "NeuraLeads_Outreach_Email_Guide.docx")
doc.save(out_path)
print(f"Document saved: {out_path}")
print(f"Sections: 11 | Pages: ~20 | Tables: 8")
