"""
Generate SME Q&A Document for NeuraLeads AI Agent
Creates a professional .docx with detailed answers to 6 key business questions.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# ─── Helpers ───────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells, bold=False, header=False, shade=None):
    """Add a row to a table with optional formatting."""
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, '1F4E79')
        elif shade:
            set_cell_shading(cell, shade)
    return row

def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_note_box(doc, text, label="Note"):
    """Add a highlighted note/callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ─── Main Document ─────────────────────────────────────────────────────────

def create_document():
    doc = Document()

    # ── Page Setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('NeuraLeads AI Agent')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('SME Technical Assessment & Strategic Guidance')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4A, 0x86, 0xC8)

    doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph()

    meta_items = [
        ('Document Type', 'Subject Matter Expert (SME) Q&A Assessment'),
        ('System', 'NeuraLeads AI Agent — Cold Email Outreach Automation Platform'),
        ('Version', '1.0'),
        ('Date', datetime.now().strftime('%B %d, %Y')),
        ('Classification', 'Internal — Stakeholder Reference'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1.', 'Maximum Cold Email Sending Capacity'),
        ('2.', 'Outreach Load Balancing Across Mailboxes'),
        ('3.', 'Alternative Contact Discovery When Apollo Fails'),
        ('4.', 'Strategies to Increase Lead Volume'),
        ('5.', 'Best Practice for Sending 100,000 Cold Emails Daily'),
        ('6.', 'Reusing the System for Other Lines of Business'),
        ('', 'Appendix A — Configuration Quick Reference'),
        ('', 'Appendix B — Warmup Profile Comparison'),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        if num:
            run = p.add_run(f'  Question {num}  ')
            run.bold = True
            run.font.size = Pt(11)
        else:
            run = p.add_run('  ')
        run = p.add_run(title_text)
        run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # QUESTION 1: Maximum Cold Email Capacity
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Question 1: Maximum Cold Email Sending Capacity', level=1)
    doc.add_heading('With the current setup, what is the maximum number of cold emails that can be sent?', level=3)

    doc.add_heading('Answer', level=2)
    p = doc.add_paragraph(
        'The maximum daily cold email capacity depends on three factors: '
        'the number of active mailboxes, each mailbox\'s daily send limit, '
        'and the warmup maturity of those mailboxes. Here is the complete breakdown:'
    )

    # ── Per-Mailbox Limits ──
    doc.add_heading('1.1 Per-Mailbox Daily Limits', level=3)
    p = doc.add_paragraph(
        'Each sender mailbox has an independent daily send limit, defaulting to 30 emails/day. '
        'This is configurable per mailbox (the seeded admin mailbox defaults to 50/day). '
        'The system enforces this via the emails_sent_today counter, which resets at midnight UTC.'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Parameter', 'Default Value', 'Configurable?', 'Location'], header=True)
    data = [
        ['Daily Send Limit (per mailbox)', '30 emails', 'Yes (per mailbox)', 'Settings UI / DB'],
        ['Admin Mailbox Default', '50 emails', 'Yes', 'Seed data'],
        ['SMTP Send Rate', '30 emails/min (2s gap)', 'Yes', 'SMTP adapter'],
        ['SMTP Timeout', '30 seconds', 'Yes', 'constants.py'],
        ['Daily Counter Reset', 'Midnight UTC', 'No', 'Scheduler'],
    ]
    for row_data in data:
        add_table_row(table, row_data)

    # ── Warmup Constraints ──
    doc.add_heading('1.2 Warmup Constraints on New Mailboxes', level=3)
    p = doc.add_paragraph(
        'New mailboxes cannot send outreach immediately. They must complete a warmup profile '
        'before becoming eligible. The system offers three warmup profiles:'
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Profile', 'Duration', 'Start Volume', 'End Volume', 'Best For'], header=True)
    warmup_data = [
        ['Conservative', '45 days', '1-3/day', '15-25/day', 'Brand-new domains'],
        ['Standard (Default)', '30 days', '2-5/day', '25-35/day', 'Most use cases'],
        ['Aggressive', '20 days', '3-8/day', '35-50/day', 'Established domains'],
    ]
    for row_data in warmup_data:
        add_table_row(table, row_data)

    p = doc.add_paragraph(
        'After warmup completes, a mailbox enters COLD_READY status and must maintain a health score '
        'of 80+ for 7 consecutive days before being promoted to ACTIVE. Only COLD_READY and ACTIVE '
        'mailboxes can send outreach emails.'
    )

    # ── Capacity Formula ──
    doc.add_heading('1.3 Capacity Calculation Formula', level=3)

    p = doc.add_paragraph()
    run = p.add_run('Daily Capacity = Number of Active Mailboxes x Per-Mailbox Daily Limit')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Scenario', 'Mailboxes', 'Limit/Each', 'Daily Capacity'], header=True)
    scenarios = [
        ['Current Default (minimal)', '1', '30', '30 emails/day'],
        ['Small Team', '5', '30', '150 emails/day'],
        ['Medium Operation', '10', '40', '400 emails/day'],
        ['Scaled Operation', '20', '50', '1,000 emails/day'],
        ['Large Scale', '50', '50', '2,500 emails/day'],
        ['Enterprise (100 mailboxes)', '100', '50', '5,000 emails/day'],
    ]
    for row_data in scenarios:
        add_table_row(table, row_data)

    # ── Additional Constraints ──
    doc.add_heading('1.4 Additional Throttling Constraints', level=3)

    add_bullet(doc, ' 10-day cooldown between emails to the same contact (prevents spamming).',
               bold_prefix='Cooldown Period:')
    add_bullet(doc, ' Maximum 4 contacts emailed per company per job posting.',
               bold_prefix='Per-Job Limit:')
    add_bullet(doc, ' Only contacts with "Valid" email validation status receive outreach.',
               bold_prefix='Email Validation Gate:')
    add_bullet(doc, ' Mailbox auto-pauses if bounce rate > 5% or complaint rate > 0.3%.',
               bold_prefix='Auto-Pause:')
    add_bullet(doc, ' Contacts can unsubscribe; system enforces suppression list.',
               bold_prefix='Suppression List:')

    add_note_box(doc,
        'With the current single-mailbox default configuration, the system sends a maximum of '
        '30 cold emails per day. To scale, add more warmed-up mailboxes. '
        'With 10 mailboxes at 40 emails/day each, you reach 400 emails/day.',
        label='Bottom Line')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # QUESTION 2: Load Balancing
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Question 2: Outreach Load Balancing Across Mailboxes', level=1)
    doc.add_heading('Does the system load-balance by allocating contacts equally across active mailboxes?', level=3)

    doc.add_heading('Answer', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Yes.')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        ' The system implements a Least-Loaded Round-Robin load balancing strategy. '
        'Each time the outreach pipeline needs to send an email, it selects the mailbox with the '
        'lowest emails_sent_today count from all eligible mailboxes. This ensures even distribution '
        'of outreach volume across all active mailboxes.'
    )

    doc.add_heading('2.1 How Mailbox Selection Works', level=3)
    p = doc.add_paragraph(
        'When the pipeline processes a contact for sending, it queries the database for the next '
        'available mailbox using these criteria (all must be true):'
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Criterion', 'Field', 'Requirement'], header=True)
    criteria = [
        ['Mailbox is active', 'is_active', 'Must be True'],
        ['Warmup complete', 'warmup_status', 'COLD_READY or ACTIVE'],
        ['Has remaining capacity', 'emails_sent_today', '< daily_send_limit'],
        ['SMTP connection working', 'connection_status', '"successful"'],
    ]
    for row_data in criteria:
        add_table_row(table, row_data)

    p = doc.add_paragraph(
        'From all eligible mailboxes, the system selects the one with the lowest emails_sent_today '
        '(ORDER BY emails_sent_today ASC, LIMIT 1). This is the "least-loaded" strategy.'
    )

    doc.add_heading('2.2 Practical Example', level=3)
    p = doc.add_paragraph('Given 3 active mailboxes, each with a daily limit of 30:')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Email #', 'Mailbox A (sent)', 'Mailbox B (sent)', 'Selected'], header=True)
    example_data = [
        ['1st email', '0', '0', 'A (tied, first returned)'],
        ['2nd email', '1', '0', 'B (lower count)'],
        ['3rd email', '1', '1', 'A (tied, first returned)'],
        ['4th email', '2', '1', 'B (lower count)'],
        ['...', '...', '...', 'Alternates evenly'],
        ['89th email', '30 (full)', '29', 'B (A is full)'],
        ['90th email', '30 (full)', '30 (full)', 'None — daily limit reached'],
    ]
    for row_data in example_data:
        add_table_row(table, row_data)

    doc.add_heading('2.3 Key Behaviors', level=3)
    add_bullet(doc, ' Each contact assignment is independent — the system does NOT maintain affinity '
               '(i.e., follow-ups may come from a different mailbox).',
               bold_prefix='No Mailbox Affinity:')
    add_bullet(doc, ' If ALL eligible mailboxes have reached their daily limit, the contact is skipped '
               'with the reason "No available sender mailbox."',
               bold_prefix='Graceful Degradation:')
    add_bullet(doc, ' The counter (emails_sent_today) is incremented immediately after each successful send, '
               'ensuring real-time accuracy for the next selection.',
               bold_prefix='Real-Time Counter Update:')
    add_bullet(doc, ' The warmup peer system uses random selection (not least-loaded), since warmup '
               'emails have different distribution goals.',
               bold_prefix='Warmup Uses Different Strategy:')

    add_note_box(doc,
        'The system ensures fair, even distribution of outreach across all active mailboxes. '
        'No single mailbox is overloaded while others sit idle. This protects sender reputation '
        'and maximizes daily throughput.',
        label='Bottom Line')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # QUESTION 3: Contact Enrichment Alternatives
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Question 3: Alternative Contact Discovery When Apollo Fails', level=1)
    doc.add_heading('If contact enrichment cannot find contacts through Apollo for a lead, what are the alternatives?', level=3)

    doc.add_heading('Answer', level=2)
    p = doc.add_paragraph(
        'The system is designed with a multi-layered fallback strategy for contact discovery. '
        'When Apollo fails or returns zero results, several automatic and manual alternatives exist:'
    )

    doc.add_heading('3.1 Automatic Fallback: Seamless.ai', level=3)
    p = doc.add_paragraph(
        'The contact enrichment pipeline supports multiple providers configured in a priority list. '
        'If Apollo fails (returns 0 results, API error, or network timeout), the system automatically '
        'tries the next provider — Seamless.ai.'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Provider', 'Status', 'Data Quality', 'Cost'], header=True)
    providers = [
        ['Apollo.io', 'Primary (implemented)', 'High (email + phone + title)', 'Paid (credits)'],
        ['Seamless.ai', 'Secondary (implemented)', 'High (email + phone + title)', 'Paid (subscription)'],
        ['Mock', 'Dev/testing only', 'Synthetic data', 'Free'],
    ]
    for row_data in providers:
        add_table_row(table, row_data)

    p = doc.add_paragraph(
        'Configuration: Set contact_providers to ["apollo", "seamless"] in Admin Settings. '
        'The pipeline iterates through the provider list in order — if Apollo returns 0 results '
        'or throws an error, Seamless is tried next.'
    )

    doc.add_heading('3.2 Company Contact Cache (Zero API Cost)', level=3)
    p = doc.add_paragraph(
        'The system maintains a smart contact cache. When contacts are discovered for a company '
        '(e.g., "Acme Corp"), those contacts are automatically linked to ALL other leads at the same '
        'company. This means:'
    )
    add_bullet(doc, 'If Apollo finds 3 contacts for "Acme Corp" on Lead A, and Lead B is also at '
               '"Acme Corp", Lead B gets those same 3 contacts automatically — zero API calls.')
    add_bullet(doc, 'This "auto-enrichment" runs during both lead sourcing and contact enrichment pipelines.')
    add_bullet(doc, 'The junction table (lead_contact_association) supports many-to-many relationships, '
               'so one contact can serve multiple leads at the same company.')

    doc.add_heading('3.3 Manual Contact Creation via API', level=3)
    p = doc.add_paragraph(
        'When automated discovery fails entirely, administrators can manually add contacts:'
    )
    add_bullet(doc, ' POST /api/v1/contacts — Create a single contact with all fields '
               '(first_name, last_name, email, title, phone, priority_level).',
               bold_prefix='API Endpoint:')
    add_bullet(doc, ' Specify lead_ids to link the contact to one or more leads simultaneously.',
               bold_prefix='Lead Linking:')
    add_bullet(doc, ' Manually created contacts get validation_status="pending" and must go through '
               'the email validation pipeline before outreach.',
               bold_prefix='Validation:')

    doc.add_heading('3.4 Failure Scenarios & System Behavior', level=3)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Scenario', 'System Behavior', 'Resolution'], header=True)
    scenarios = [
        ['Apollo returns 0 results', 'Tries Seamless automatically', 'Automatic fallback'],
        ['Apollo API error (network)', 'Logs error, tries Seamless', 'Automatic fallback'],
        ['Apollo credits exhausted', 'Pipeline STOPS immediately', 'Add credits or switch to Seamless-only'],
        ['Both Apollo + Seamless fail', 'Checks company contact cache', 'Auto-links existing contacts'],
        ['All sources fail + no cache', 'Lead marked "skipped"', 'Manual contact creation required'],
    ]
    for row_data in scenarios:
        add_table_row(table, row_data)

    doc.add_heading('3.5 Contact Priority System (P1-P5)', level=3)
    p = doc.add_paragraph(
        'Both Apollo and Seamless classify discovered contacts using the same priority system:'
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Priority', 'Label', 'Title Keywords'], header=True)
    priorities = [
        ['P1', 'Job Poster / Hiring Manager', 'Hiring Manager, Talent Acquisition'],
        ['P2', 'HR/TA Recruiter', 'Recruiter, HR Coordinator, Talent'],
        ['P3', 'HR Manager', 'HR Manager, HRBP, HR Director, VP HR'],
        ['P4', 'Operations Leader', 'Operations, Plant Manager, Production'],
        ['P5', 'Functional Manager', 'All other titles (default)'],
    ]
    for row_data in priorities:
        add_table_row(table, row_data)

    add_note_box(doc,
        'The recommended setup is to configure both Apollo and Seamless as providers '
        '(["apollo", "seamless"]). This gives automatic fallback at zero extra effort. '
        'For leads where both fail, the company contact cache often fills the gap. '
        'Manual creation via API is the last resort.',
        label='Recommendation')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # QUESTION 4: Increasing Lead Volume
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Question 4: Strategies to Increase Lead Volume', level=1)
    doc.add_heading('If the lead pipeline returns very few new leads, what are the alternative ways to get more?', level=3)

    doc.add_heading('Answer', level=2)
    p = doc.add_paragraph(
        'Multiple strategies exist to increase lead volume, ranging from zero-code configuration '
        'changes to adding new data sources. Below are all options ranked by effort and impact:'
    )

    doc.add_heading('4.1 Immediate Wins (No Code Changes)', level=3)

    doc.add_heading('A. Broaden Target Job Titles', level=4)
    p = doc.add_paragraph(
        'The system currently searches for 37 job titles. Expanding this list to include '
        'subordinate/senior variants can dramatically increase results:'
    )
    add_bullet(doc, 'Add "Operations Coordinator", "Operations Specialist", "Operations Director", "VP Operations"')
    add_bullet(doc, 'Add "Warehouse Supervisor", "Logistics Coordinator", "Supply Chain Analyst"')
    add_bullet(doc, 'Add "Maintenance Supervisor", "Quality Supervisor", "Safety Coordinator"')
    add_bullet(doc, 'Configuration: Admin Settings > Job Sources > Target Job Titles')

    doc.add_heading('B. Expand Target Industries', level=4)
    p = doc.add_paragraph(
        'Currently targeting 22 non-IT industries. Adding adjacent industries increases the addressable market:'
    )
    add_bullet(doc, 'Add: Transportation, Agriculture, Chemical Manufacturing, Medical Devices, Pharmaceuticals')
    add_bullet(doc, 'Configuration: Admin Settings > Job Sources > Target Industries')

    doc.add_heading('C. Increase Pipeline Run Frequency', level=4)
    p = doc.add_paragraph(
        'The default is one pipeline run per day. Increasing to 2-3 runs captures newly posted jobs sooner. '
        'Deduplication prevents duplicate leads automatically.'
    )

    doc.add_heading('D. Manual CSV Import', level=4)
    p = doc.add_paragraph(
        'The system supports bulk lead import via CSV upload (POST /api/v1/leads/import/csv). '
        'This is ideal for researcher-curated leads from LinkedIn, past databases, or referrals.'
    )
    add_bullet(doc, 'CSV columns: Company Name, Job Title, State, Source, Salary Min/Max, Job Link')
    add_bullet(doc, 'Preview endpoint shows duplicates before import')
    add_bullet(doc, 'Supports hundreds of leads per upload')

    doc.add_heading('4.2 Medium-Term Improvements', level=3)

    doc.add_heading('A. Enable Apollo.io as a Lead Source', level=4)
    p = doc.add_paragraph(
        'Apollo.io provides organization-level hiring signals and can serve as both a lead source '
        'AND contact discovery provider. Benefits:'
    )
    add_bullet(doc, 'Returns 200-500+ leads per run with company data pre-populated')
    add_bullet(doc, 'Contact data often included (reducing enrichment API calls)')
    add_bullet(doc, 'Runs in parallel with JSearch (3 concurrent worker threads)')
    add_bullet(doc, 'Requirement: Paid Apollo.io API key')

    doc.add_heading('B. Relax Deduplication Window', level=4)
    p = doc.add_paragraph(
        'The current 30-day lookback window deduplicates against existing leads. '
        'If your database already has many leads, consider archiving older leads '
        'to allow re-discovery of roles that have been re-posted.'
    )

    doc.add_heading('C. Broaden Geographic Scope', level=4)
    p = doc.add_paragraph(
        'If currently filtering by specific states, expanding to all 50 US states '
        'or adding Canadian provinces can increase the lead pool significantly.'
    )

    doc.add_heading('4.3 Current Lead Sources Comparison', level=3)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Source', 'Status', 'Typical Volume', 'Cost', 'Recommendation'], header=True)
    sources = [
        ['JSearch (RapidAPI)', 'Active', '100-500+/run', 'Free tier: 500 req/mo', 'Primary source'],
        ['Apollo.io', 'Available', '200-500+/run', 'Paid (credits)', 'Enable as secondary'],
        ['Indeed', 'Available', '50-200/run', 'Publisher Program', 'Low value (JSearch covers it)'],
        ['CSV Import', 'Active', 'Unlimited', 'Free (manual effort)', 'Supplement as needed'],
        ['Mock', 'Dev only', '80-150/run', 'Free', 'Never use in production'],
    ]
    for row_data in sources:
        add_table_row(table, row_data)

    add_note_box(doc,
        'The highest-impact, lowest-effort actions are: (1) expand target job titles to 60+, '
        '(2) add 5-10 adjacent industries, and (3) enable Apollo.io as a second source. '
        'These three changes alone can 3-5x your lead volume.',
        label='Recommendation')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # QUESTION 5: Sending 100,000 Cold Emails Daily
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Question 5: Best Practice for Sending 100,000 Cold Emails Daily', level=1)
    doc.add_heading('What is the recommended approach to send 100,000 cold emails daily through this system?', level=3)

    doc.add_heading('Answer', level=2)
    p = doc.add_paragraph(
        'Reaching 100,000 cold emails per day is an enterprise-scale operation that requires '
        'significant infrastructure, careful domain management, and a phased ramp-up strategy. '
        'Below is a detailed blueprint.'
    )

    doc.add_heading('5.1 Infrastructure Requirements', level=3)

    p = doc.add_paragraph()
    run = p.add_run('Core Formula: ')
    run.bold = True
    run = p.add_run('100,000 emails/day ÷ 50 emails/mailbox/day = 2,000 active mailboxes')

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Component', 'Requirement', 'Rationale'], header=True)
    infra = [
        ['Sender Mailboxes', '2,000-2,500 (with 20% buffer)', '50 emails/day/mailbox is the safe max for reputation'],
        ['Domains', '200-250 unique domains', '10 mailboxes per domain to avoid domain-level throttling'],
        ['SMTP Providers', '3-4 providers (Office365, Google, Zoho, private SMTP)', 'Diversify to avoid single-provider blocks'],
        ['Email Validation Credits', '100,000+/month', 'Every contact must be validated before outreach'],
        ['Contact Discovery Credits', '25,000-50,000/month', 'At 4 contacts/lead, need 25K leads for 100K emails'],
        ['VPS/Servers', '4-8 VPS instances', 'Distribute load, avoid IP-level blocks'],
        ['IP Addresses', '20-50 dedicated IPs', 'Spread sending across IPs to maintain reputation'],
    ]
    for row_data in infra:
        add_table_row(table, row_data)

    doc.add_heading('5.2 Domain & Mailbox Strategy', level=3)

    add_bullet(doc, ' Never use your primary business domain for cold outreach. '
               'Register 200+ secondary domains (e.g., acme-solutions.com, acme-team.com).',
               bold_prefix='Separate Domains:')
    add_bullet(doc, ' Set up 10 mailboxes per domain (john@, sarah@, mike@, etc.).',
               bold_prefix='10 Mailboxes per Domain:')
    add_bullet(doc, ' Every domain must have valid SPF, DKIM, and DMARC records. '
               'The system\'s DNS checker validates this automatically.',
               bold_prefix='DNS Authentication:')
    add_bullet(doc, ' Use the system\'s built-in warmup engine. At Standard profile (30 days), '
               'each batch of new mailboxes needs 30 days before they can send outreach.',
               bold_prefix='Warmup Every Mailbox:')

    doc.add_heading('5.3 Phased Ramp-Up Schedule', level=3)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Phase', 'Timeline', 'Mailboxes', 'Daily Volume', 'Focus'], header=True)
    phases = [
        ['Phase 1', 'Weeks 1-4', '50 (warming)', '500-1,000', 'Domain setup, warmup, testing'],
        ['Phase 2', 'Weeks 5-8', '200 (warming) + 50 active', '2,500-5,000', 'Scale mailboxes, validate content'],
        ['Phase 3', 'Weeks 9-12', '500 active + 500 warming', '10,000-25,000', 'Multi-VPS, monitor deliverability'],
        ['Phase 4', 'Weeks 13-16', '1,000 active + 500 warming', '25,000-50,000', 'Optimize bounce/complaint rates'],
        ['Phase 5', 'Weeks 17-20', '2,000 active + 500 warming', '75,000-100,000', 'Full capacity, ongoing monitoring'],
    ]
    for row_data in phases:
        add_table_row(table, row_data)

    doc.add_heading('5.4 Required System Modifications', level=3)

    p = doc.add_paragraph('The current single-VPS architecture must be enhanced:')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Modification', 'Current State', 'Required Change'], header=True)
    mods = [
        ['Worker Architecture', 'Single process (APScheduler)', 'Celery + Redis with dedicated workers per VPS'],
        ['Database', 'Single MySQL on same VPS', 'Dedicated MySQL RDS or managed DB (separate from app)'],
        ['Queue System', 'Redis installed but unused', 'Redis as Celery broker + result backend'],
        ['Horizontal Scaling', 'Single VPS', '4-8 VPS instances with shared DB + Redis'],
        ['Monitoring', 'No APM/metrics', 'Sentry + Prometheus + Grafana for real-time alerts'],
        ['SMTP Connection Pool', 'New SMTP connection per email', 'Connection pooling (reuse per mailbox)'],
        ['Batch Processing', '100 contacts per pipeline run', 'Increase to 1,000+ with parallel workers'],
        ['IP Rotation', 'Single server IP', 'Multiple IPs via SMTP relay services (SendGrid, Mailgun)'],
    ]
    for row_data in mods:
        add_table_row(table, row_data)

    doc.add_heading('5.5 Deliverability Management', level=3)
    add_bullet(doc, ' Keep bounce rate under 2% (system auto-pauses at 5%).',
               bold_prefix='Bounce Rate:')
    add_bullet(doc, ' Keep complaint rate under 0.1% (system auto-pauses at 0.3%).',
               bold_prefix='Complaint Rate:')
    add_bullet(doc, ' Target 15%+ open rate; below 5% indicates deliverability issues.',
               bold_prefix='Open Rate Monitoring:')
    add_bullet(doc, ' At 100K/day scale, expect 5-10% mailbox rotation monthly '
               'due to reputation degradation.',
               bold_prefix='Mailbox Rotation:')
    add_bullet(doc, ' Consider warm IP pools from services like Mailgun, SendGrid, or Amazon SES '
               'for supplementary volume.',
               bold_prefix='Supplementary Infrastructure:')

    doc.add_heading('5.6 Cost Estimate (Monthly)', level=3)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Item', 'Quantity', 'Estimated Monthly Cost'], header=True)
    costs = [
        ['Domains (200+)', '200 x $12/yr', '$200/month (amortized)'],
        ['Mailbox Hosting (Google/O365)', '2,000 mailboxes', '$6,000-12,000/month'],
        ['VPS Instances (4-8)', '8 x $40/mo', '$320/month'],
        ['Managed Database', '1 x MySQL RDS', '$100-200/month'],
        ['Email Validation (NeverBounce)', '100K+ validations', '$300-500/month'],
        ['Apollo.io (Contact Discovery)', 'Professional plan', '$400-800/month'],
        ['Monitoring (Sentry + Grafana)', 'Team plans', '$100-200/month'],
        ['', 'TOTAL', '$7,500-14,000/month'],
    ]
    for row_data in costs:
        add_table_row(table, row_data)

    add_note_box(doc,
        'Sending 100,000 cold emails/day is feasible but requires 5-month ramp-up, '
        '2,000+ warmed mailboxes, multi-VPS architecture, and $7,500-14,000/month budget. '
        'The current system architecture supports this with the modifications listed above. '
        'Start with Phase 1 (50 mailboxes) to validate content and deliverability before scaling.',
        label='Bottom Line')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # QUESTION 6: Reusability for Other LOBs
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Question 6: Reusing the System for Other Lines of Business', level=1)
    doc.add_heading(
        'Can you reuse this system for other nature of leads and cold-emailing — '
        'for example IT services, AI/ML services, RCM, QA/Automation?', level=3
    )

    doc.add_heading('Answer', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Yes, the system is approximately 90% reusable across different lines of business.')
    run.bold = True
    p.add_run(
        ' The core architecture — lead sourcing, contact discovery, email validation, outreach pipeline, '
        'warmup engine, and mailbox management — is industry-agnostic. However, 10% of the system '
        'contains hardcoded staffing/recruiting references that must be customized.'
    )

    doc.add_heading('6.1 What Is Fully Configurable (No Code Changes)', level=3)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Component', 'Configuration Method', 'Change Required'], header=True)
    config_items = [
        ['Target Industries', 'Admin Settings UI', 'Replace with LOB-specific industries'],
        ['Target Job Titles', 'Admin Settings UI', 'Replace with LOB-specific decision-makers'],
        ['Excluded Keywords', 'Admin Settings UI', 'Adjust exclusion filters for new LOB'],
        ['Email Templates', 'Admin Settings UI', 'Create new templates with LOB-specific messaging'],
        ['AI Provider', 'Admin Settings UI', 'Select provider (Groq/OpenAI/Anthropic/Gemini)'],
        ['Contact Provider', 'Admin Settings UI', 'Apollo, Seamless, or both'],
        ['Business Rules', 'Admin Settings UI', 'Send limits, cooldowns, contacts per lead'],
        ['Company Size Preferences', 'Admin Settings UI', 'Adjust employee count targets'],
        ['Warmup Settings', 'Admin Settings UI', '30+ configurable parameters'],
        ['Email Signature', 'Per-Mailbox Settings', 'Custom HTML signature per mailbox'],
    ]
    for row_data in config_items:
        add_table_row(table, row_data)

    doc.add_heading('6.2 What Requires Minor Code Changes (~4-6 hours)', level=3)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Component', 'Current State', 'Change Required'], header=True)
    code_changes = [
        ['AI System Prompt', 'Hardcoded "staffing/recruitment" context in Groq adapter',
         'Make configurable via database setting (ai_system_prompt)'],
        ['Fallback Email Body', 'Hardcoded "staffing solutions" message in outreach pipeline',
         'Make configurable via database setting or require active template'],
        ['Fallback Subject Line', 'Hardcoded "Exciting Opportunity" format',
         'Make configurable via database setting'],
        ['AI Fallback Response', 'Hardcoded "staffing services" in error fallback',
         'Make configurable or use generic text'],
    ]
    for row_data in code_changes:
        add_table_row(table, row_data)

    doc.add_heading('6.3 LOB-Specific Configuration Examples', level=3)

    doc.add_heading('IT / Application Development Services', level=4)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Setting', 'Value'], header=True)
    it_config = [
        ['Target Industries', 'Software, Technology, Cloud Computing, SaaS, FinTech, E-commerce, Digital Media'],
        ['Target Job Titles', 'CTO, VP Engineering, Engineering Director, Tech Lead, DevOps Manager, IT Director, CIO'],
        ['Exclude Keywords', 'Remove IT exclusions; add staffing/recruiting exclusions instead'],
        ['Email Template Focus', 'Software development capabilities, tech stack expertise, case studies'],
        ['Contact Priority P1', '"CTO", "VP Engineering", "Chief Technology Officer"'],
    ]
    for row_data in it_config:
        add_table_row(table, row_data)

    doc.add_heading('AI / Machine Learning Services', level=4)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Setting', 'Value'], header=True)
    ai_config = [
        ['Target Industries', 'Technology, Healthcare, Finance, Research, Automotive, Defense, Retail Analytics'],
        ['Target Job Titles', 'Chief Data Officer, AI Research Director, ML Engineering Manager, Head of Analytics, VP Data Science'],
        ['Exclude Keywords', 'Staffing, recruiting, temp agency'],
        ['Email Template Focus', 'AI/ML capabilities, model development, data pipeline expertise'],
        ['Contact Priority P1', '"Chief Data Officer", "VP AI", "Head of Machine Learning"'],
    ]
    for row_data in ai_config:
        add_table_row(table, row_data)

    doc.add_heading('Revenue Cycle Management (RCM)', level=4)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Setting', 'Value'], header=True)
    rcm_config = [
        ['Target Industries', 'Healthcare, Hospitals, Medical Practices, Health Insurance, Ambulatory Care'],
        ['Target Job Titles', 'Revenue Cycle Director, Billing Manager, Claims Manager, CFO, VP Finance, Compliance Officer'],
        ['Exclude Keywords', 'IT staffing, software development, tech recruiting'],
        ['Email Template Focus', 'Revenue optimization, claims processing, compliance, denial management'],
        ['Contact Priority P1', '"Revenue Cycle Director", "VP Revenue Cycle", "CFO"'],
    ]
    for row_data in rcm_config:
        add_table_row(table, row_data)

    doc.add_heading('QA / Automation Testing Services', level=4)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Setting', 'Value'], header=True)
    qa_config = [
        ['Target Industries', 'Software, SaaS, FinTech, Healthcare IT, Banking Technology, E-commerce'],
        ['Target Job Titles', 'QA Manager, Test Director, QA Lead, VP Engineering, Release Manager, DevOps Director'],
        ['Exclude Keywords', 'Staffing, recruiting, temp agency'],
        ['Email Template Focus', 'Test automation, CI/CD integration, quality metrics, shift-left testing'],
        ['Contact Priority P1', '"QA Director", "VP Quality", "Head of Testing"'],
    ]
    for row_data in qa_config:
        add_table_row(table, row_data)

    doc.add_heading('6.4 Deployment Options for Multiple LOBs', level=3)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Option', 'Effort', 'Cost', 'Best For'], header=True)
    deploy_opts = [
        ['Reconfigure Single Instance', '1-2 hours', 'Zero', 'Switching between LOBs sequentially'],
        ['Separate Deployments per LOB', '2-4 hours each', '$40-80/mo per instance', '2-3 concurrent LOBs'],
        ['Parameterize AI Prompts', '4-6 hours (one-time)', 'Zero', 'Single instance, quick LOB switching'],
        ['Full Multi-Campaign System', '2-3 weeks dev', 'Dev time only', 'Enterprise: 5+ concurrent LOBs'],
    ]
    for row_data in deploy_opts:
        add_table_row(table, row_data)

    doc.add_heading('6.5 Architecture Strengths for Reuse', level=3)

    add_bullet(doc, ' All external integrations (job sources, contact discovery, email validation, '
               'AI content) use an abstract adapter pattern. New providers can be added without '
               'changing pipeline logic.',
               bold_prefix='Adapter Pattern:')
    add_bullet(doc, ' Over 90% of business logic is configurable via the Admin Settings UI or database, '
               'not hardcoded in source code.',
               bold_prefix='Configuration-Driven:')
    add_bullet(doc, ' Email content is generated by configurable AI models (Groq, OpenAI, Anthropic, Gemini) '
               'using customizable templates with placeholder support.',
               bold_prefix='AI-Powered Content:')
    add_bullet(doc, ' Lead sourcing, contact enrichment, email validation, and outreach are independent '
               'pipeline stages — each can be customized without affecting others.',
               bold_prefix='Pipeline Independence:')
    add_bullet(doc, ' The warmup engine, mailbox management, DNS checking, and reputation monitoring '
               'are completely industry-agnostic.',
               bold_prefix='Warmup Engine is LOB-Agnostic:')

    add_note_box(doc,
        'The system is highly reusable. For a single LOB switch, reconfigure settings in 1-2 hours '
        '(zero code changes if an active email template is set). For concurrent multi-LOB operations, '
        'deploy separate instances ($40-80/month each). For enterprise-scale multi-LOB in a single system, '
        'a 2-3 week development effort adds full campaign management.',
        label='Bottom Line')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX A
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Appendix A: Configuration Quick Reference', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Setting', 'Default', 'Location', 'Impact'], header=True)
    config_ref = [
        ['DAILY_SEND_LIMIT', '30/mailbox', 'Settings UI', 'Max emails per mailbox per day'],
        ['COOLDOWN_DAYS', '10 days', 'Settings UI', 'Min gap between emails to same contact'],
        ['MAX_CONTACTS_PER_COMPANY_PER_JOB', '4', 'Settings UI', 'Max contacts emailed per job lead'],
        ['MIN_SALARY_THRESHOLD', '$30,000', 'Settings UI', 'Filters low-salary leads'],
        ['TARGET_INDUSTRIES', '22 non-IT', 'Settings UI', 'Industry filter for lead sourcing'],
        ['TARGET_JOB_TITLES', '37 titles', 'Settings UI', 'Job title search keywords'],
        ['CONTACT_PROVIDER', 'mock', 'Settings UI', 'Apollo, Seamless, or mock'],
        ['EMAIL_VALIDATION_PROVIDER', 'mock', 'Settings UI', '7 providers available'],
        ['AI Provider', 'Groq', 'Settings UI', 'Content generation engine'],
        ['Warmup Profile', 'Standard (30d)', 'Per mailbox', 'Conservative/Standard/Aggressive'],
    ]
    for row_data in config_ref:
        add_table_row(table, row_data)

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX B
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('Appendix B: Warmup Profile Comparison', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ['Phase', 'Days', 'Conservative', 'Standard', 'Aggressive', 'Notes'], header=True)
    warmup_detail = [
        ['Phase 1', '1-10 / 1-7 / 1-5', '1-3/day', '2-5/day', '3-8/day', 'Initial reputation building'],
        ['Phase 2', '11-20 / 8-14 / 6-10', '3-8/day', '5-15/day', '8-20/day', 'Gradual increase'],
        ['Phase 3', '21-30 / 15-21 / 11-15', '8-15/day', '15-25/day', '20-35/day', 'Approaching full capacity'],
        ['Phase 4', '31-45 / 22-30 / 16-20', '15-25/day', '25-35/day', '35-50/day', 'Full warmup volume'],
        ['Total Duration', '', '45 days', '30 days', '20 days', ''],
        ['Auto-Pause', '', 'Bounce >5%', 'Bounce >5%', 'Bounce >5%', 'Same thresholds for all'],
        ['To ACTIVE', '', '7+ days at 80+ health', '7+ days at 80+ health', '7+ days at 80+ health', 'Post-warmup gate'],
    ]
    for row_data in warmup_detail:
        add_table_row(table, row_data)

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Document generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('NeuraLeads AI Agent — SME Technical Assessment & Strategic Guidance')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


if __name__ == '__main__':
    doc = create_document()
    output_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    output_path = os.path.join(output_dir, 'Exzelon_RA_Agent_SME_Assessment.docx')
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Size: {os.path.getsize(output_path):,} bytes")
