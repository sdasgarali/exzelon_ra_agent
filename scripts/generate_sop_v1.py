"""Generate Exzelon RA Agent SOP v1 document (.docx) with embedded screenshots.

Run from project root:
    python scripts/generate_sop_v1.py

Requires: python-docx, screenshots in scripts/sop_screenshots/
Produces: Exzelon_RA_Agent_SOP_v1.docx in project root (C:\Ali\Rizwan-Taiyab\Exzelon\AI-Agent-RA\)
"""
import os
import sys

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
SCREENSHOT_DIR = os.path.join(SCRIPT_DIR, "sop_screenshots")
# Output goes to the parent directory (C:\Ali\Rizwan-Taiyab\Exzelon\AI-Agent-RA\)
OUTPUT_DIR = os.path.dirname(PROJECT_ROOT)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "Exzelon_RA_Agent_SOP_v1.docx")

# ---------------------------------------------------------------------------
# Color palette
# ---------------------------------------------------------------------------
COLOR_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)
COLOR_DARK = RGBColor(0x1F, 0x2A, 0x37)
COLOR_GREEN = RGBColor(0x05, 0x7A, 0x55)
COLOR_RED = RGBColor(0xC8, 0x16, 0x16)
COLOR_AMBER = RGBColor(0x92, 0x40, 0x0E)
COLOR_PURPLE = RGBColor(0x6B, 0x21, 0xA8)
COLOR_GRAY = RGBColor(0x6B, 0x72, 0x80)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_HEADER_BG = "1E40AF"
COLOR_GREEN_LIGHT = "D1FAE5"
COLOR_BLUE_LIGHT = "DBEAFE"
COLOR_AMBER_LIGHT = "FEF3C7"
COLOR_RED_LIGHT = "FEE2E2"
COLOR_PURPLE_LIGHT = "F3E8FF"
COLOR_GRAY_LIGHT = "F3F4F6"
COLOR_INDIGO_BG = "EEF2FF"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, alignment=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None, compact=False,
                     header_bg=None, row_colors=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_bg = header_bg or COLOR_HEADER_BG
    font_size = Pt(8) if compact else Pt(9)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, hdr_bg)
        set_cell_text(cell, header, bold=True, color=COLOR_WHITE, size=font_size)

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_colors and row_idx < len(row_colors) and row_colors[row_idx]:
                set_cell_shading(cell, row_colors[row_idx])
            elif not row_colors and row_idx % 2 == 1:
                set_cell_shading(cell, COLOR_GRAY_LIGHT)

            text = str(value)
            if text.startswith("[GREEN]"):
                set_cell_text(cell, text[7:], color=COLOR_GREEN, size=font_size)
            elif text.startswith("[RED]"):
                set_cell_text(cell, text[5:], color=COLOR_RED, size=font_size)
            elif text.startswith("[AMBER]"):
                set_cell_text(cell, text[7:], color=COLOR_AMBER, size=font_size)
            elif text.startswith("[PURPLE]"):
                set_cell_text(cell, text[8:], color=COLOR_PURPLE, size=font_size)
            elif text.startswith("[BOLD]"):
                set_cell_text(cell, text[6:], bold=True, size=font_size)
            else:
                set_cell_text(cell, text, size=font_size)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width
    doc.add_paragraph()
    return table


def add_callout_box(doc, text, style="info"):
    colors = {
        "info": (COLOR_BLUE_LIGHT, COLOR_PRIMARY),
        "success": (COLOR_GREEN_LIGHT, COLOR_GREEN),
        "warning": (COLOR_AMBER_LIGHT, COLOR_AMBER),
        "danger": (COLOR_RED_LIGHT, COLOR_RED),
        "purple": (COLOR_PURPLE_LIGHT, COLOR_PURPLE),
    }
    bg, fg = colors.get(style, colors["info"])
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.color.rgb = fg
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    doc.add_paragraph()


def add_screenshot(doc, filename, caption=None, width_inches=6.2):
    filepath = os.path.join(SCREENSHOT_DIR, filename)
    if not os.path.exists(filepath):
        p = doc.add_paragraph()
        run = p.add_run(f"[Screenshot not available: {filename}]")
        run.font.color.rgb = COLOR_RED
        run.italic = True
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(filepath, width=Inches(width_inches))

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_GRAY
        run.italic = True
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_toc(doc):
    """Insert a Table of Contents field that Word will render on open."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)

    run2 = p.add_run()
    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instr)

    run3 = p.add_run()
    fld_char_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fld_char_separate)

    run4 = p.add_run("Right-click and select 'Update Field' to populate the Table of Contents")
    run4.font.color.rgb = COLOR_GRAY
    run4.italic = True

    run5 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fld_char_end)


# ---------------------------------------------------------------------------
# Document Sections
# ---------------------------------------------------------------------------
def add_title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_heading("Exzelon RA Agent", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = COLOR_PRIMARY
        run.font.size = Pt(34)

    subtitle = doc.add_heading("Standard Operating Procedure (SOP)", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = COLOR_DARK
        run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Complete End-to-End Operational Guide")
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_GRAY
    run.italic = True

    doc.add_paragraph()

    # Metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("System URL", "https://ra.partnerwithus.tech"),
        ("Version", "1.1 (Updated March 2026)"),
        ("Classification", "CONFIDENTIAL"),
        ("Platform", "FastAPI + Next.js 14 + MySQL"),
        ("Pages", "20 Dashboard Pages | 32 API Modules | 36 Database Tables"),
    ]
    for i, (label, value) in enumerate(meta_data):
        set_cell_text(meta_table.rows[i].cells[0], label, bold=True, size=Pt(10))
        set_cell_text(meta_table.rows[i].cells[1], value, size=Pt(10))
        if i % 2 == 0:
            set_cell_shading(meta_table.rows[i].cells[0], COLOR_GRAY_LIGHT)
            set_cell_shading(meta_table.rows[i].cells[1], COLOR_GRAY_LIGHT)

    doc.add_paragraph()

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("CONFIDENTIAL - INTERNAL USE ONLY")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_RED
    run.bold = True

    doc.add_page_break()


def add_toc_section(doc):
    add_heading(doc, "Table of Contents")
    add_toc(doc)
    doc.add_page_break()


def section_01_overview(doc):
    add_heading(doc, "1. System Overview & Login")

    doc.add_paragraph(
        "The Exzelon Research Analyst (RA) Agent is an enterprise-grade automated cold-email "
        "outreach platform that replaces manual spreadsheets, mail merge, and multi-tool workflows "
        "with a single web-based system. It automates the entire lifecycle: job research, contact "
        "discovery, email validation, multi-step campaigns, reply management, and CRM deal tracking."
    )

    add_heading(doc, "1.1 What This System Replaces", level=2)

    headers = ["Old Manual Process", "New Automated System"]
    rows = [
        ["Search Indeed/LinkedIn/Glassdoor manually", "Lead Sourcing Pipeline -- auto-fetches from 10 job APIs"],
        ["Copy job data into Job_requirements.xls", "Leads page -- all leads stored in database with search/filter"],
        ["Look up contacts on Apollo/Seamless manually", "Contact Enrichment Pipeline -- 8 adapters auto-discover decision-makers"],
        ["Export to contact_details.csv", "Contacts page -- 10-column sortable table with server-side sorting"],
        ["Validate emails via NeverBounce one by one", "Email Validation Pipeline -- 7 providers, bulk validation"],
        ["MS Word Mail Merge -> send 30 emails", "Campaign Engine -- multi-step sequences with A/B testing"],
        ["Track everything in Excel/Google Sheets", "Dashboard -- real-time KPIs, charts, analytics, CRM deals"],
        ["Manage client_info table manually", "Clients page -- auto-categorized (Regular/Occasional)"],
        ["Reply tracking in Outlook manually", "Unified Inbox -- AI sentiment analysis, auto-reply suggestions"],
        ["No CRM pipeline", "CRM Deal Pipeline -- Kanban board with auto-advancement"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "1.2 Logging In", level=2)
    add_screenshot(doc, "01_login.png", "Figure 1.1: Login Page")

    add_numbered_list(doc, [
        "Open https://ra.partnerwithus.tech in your browser.",
        "Enter your Email and Password.",
        "Click Sign In.",
        "After successful login, you are redirected to the Dashboard.",
    ])

    add_heading(doc, "1.3 User Roles", level=2)

    headers = ["Role", "Access Level", "Can Manage"]
    rows = [
        ["Super Admin", "Full access to everything", "Users, Roles, Settings, Backups, Restore, all operations"],
        ["Admin", "Manage operations and settings", "Settings, Users, Backups (no Roles page, no Restore)"],
        ["Operator", "Day-to-day pipeline operations", "Run pipelines, campaigns, outreach, mailboxes"],
        ["Viewer", "Read-only access", "View data only -- cannot run pipelines or modify anything"],
    ]
    add_styled_table(doc, headers, rows)

    add_callout_box(doc,
        "PERMISSION SYSTEM: The platform uses a 3-layer permission model: "
        "(1) Default hardcoded permissions, (2) Role-based matrix (configurable by Super Admin), "
        "(3) User-specific overrides. Super Admin always bypasses all checks.",
        style="info"
    )


def section_02_dashboard(doc):
    add_heading(doc, "2. Dashboard -- Command Center")

    doc.add_paragraph(
        "The Dashboard is your single-screen overview of the entire operation. It loads "
        "automatically after login and provides KPIs, performance metrics, quick action buttons, "
        "activity trends, pipeline funnel visualization, and system health indicators."
    )

    add_screenshot(doc, "02_dashboard.png", "Figure 2.1: Dashboard Overview")

    add_heading(doc, "2.1 KPI Cards (Top Row)", level=2)
    headers = ["Card", "What It Shows"]
    rows = [
        ["Companies Identified", "Total unique companies found through lead sourcing"],
        ["Total Contacts", "Total decision-maker contacts discovered"],
        ["Valid Emails", "Contacts with validated email addresses"],
        ["Emails Sent", "Total outreach emails sent"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "2.2 Performance Metrics", level=2)
    add_bullet_list(doc, [
        "Bounce Rate -- Percentage of emails that bounced. Target: < 2%. Progress bar turns red if above 2%.",
        "Reply Rate -- Percentage of contacts who replied to outreach.",
        "Total Leads -- All-time lead count from job sourcing.",
    ])

    add_heading(doc, "2.3 Quick Actions (Operator/Admin/Super Admin Only)", level=2)

    headers = ["Button", "Color", "Action"]
    rows = [
        ["Run Lead Sourcing", "Indigo", "Opens confirmation dialog -> scrapes 10 job source APIs"],
        ["Enrich Contacts", "Purple", "Opens Lead Selector modal -> choose leads to enrich via 8 contact adapters"],
        ["Validate Emails", "Cyan", "Opens Contact Selector modal -> bulk validate via 7 providers"],
        ["Export Mailmerge", "Orange", "Opens Lead Selector modal -> generate CSV or programmatic send"],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph(
        "Selector Modals allow you to search, filter by status, paginate, select individual rows "
        "or all, and then run for selected items or all items (with a warning confirmation)."
    )

    add_heading(doc, "2.4 Charts & Visualizations", level=2)
    add_bullet_list(doc, [
        "30-Day Activity Trends -- Area chart showing daily leads sourced (indigo) and emails sent (orange).",
        "Pipeline Funnel -- Visual flow: Leads -> Contacts -> Valid Emails -> Emails Sent -> Replied, "
        "with conversion percentages between each stage.",
        "Outreach Outcomes -- Sent, Replied, Bounced, Skipped stat boxes with reply rate and bounce rate bars.",
        "Lead Status Distribution -- Horizontal bar chart showing leads per status.",
        "Contact Validation Breakdown -- Donut chart: Valid (green), Invalid (red), Catch-All (yellow), Unknown (gray).",
    ])

    add_heading(doc, "2.5 System Health", level=2)
    headers = ["Card", "Shows"]
    rows = [
        ["Ready to Send", "Active mailboxes with available daily quota / total"],
        ["Warming Up", "Mailboxes currently in warmup phase"],
        ["Active Templates", "Email templates marked as active / total"],
        ["Lead Sources", "Number of distinct job sources configured"],
    ]
    add_styled_table(doc, headers, rows)


def section_03_lead_sourcing(doc):
    add_heading(doc, "3. Phase 1 -- Lead Sourcing (Job Research & Data Capture)")

    add_callout_box(doc,
        "REPLACES: Manually searching Indeed, LinkedIn, Glassdoor and copying job data into "
        "Job_requirements.xls. The system now auto-fetches from 10 official job APIs in parallel.",
        style="info"
    )

    add_heading(doc, "3.1 Job Source Adapters (10 Sources)", level=2)

    doc.add_paragraph(
        "All 10 adapters are independently configurable from Settings > Job Sources. Each can "
        "be enabled/disabled, tested, and configured with API keys -- no code changes needed."
    )

    headers = ["Source", "Type", "Free Tier", "Paid Plan", "Status"]
    rows = [
        ["JSearch", "RapidAPI aggregator (Indeed, LinkedIn, Glassdoor)", "500 req/mo", "$50/mo", "[GREEN]Active"],
        ["Apollo", "Apollo.io org search", "10K credits/mo", "$500/mo", "[GREEN]Active"],
        ["TheirStack", "Career page scraper API", "100 req/mo", "$49/mo", "[GREEN]Active"],
        ["SerpAPI", "Google Jobs proxy", "100 searches/mo", "$50/mo", "[GREEN]Active"],
        ["Adzuna", "40+ country job aggregator", "250 req/mo", "$99/mo", "[GREEN]Active"],
        ["SearchAPI.io", "Google Jobs proxy (alt)", "100 searches/mo", "$40/mo", "[GREEN]Implemented"],
        ["USAJobs", "US federal government jobs", "Unlimited (free)", "FREE", "[GREEN]Implemented"],
        ["Jooble", "71-country aggregator", "Generous free tier", "FREE", "[GREEN]Implemented"],
        ["JobDataFeeds", "Bulk data feed (Techmap)", "N/A", "$400/mo", "[GREEN]Implemented"],
        ["Coresignal", "399M+ jobs + recruiter contacts", "N/A", "$800-1,500/mo", "[PURPLE]Premium"],
    ]
    add_styled_table(doc, headers, rows, compact=True)

    add_callout_box(doc,
        "CORESIGNAL DIFFERENTIATOR: The only source that bundles recruiter/hiring manager "
        "contact details directly with job postings, eliminating the separate contact "
        "enrichment step for those leads.",
        style="purple"
    )

    add_heading(doc, "3.2 Running Lead Sourcing", level=2)

    doc.add_paragraph("Method A: From Dashboard")
    add_numbered_list(doc, [
        "Go to Dashboard.",
        "Click the indigo Run Lead Sourcing button.",
        "Confirm in the dialog (shows which sources are enabled).",
        "You are redirected to the Pipelines page to monitor progress.",
    ])

    doc.add_paragraph("Method B: From Pipelines Page")
    add_numbered_list(doc, [
        "Go to Pipelines in the sidebar.",
        "Click Run Pipeline on the Lead Sourcing card (indigo).",
        "Confirm and watch progress in the Pipeline Run History table.",
    ])

    add_heading(doc, "3.3 What Gets Captured Per Lead", level=2)
    headers = ["Field", "Description"]
    rows = [
        ["Client/Company Name", "The hiring company"],
        ["Job Title", "The position being hired for"],
        ["State", "US state where the job is located"],
        ["City", "City (for enhanced deduplication)"],
        ["Posting Date", "When the job was posted"],
        ["Job Link", "URL to the original job posting"],
        ["Salary Min/Max", "Salary range if available"],
        ["Source", "Which job API it came from"],
        ["External Job ID", "Unique ID from the source (Layer 1 dedup)"],
        ["Employer LinkedIn URL", "Company LinkedIn URL (Layer 2 dedup)"],
        ["Employer Website", "Company website URL"],
        ["Job Publisher", "Sub-source (LinkedIn, Indeed, Glassdoor, etc.)"],
        ["Lead Status", "Tracking status (new -> enriched -> validated -> sent)"],
    ]
    add_styled_table(doc, headers, rows, compact=True)

    add_heading(doc, "3.4 3-Layer Deduplication", level=2)
    doc.add_paragraph(
        "The pipeline uses a sophisticated 3-layer deduplication strategy to prevent storing "
        "duplicate leads, plus a 30-day rolling database window:"
    )
    add_bullet_list(doc, [
        "Layer 1 (Exact Match): Dedup by external_job_id across all adapters in a single run.",
        "Layer 2 (Company Match): Match on employer_linkedin_url to catch same-company postings.",
        "Layer 3 (Normalized Key): Normalized company_name + job_title + state + city key match.",
        "Database Window: 30-day rolling window prevents re-processing of recently seen leads.",
    ])

    add_heading(doc, "3.5 Automatic Business Rules Applied", level=2)
    add_bullet_list(doc, [
        "Filters to 22 Non-IT industries only.",
        "Excludes IT keywords (software developer, data scientist, devops, etc.).",
        "Excludes US staffing/recruitment agencies.",
        "Filters by minimum salary threshold ($30K+).",
        "Targets US-based jobs only.",
        "Auto-creates/updates client_info records for new companies.",
        "6 ThreadPoolExecutor workers run adapters in parallel.",
        "Automated 3x/day at 6am, 12pm, 6pm UTC via APScheduler.",
    ])

    add_heading(doc, "3.6 Viewing Sourced Leads", level=2)
    add_screenshot(doc, "03_leads.png", "Figure 3.1: Leads Page")
    doc.add_paragraph(
        "Navigate to Leads in the sidebar. The page shows a searchable, filterable table "
        "with columns for ID, Company, Job Title, State, Salary, Source, Status, Contacts count, "
        "and Posted date. You can search by company/title/state and filter by Status, State, "
        "Source, Salary Range, and Posting Date."
    )


def section_04_contact_enrichment(doc):
    add_heading(doc, "4. Phase 2 -- Contact Enrichment (Identify Decision-Makers)")

    add_callout_box(doc,
        "REPLACES: Manually looking up contacts on Apollo.io and Seamless.ai, finding "
        "HR/recruiter contacts, exporting to contact_details.csv.",
        style="info"
    )

    add_heading(doc, "4.1 Contact Discovery Adapters (8 Sources)", level=2)
    headers = ["Adapter", "Pricing", "Data"]
    rows = [
        ["Apollo", "~$500/mo (10K credits free)", "Decision-maker search by company + title"],
        ["Seamless.AI", "~$99/mo", "B2B contact database"],
        ["Hunter.io", "Free: 25 -> $49/mo", "Email finder by domain"],
        ["Snov.io", "Free: 50 credits -> $39/mo", "Email discovery + enrichment"],
        ["RocketReach", "Free: 5 -> $99/mo", "B2B contact database"],
        ["People Data Labs", "Free: 100 -> $0.01/match", "1.5B+ person profiles"],
        ["Proxycurl", "Free: 10 -> $0.01/call", "LinkedIn profile data"],
        ["Mock", "Free (dev/testing)", "Generates test data"],
    ]
    add_styled_table(doc, headers, rows, compact=True)

    add_heading(doc, "4.2 Contact Priority Levels", level=2)
    headers = ["Priority", "Label", "Description"]
    rows = [
        ["[RED]P1", "Job Poster", "The person who posted the job (highest priority)"],
        ["[AMBER]P2", "HR / Recruiter", "HR / Talent Acquisition / Recruiter"],
        ["P3", "HR Manager", "HR Manager / HRBP / Director"],
        ["P4", "Ops Leader", "Operations / Plant / Production / Business Leaders"],
        ["P5", "Functional Manager", "Role-specific functional managers (lowest priority)"],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph(
        "Rules: Same-state contacts are preferred unless HR is centralized. "
        "Maximum 4 contacts per company per job."
    )

    add_heading(doc, "4.3 Running Contact Enrichment", level=2)
    add_numbered_list(doc, [
        "Go to Dashboard and click the purple Enrich Contacts button.",
        "The Lead Selector modal opens -- search/filter/select specific leads.",
        "Click Run for Selected (N) or Run for All Leads.",
        "If 'Run for All', a warning confirmation appears.",
        "Monitor progress on the Pipelines page.",
    ])

    add_heading(doc, "4.4 Viewing Contacts", level=2)
    add_screenshot(doc, "04_contacts.png", "Figure 4.1: Contacts Page")
    doc.add_paragraph(
        "The Contacts page shows a 10-column sortable table with server-side sorting. "
        "All columns are sortable: Name, Company, Email, Phone, Priority, Validation, "
        "Lead ID, Source, Status, Unsubscribed. Filter by Priority, Validation Status, "
        "Source (all 8 adapters), and Outreach Status."
    )


def section_05_validation(doc):
    add_heading(doc, "5. Phase 3 -- Email Validation")

    add_callout_box(doc,
        "REPLACES: Manually running bulk validation through NeverBounce, ZeroBounce, or Hunter; "
        "removing invalid emails from spreadsheets.",
        style="info"
    )

    add_heading(doc, "5.1 Validation Providers (7 Options)", level=2)
    headers = ["Provider", "Pricing"]
    rows = [
        ["NeverBounce", "~$99/mo"],
        ["ZeroBounce", "~$50/mo"],
        ["Hunter.io", "Free: 100 -> $49/mo"],
        ["Clearout", "~$99/mo"],
        ["Emailable", "~$99/mo"],
        ["MailboxValidator", "~$99/mo"],
        ["Reacher", "Self-hosted (free)"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "5.2 Running Email Validation", level=2)
    add_screenshot(doc, "06_validation.png", "Figure 5.1: Email Validation Page")

    add_numbered_list(doc, [
        "Go to Dashboard and click the cyan Validate Emails button.",
        "The Contact Selector modal opens -- filter by 'Pending' to validate only new contacts.",
        "Select contacts or click Run for All Contacts.",
        "Monitor progress on the Pipelines page.",
    ])

    add_heading(doc, "5.3 Validation Statuses", level=2)
    headers = ["Status", "Meaning", "Action"]
    rows = [
        ["[GREEN]Valid", "Email is deliverable", "Include in outreach"],
        ["[RED]Invalid", "Email does not exist or bounces", "Excluded from outreach automatically"],
        ["[AMBER]Catch-All", "Domain accepts all emails", "Configurable via Settings"],
        ["Unknown", "Provider couldn't determine status", "Review manually"],
        ["Pending", "Not yet validated", "Run validation"],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph(
        "The system enforces a < 2% bounce rate target by only sending to Valid contacts "
        "(and optionally Catch-All based on your Settings configuration)."
    )


def section_06_campaigns(doc):
    add_heading(doc, "6. Campaigns -- Multi-Step Email Sequences")

    add_callout_box(doc,
        "NEW FEATURE: The Campaign Engine replaces simple one-shot emails with multi-step "
        "sequences, A/B testing, conditional branching, timezone-aware send windows, and "
        "spintax text variation.",
        style="success"
    )

    add_screenshot(doc, "07_campaigns.png", "Figure 6.1: Campaigns Page")

    add_heading(doc, "6.1 Campaign Features", level=2)
    add_bullet_list(doc, [
        "Multi-step sequences with email, wait, and condition (if/then branching) steps.",
        "A/B testing with weighted variant assignment and chi-squared auto-optimization.",
        "Spintax text variation: {option1|option2} with nested pattern support and deterministic seeding.",
        "Timezone-aware send windows using contact's US state.",
        "Round-robin mailbox selection based on health score.",
        "Reply, bounce, and unsubscribe tracking per campaign contact.",
        "Auto-enrollment of new contacts into active campaigns.",
        "Campaign queue processed every 2 minutes by APScheduler.",
    ])

    add_heading(doc, "6.2 Creating a Campaign", level=2)
    add_numbered_list(doc, [
        "Go to Campaigns in the sidebar.",
        "Click Create Campaign.",
        "Set campaign name, description, send window (start/end hours), timezone.",
        "Add sequence steps: Email (with subject + body + A/B variants), Wait (delay), or Condition.",
        "Assign mailboxes to the campaign (round-robin selection).",
        "Save and activate when ready.",
        "Enroll contacts manually or enable auto-enrollment.",
    ])

    add_heading(doc, "6.3 A/B Testing", level=2)
    doc.add_paragraph(
        "Each email step can have multiple variants with weighted assignment (e.g., 50/50 or 70/30). "
        "The system uses chi-squared statistical testing to automatically determine winners and "
        "shift traffic to the best-performing variant. Metrics tracked per variant: open rate, "
        "click rate, reply rate."
    )

    add_heading(doc, "6.4 Spintax Engine", level=2)
    doc.add_paragraph(
        "Write email copy with text variations using curly braces and pipes: "
        "{Hi|Hello|Hey} {contact_first_name}, {I noticed|I saw|I came across} your "
        "{posting|opening|listing}. Supports nested patterns for unique combinations. "
        "Each contact gets a deterministic variant based on a seeded hash."
    )


def section_07_outreach(doc):
    add_heading(doc, "7. Outreach (Single-Send Email)")

    doc.add_paragraph(
        "The Outreach page handles single-send (non-campaign) email sending and mail merge exports. "
        "For multi-step sequences, use Campaigns instead."
    )

    add_screenshot(doc, "08_outreach.png", "Figure 7.1: Outreach Page")

    add_heading(doc, "7.1 Two Outreach Modes", level=2)
    headers = ["Mode", "Description", "When to Use"]
    rows = [
        ["Mailmerge Export", "Generates a CSV file ready for MS Word", "When you prefer Outlook/Word sending"],
        ["Programmatic Send", "System sends through configured mailboxes", "For fully automated sending"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "7.2 Merge Fields Available", level=2)
    headers = ["Field", "Replaced With"]
    rows = [
        ["{{contact_first_name}}", "Contact's first name"],
        ["{{sender_first_name}}", "Sender mailbox display name"],
        ["{{job_title}}", "Job title from the lead"],
        ["{{job_location}}", "State/location from the lead"],
        ["{{company_name}}", "Company name"],
        ["{{signature}}", "Sender's email signature (HTML)"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "7.3 Business Rules Enforced", level=2)
    headers = ["Rule", "Value"]
    rows = [
        ["Daily send limit per mailbox", "30 emails (configurable)"],
        ["Cooldown between emails to same contact", "10 days (configurable)"],
        ["Max contacts per company per job", "4 (configurable)"],
        ["Only valid emails", "Enforced automatically"],
        ["Exclude staffing agencies", "Enforced automatically"],
    ]
    add_styled_table(doc, headers, rows)


def section_08_inbox(doc):
    add_heading(doc, "8. Unified Inbox (Unibox)")

    add_callout_box(doc,
        "NEW FEATURE: Centralized reply management with AI-powered sentiment analysis "
        "and reply suggestions. Replaces manual Outlook reply tracking.",
        style="success"
    )

    add_screenshot(doc, "09_inbox.png", "Figure 8.1: Unified Inbox")

    add_heading(doc, "8.1 Three-Panel Layout", level=2)
    add_bullet_list(doc, [
        "Left Panel: Thread list with contact name, subject, timestamp, and category badge.",
        "Center Panel: Full conversation thread with all sent/received messages.",
        "Right Panel: Contact details, lead info, and action buttons.",
    ])

    add_heading(doc, "8.2 AI Features", level=2)
    add_bullet_list(doc, [
        "Sentiment Analysis: Auto-classifies replies as interested, not_interested, OOO, question, "
        "referral, or do_not_contact using rule-based matching + LLM fallback.",
        "Reply Suggestions: AI generates contextual reply drafts based on conversation history.",
        "Auto-Forward to CRM: Interested replies automatically create CRM deals (configurable).",
    ])

    add_heading(doc, "8.3 Category Labels", level=2)
    headers = ["Category", "Color", "Meaning"]
    rows = [
        ["[GREEN]Interested", "Green", "Contact expressed interest -- priority follow-up"],
        ["[RED]Not Interested", "Red", "Contact declined -- respect and log"],
        ["[AMBER]OOO", "Amber", "Out of office -- retry later"],
        ["Question", "Blue", "Contact has questions -- respond promptly"],
        ["Referral", "Purple", "Contact referred someone else -- follow up on referral"],
        ["[RED]Do Not Contact", "Red", "Suppressed -- no further outreach"],
    ]
    add_styled_table(doc, headers, rows)


def section_09_deals(doc):
    add_heading(doc, "9. CRM Deal Pipeline")

    add_callout_box(doc,
        "NEW FEATURE: Kanban-style deal tracking with auto-creation from interested replies, "
        "auto-stage advancement, activity timeline, tasks, and forecasting.",
        style="success"
    )

    add_screenshot(doc, "10_deals.png", "Figure 9.1: CRM Deal Pipeline")

    add_heading(doc, "9.1 Pipeline Stages (7 Default)", level=2)
    headers = ["Stage", "Description"]
    rows = [
        ["New Lead", "Initial contact established"],
        ["Contacted", "Outreach sent, awaiting response"],
        ["Qualified", "Contact confirmed interest"],
        ["Proposal", "Proposal or candidate profiles shared"],
        ["Negotiation", "Terms being discussed"],
        ["[GREEN]Won", "Deal successfully closed"],
        ["[RED]Lost", "Deal lost or contact declined"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "9.2 Deal Automation", level=2)
    add_bullet_list(doc, [
        "Auto-Create: When an inbox reply is categorized as 'interested', a deal is automatically "
        "created (configurable via Settings > Automation).",
        "Auto-Advance: Deal stages advance based on email events (reply = Qualified, etc.).",
        "Stale Detection: Deals inactive for configurable period are flagged.",
        "Activity Timeline: All deal interactions logged with timestamps.",
        "Deal Tasks: Sub-tasks with assignee, due date, priority, and status.",
    ])

    add_heading(doc, "9.3 Deal Statistics", level=2)
    add_bullet_list(doc, [
        "Win Rate: Percentage of deals that reach the 'Won' stage.",
        "Average Deal Size: Mean deal value across all won deals.",
        "Pipeline Value: Total value of open deals across all stages.",
        "Forecast: Probability-weighted pipeline value.",
    ])


def section_10_analytics(doc):
    add_heading(doc, "10. Analytics")

    add_screenshot(doc, "11_analytics.png", "Figure 10.1: Analytics Page")

    doc.add_paragraph(
        "The Analytics page provides team performance, campaign comparison, and revenue "
        "metrics for data-driven decision-making."
    )

    add_heading(doc, "10.1 Analytics Features", level=2)
    add_bullet_list(doc, [
        "Team Leaderboard: Per-user send counts, reply rates, and conversion metrics.",
        "Campaign Comparison: Side-by-side performance of different campaigns.",
        "Revenue Metrics: Cost per lead, ROI by campaign, monthly revenue trends.",
        "Cost Tracking: Per-adapter API costs, total spend, cost-per-acquisition.",
    ])


def section_11_icp_wizard(doc):
    add_heading(doc, "11. ICP Wizard (AI-Powered)")

    add_screenshot(doc, "12_icp_wizard.png", "Figure 11.1: ICP Wizard Page")

    doc.add_paragraph(
        "The ICP (Ideal Customer Profile) Wizard uses AI to generate targeting profiles "
        "based on your historical data, industry knowledge, and business objectives."
    )

    add_heading(doc, "11.1 How It Works", level=2)
    add_numbered_list(doc, [
        "Click Generate ICP Profile.",
        "AI analyzes your successful placements, reply rates, and industry performance.",
        "A profile is generated with recommended: industries, job titles, US states, "
        "company sizes, and scoring weights.",
        "Save the profile for use in lead sourcing configuration.",
        "Apply the ICP to automatically configure Settings > Job Sources.",
    ])

    add_callout_box(doc,
        "The ICP Wizard includes a rule-based fallback that works without any AI provider "
        "configured, ensuring the feature is always available.",
        style="info"
    )


def section_12_lead_management(doc):
    add_heading(doc, "12. Lead Management")

    add_heading(doc, "12.1 Lead Statuses", level=2)
    headers = ["Status", "Meaning"]
    rows = [
        ["New", "Just sourced from job boards"],
        ["Enriched", "Contacts have been found for this lead"],
        ["Validated", "Contact emails have been validated"],
        ["Open", "Active lead, ready for outreach"],
        ["Hunting", "Actively pursuing this lead"],
        ["Sent", "Outreach emails have been sent"],
        ["Skipped", "Skipped due to business rules"],
        ["Closed - Hired", "Successfully placed a candidate"],
        ["Closed - Not Hired", "Lead closed without placement"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "12.2 Bulk Operations", level=2)
    headers = ["Action", "Description"]
    rows = [
        ["Enrich Contacts", "Run contact enrichment for selected leads"],
        ["Run Outreach", "Send outreach for selected leads"],
        ["Update Status", "Change status of all selected leads"],
        ["Archive", "Soft-delete selected leads (can be restored)"],
        ["Export CSV", "Download filtered results as CSV"],
        ["Import CSV", "Upload a CSV file with lead data"],
    ]
    add_styled_table(doc, headers, rows)


def section_13_clients(doc):
    add_heading(doc, "13. Client / Company Management")

    add_screenshot(doc, "05_clients.png", "Figure 13.1: Clients Page")

    add_heading(doc, "13.1 Client Categories (Auto-Classified)", level=2)
    headers = ["Category", "Definition", "Original SOP Equivalent"]
    rows = [
        ["Regular", "Company has >3 unique posting dates in last 3 months", "Database A -- Active Clients"],
        ["Occasional", "Company has <=3 unique posting dates in last 3 months", "Database B -- One-Time Clients"],
        ["Prospect", "New company, not yet categorized", "--"],
        ["Dormant", "Company has been inactive", "--"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "13.2 Client Page Features", level=2)
    add_bullet_list(doc, [
        "Website and LinkedIn columns with clickable links.",
        "Bulk Enrich button for company data enrichment (Clearbit, OpenCorporates).",
        "Industry, Service Count, Start/End Date tracking.",
        "LLM fallback enrichment for companies without API data.",
    ])


def section_14_mailboxes(doc):
    add_heading(doc, "14. Mailbox Management")

    add_screenshot(doc, "14_mailboxes.png", "Figure 14.1: Mailboxes Page")

    add_heading(doc, "14.1 Mailbox Configuration", level=2)
    headers = ["Field", "Description"]
    rows = [
        ["Email", "Sender email address"],
        ["Display Name", "Name shown in 'From' field"],
        ["Provider", "Microsoft 365, Gmail, or Custom SMTP"],
        ["OAuth2", "Microsoft 365 OAuth2 flow supported (no password needed)"],
        ["Warmup Status", "warming_up, cold_ready, active, paused, inactive, blacklisted, recovering"],
        ["Daily Send Limit", "Max emails per day (default: 30)"],
        ["Health Score", "0-100 based on bounce rate, reply rate, warmup progress"],
        ["Signature", "HTML email signature with live preview"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "14.2 Health-Aware Mailbox Selection", level=2)
    doc.add_paragraph(
        "When sending emails, the system selects the optimal mailbox using a weighted formula: "
        "Score = Health * 0.4 + Quota Remaining * 0.3 + Warmup Age * 0.15 + Deliverability * 0.15. "
        "This ensures emails are sent from the healthiest, most reputable mailboxes."
    )

    add_heading(doc, "14.3 OAuth2 Support (Microsoft 365)", level=2)
    doc.add_paragraph(
        "Microsoft 365 mailboxes can be connected via OAuth2 flow -- no passwords needed. "
        "Configure MS365_OAUTH_CLIENT_ID, MS365_OAUTH_CLIENT_SECRET, and MS365_OAUTH_TENANT_ID "
        "in Settings > Outreach, then use the 'Connect with Microsoft' button when adding a mailbox."
    )


def section_15_templates(doc):
    add_heading(doc, "15. Email Templates")

    add_screenshot(doc, "13_templates.png", "Figure 15.1: Templates Page")

    doc.add_paragraph(
        "Create reusable email templates with merge fields. Only one template can be "
        "active at a time for outreach (campaigns use their own step-level templates). "
        "Templates support HTML body with live preview and plain-text fallback."
    )

    add_heading(doc, "15.1 Sample Template (Matching Original SOP)", level=2)
    doc.add_paragraph(
        'Subject: U.S. Staffing Services -- Introduction\n\n'
        'Body:\n'
        'Hi {{contact_first_name}},\n\n'
        'This is {{sender_first_name}} from Exzelon. I noticed you\'re hiring for a '
        '{{job_title}} role in {{job_location}}.\n\n'
        'We specialize in connecting employers with pre-vetted, local professionals across '
        'the U.S. -- and we operate on a risk-free recruitment model. You can review resumes '
        'and interview candidates at no cost and only pay if you decide to hire.\n\n'
        'Would you like me to share a few local resumes for your review?\n\n'
        '{{signature}}'
    )


def section_16_warmup(doc):
    add_heading(doc, "16. Warmup Engine")

    add_screenshot(doc, "15_warmup.png", "Figure 16.1: Warmup Engine")

    doc.add_paragraph(
        "The Warmup Engine manages email domain reputation by gradually increasing send "
        "volume and monitoring deliverability. It runs continuously via APScheduler."
    )

    add_heading(doc, "16.1 Warmup Tabs", level=2)
    headers = ["Tab", "Purpose"]
    rows = [
        ["Overview", "Mailbox warmup status table -- day, phase, health score, bounce/reply rates"],
        ["Analytics", "Charts: email volume trends, health score trends, bounce vs reply rates"],
        ["Emails", "Peer-to-peer warmup email history (AI-generated)"],
        ["DNS & Blacklist", "SPF/DKIM/DMARC checking; blacklist monitoring (Spamhaus, Barracuda, etc.)"],
        ["Profiles", "3 profiles: Conservative (45d), Standard (30d), Aggressive (20d)"],
        ["Alerts", "System alerts for warmup issues"],
        ["Settings", "Warmup phases, thresholds, health score weights, auto-pause rules"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "16.2 Warmup Features", level=2)
    add_bullet_list(doc, [
        "Peer-to-peer warmup emails between your own mailboxes.",
        "AI-generated warmup content (via Groq/OpenAI/Anthropic/Gemini).",
        "Auto-reply to warmup emails with AI-generated responses.",
        "IMAP read emulation (marks peer emails as read to simulate engagement).",
        "Spam score checking (100+ trigger words, 0-100 score).",
        "Inbox placement tracking.",
        "Smart send scheduling for optimal delivery times.",
        "Auto-recovery for degraded mailboxes.",
    ])


def section_17_pipelines(doc):
    add_heading(doc, "17. Pipeline Orchestration & Monitoring")

    add_screenshot(doc, "16_pipelines.png", "Figure 17.1: Pipelines Page")

    add_heading(doc, "17.1 The 4-Stage Pipeline", level=2)
    doc.add_paragraph(
        "The page shows a visual workflow: Lead Sourcing (indigo) -> Contact Enrichment (purple) "
        "-> Email Validation (cyan) -> Outreach (orange). Each stage shows its cumulative count."
    )

    add_heading(doc, "17.2 Pipeline Run History", level=2)
    headers = ["Column", "Description"]
    rows = [
        ["Run ID", "Unique identifier"],
        ["Pipeline", "Which pipeline (color-coded badge)"],
        ["Status", "Running (animated), Completed (green), Failed (red), Cancelled (gray)"],
        ["Started", "Timestamp"],
        ["Duration", "How long it took"],
        ["Adapter", "Which provider was used"],
        ["Records", "Total records processed"],
        ["Success/Failed", "Green count / Red count"],
        ["Triggered By", "User email or 'system' (for scheduled runs)"],
        ["Actions", "Cancel button for running jobs"],
    ]
    add_styled_table(doc, headers, rows, compact=True)

    add_heading(doc, "17.3 Cancellation", level=2)
    doc.add_paragraph(
        "Running pipelines can be cancelled via the Cancel button in the history table. "
        "The system uses cooperative cancellation -- the job stops after completing its "
        "current batch. Progress percentage is tracked and updated incrementally."
    )


def section_18_automation(doc):
    add_heading(doc, "18. Automation Control Center")

    add_callout_box(doc,
        "NEW FEATURE: Centralized control for all 20+ scheduled automation jobs with "
        "master toggle, per-job enable/disable, and activity event log.",
        style="success"
    )

    add_screenshot(doc, "17_automation.png", "Figure 18.1: Automation Control Center")

    add_heading(doc, "18.1 Automation Controls", level=2)
    add_bullet_list(doc, [
        "Master Toggle: Global on/off switch for all automation.",
        "Lead Sourcing: Scheduled 3x/day (6am/12pm/6pm UTC).",
        "Campaign Processing: Queue processed every 2 minutes.",
        "Auto-Enrollment: New contacts auto-enrolled into active campaigns.",
        "Deal Auto-Create: Interested replies create CRM deals.",
        "Deal Auto-Advance: Deal stages advance on email events.",
        "Warmup Cycles: Daily warmup email sending.",
        "DNS Monitoring: Periodic SPF/DKIM/DMARC checks.",
        "Blacklist Monitoring: Periodic blacklist scans.",
        "Stale Deal Detection: Flag inactive deals.",
    ])

    add_heading(doc, "18.2 Activity Event Log", level=2)
    doc.add_paragraph(
        "Every automation run is logged to the AutomationEvent table with: event type, "
        "timestamp, status (success/failure), details, and trigger source. View the log "
        "on the Automation page with filtering by event type and date range."
    )


def section_19_settings(doc):
    add_heading(doc, "19. Settings & Configuration")

    doc.add_paragraph(
        "Settings is organized into 7+ tabs, each controlling a different aspect of the system. "
        "Access is role-based with per-tab permissions."
    )

    add_screenshot(doc, "21_settings.png", "Figure 19.1: Settings Page")

    add_heading(doc, "19.1 Job Sources Tab", level=2)
    add_screenshot(doc, "21a_settings_job_sources.png", "Figure 19.2: Settings > Job Sources", 5.8)

    doc.add_paragraph(
        "Configure all 10 job source adapters. Each adapter has: enable/disable checkbox, "
        "API key input field, Test Connection button, and pricing info. Changes take effect "
        "on the next pipeline run -- no restart needed."
    )

    add_heading(doc, "19.2 AI / LLM Tab", level=2)
    add_screenshot(doc, "21b_settings_ai_llm.png", "Figure 19.3: Settings > AI / LLM", 5.8)

    headers = ["Setting", "Description"]
    rows = [
        ["AI Provider", "Groq (free, fast), OpenAI, Anthropic, Gemini"],
        ["API Key", "Provider-specific API key"],
        ["Model", "Specific model to use"],
    ]
    add_styled_table(doc, headers, rows)
    doc.add_paragraph("Used for: warmup content, auto-reply, email generation, ICP wizard, "
                      "sentiment analysis, reply suggestions, sequence generation, copilot.")

    add_heading(doc, "19.3 Contacts Tab", level=2)
    add_screenshot(doc, "21c_settings_contacts.png", "Figure 19.4: Settings > Contacts", 5.8)
    doc.add_paragraph(
        "Configure which contact discovery provider to use (Apollo, Seamless, or Mock) and "
        "enter the corresponding API key. Additional providers (Hunter, Snov.io, RocketReach, "
        "PDL, Proxycurl) can also be configured here."
    )

    add_heading(doc, "19.4 Validation Tab", level=2)
    add_screenshot(doc, "21d_settings_validation.png", "Figure 19.5: Settings > Validation", 5.8)
    doc.add_paragraph(
        "Select one of 7 email validation providers and enter the API key."
    )

    add_heading(doc, "19.5 Outreach Tab", level=2)
    add_screenshot(doc, "21e_settings_outreach.png", "Figure 19.6: Settings > Outreach", 5.8)
    doc.add_paragraph(
        "Configure email send mode (Programmatic or Mailmerge), SMTP settings, and "
        "Microsoft 365 OAuth2 credentials."
    )

    add_heading(doc, "19.6 Business Rules Tab", level=2)
    add_screenshot(doc, "21f_settings_business_rules.png", "Figure 19.7: Settings > Business Rules", 5.8)

    headers = ["Setting", "Default", "Description"]
    rows = [
        ["Daily Send Limit", "30", "Max emails per mailbox per day"],
        ["Cooldown Days", "10", "Min days between emails to same contact"],
        ["Max Contacts per Company/Job", "4", "Max contacts to email per job posting"],
        ["Min Salary Threshold", "$30,000", "Exclude jobs below this salary"],
        ["Catch-All Policy", "Configurable", "Include or exclude catch-all emails"],
        ["Backup Retention Days", "3", "Auto-delete backups older than this (1-90)"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "19.7 Automation Tab", level=2)
    add_screenshot(doc, "21g_settings_automation.png", "Figure 19.8: Settings > Automation", 5.8)
    doc.add_paragraph(
        "Toggle 20+ individual automation jobs on/off. Includes lead sourcing schedule, "
        "campaign processing, auto-enrollment, deal automation, warmup cycles, DNS monitoring, "
        "blacklist monitoring, and more."
    )


def section_20_users(doc):
    add_heading(doc, "20. User Management")

    add_screenshot(doc, "18_users.png", "Figure 20.1: User Management Page")

    add_heading(doc, "20.1 Managing Users (Admin+)", level=2)
    add_bullet_list(doc, [
        "Create User: Click Add User -> enter email, password, full name, select role.",
        "Edit User: Click edit icon -> change name, role, active status.",
        "Delete User: Click delete icon -> confirm deletion.",
        "Protections: Cannot delete your own account. Only Super Admin can modify other SAs.",
    ])


def section_21_roles(doc):
    add_heading(doc, "21. Roles & Permissions (Super Admin Only)")

    add_screenshot(doc, "19_roles.png", "Figure 21.1: Roles & Permissions Page")

    add_heading(doc, "21.1 Permission Levels", level=2)
    headers = ["Level", "Label", "Color", "Meaning"]
    rows = [
        ["[GREEN]full", "Full Access", "Green", "CRUD + configure"],
        ["read_write", "Read & Write", "Blue", "CRU (no delete/configure)"],
        ["[AMBER]read", "Read Only", "Yellow", "View only"],
        ["[RED]no_access", "No Access", "Red", "Hidden, API returns 403"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "21.2 Features", level=2)
    add_bullet_list(doc, [
        "Per-Module Configuration: Each module (Dashboard, Leads, Contacts, etc.) can have "
        "different access per role.",
        "Independent Tab Permissions: Settings module allows per-tab access control.",
        "User-Specific Overrides: Grant or restrict individual users beyond their role defaults.",
        "Reset to Defaults: One-click reset button.",
    ])


def section_22_backups(doc):
    add_heading(doc, "22. Data Backups")

    add_screenshot(doc, "20_backups.png", "Figure 22.1: Data Backups Page")

    add_heading(doc, "22.1 Automated Backups", level=2)
    add_bullet_list(doc, [
        "Daily automatic backup at 2:00 AM UTC.",
        "Retention: Configurable (default 3 days, range 1-90).",
        "Format: Compressed SQL dump (.sql.gz).",
    ])

    add_heading(doc, "22.2 Operations", level=2)
    headers = ["Action", "Role Required", "Description"]
    rows = [
        ["Create Backup", "Admin+", "Create a new backup on demand"],
        ["Download", "Admin+", "Download backup file locally"],
        ["Delete", "Super Admin", "Permanently delete a backup"],
        ["Restore", "Super Admin", "Replace database with backup (type RESTORE to confirm)"],
    ]
    add_styled_table(doc, headers, rows)

    add_callout_box(doc,
        "WARNING: Restore replaces ALL current data with the backup contents. A pre-restore "
        "snapshot is created automatically as a safety net.",
        style="danger"
    )


def section_23_integrations(doc):
    add_heading(doc, "23. Integrations & Webhooks")

    add_heading(doc, "23.1 Webhook System", level=2)
    doc.add_paragraph(
        "Subscribe to system events with HMAC-SHA256 signed webhook delivery. "
        "Configure target URLs, event filters, and view delivery history with retry tracking."
    )

    headers = ["Event", "Trigger"]
    rows = [
        ["email.sent", "Outreach email delivered"],
        ["email.opened", "Email opened (tracking pixel)"],
        ["email.clicked", "Link clicked (tracking redirect)"],
        ["email.replied", "Reply received"],
        ["email.bounced", "Email bounced"],
        ["contact.unsubscribed", "Contact opted out"],
        ["campaign.completed", "Campaign finished all steps"],
        ["lead.created", "New lead sourced"],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph(
        "Retry policy: Exponential backoff with 3 attempts (1 min, 5 min, 15 min). "
        "Each delivery is signed with X-Webhook-Signature header for verification."
    )

    add_heading(doc, "23.2 CRM Integrations", level=2)
    add_bullet_list(doc, [
        "HubSpot: Push deals, pull contacts, custom property mapping.",
        "Salesforce: OAuth2 flow, Salesforce object sync.",
        "Bidirectional sync with CRM Sync Engine.",
        "Auto-forward interested inbox replies to CRM.",
    ])

    add_heading(doc, "23.3 Zapier/Make Integration", level=2)
    doc.add_paragraph(
        "REST hook subscribe/unsubscribe pattern with sample payloads for Zapier and Make "
        "(Integromat) compatibility. API key authentication via X-API-Key header."
    )

    add_heading(doc, "23.4 Notification Channels", level=2)
    add_bullet_list(doc, [
        "Slack: Webhook-based notifications for campaign events and errors.",
        "Microsoft Teams: Webhook adapter for Teams channels.",
        "Twilio: SMS and calling adapter (stubs for future implementation).",
    ])


def section_24_business_rules(doc):
    add_heading(doc, "24. Business Rules Reference")

    doc.add_paragraph(
        "This section consolidates all business rules enforced automatically by the system."
    )

    add_heading(doc, "24.1 Lead Sourcing Rules", level=2)
    headers = ["Rule", "Value", "Enforcement"]
    rows = [
        ["Industries", "22 Non-IT industries only", "IT industries excluded"],
        ["IT Exclusion", "14+ IT keywords", "IT job titles filtered out"],
        ["Staffing Exclusion", "7+ staffing keywords", "Staffing agencies excluded"],
        ["Geography", "United States only", "Non-US jobs filtered"],
        ["Salary Minimum", "$30,000+", "Jobs below threshold excluded"],
        ["Deduplication", "3-layer strategy + 30-day DB window", "Duplicates prevented"],
        ["Job Titles", "55 target titles across all categories", "Off-target titles excluded"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "24.2 Outreach Rules", level=2)
    headers = ["Rule", "Value", "Original SOP"]
    rows = [
        ["Daily Send Limit", "30 per mailbox", "'Sending Email 30 TO draft'"],
        ["Cooldown Period", "10 days", "'Same contact within 10 days'"],
        ["Max per Company/Job", "4 contacts", "'More than 4 contacts per company per job'"],
        ["Valid Emails Only", "Enforced", "'Do NOT email invalid or unverified emails'"],
        ["No Staffing Agencies", "Enforced", "'Staffing agencies' excluded"],
        ["Track Last Outreach", "Automatic", "'Always track last outreach date'"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "24.3 Automated Enforcement", level=2)
    headers = ["Original SOP Rule", "System Enforcement"]
    rows = [
        ["Do NOT email invalid or unverified emails", "Only 'Valid' contacts receive outreach"],
        ["Do NOT email staffing agencies", "Staffing keywords excluded in sourcing"],
        ["Do NOT email same contact within 10 days", "Cooldown period enforced automatically"],
        ["Do NOT email more than 4 contacts per company per job", "Max contacts limit enforced"],
        ["Avoid Non-US jobs", "US-only filter in sourcing"],
        ["Avoid roles below $30K salary", "Min salary threshold filter"],
        ["Always track last outreach date", "Outreach events tracked with timestamps"],
    ]
    add_styled_table(doc, headers, rows)


def section_25_glossary(doc):
    add_heading(doc, "25. Glossary")

    headers = ["Term", "Definition"]
    rows = [
        ["Lead", "A job posting found via job APIs -- represents a hiring opportunity at a company"],
        ["Contact", "A decision-maker at a company (HR manager, recruiter, ops leader, etc.)"],
        ["Client", "A company/organization that has been identified as hiring"],
        ["Pipeline", "An automated multi-step process (sourcing -> enrichment -> validation -> outreach)"],
        ["Campaign", "A multi-step email sequence with A/B testing and conditional branching"],
        ["Mailbox", "A sender email account used to send outreach emails"],
        ["Warmup", "The process of gradually increasing email volume to build domain reputation"],
        ["Bounce", "An email that could not be delivered"],
        ["Catch-All", "A domain that accepts all emails regardless of address existence"],
        ["Dry Run", "A test execution that simulates without actually sending emails"],
        ["Merge Fields", "Template placeholders like {{contact_first_name}} replaced with real data"],
        ["Spintax", "Text variation syntax {option1|option2} for unique email content per recipient"],
        ["P1-P5", "Contact priority levels from highest (Job Poster) to lowest (Functional Manager)"],
        ["ICP", "Ideal Customer Profile -- AI-generated targeting profile"],
        ["Deal", "A CRM sales opportunity tracked through pipeline stages"],
        ["Webhook", "HTTP callback triggered by system events (email.sent, lead.created, etc.)"],
        ["KPI", "Key Performance Indicator -- metrics like bounce rate, reply rate, send count"],
        ["RBAC", "Role-Based Access Control -- different users have different permission levels"],
        ["Adapter", "A modular integration component (one per external API/service)"],
        ["Dedup", "Deduplication -- preventing duplicate records from being stored"],
        ["APScheduler", "Background job scheduler that runs automation tasks on a schedule"],
    ]
    add_styled_table(doc, headers, rows, compact=True)


def appendix_a_workflow(doc):
    add_heading(doc, "Appendix A: Complete Workflow -- Step by Step")

    add_heading(doc, "A.1 Day 1: Initial Setup (One-Time)", level=2)
    add_numbered_list(doc, [
        "Login as Super Admin.",
        "Go to Settings > Job Sources -- configure API keys, target states, industries, job titles.",
        "Go to Settings > Contacts -- configure contact provider API keys.",
        "Go to Settings > Validation -- configure email validation provider API key.",
        "Go to Settings > Outreach -- configure email send mode and SMTP/M365 settings.",
        "Go to Settings > Business Rules -- verify limits (30/day, 10-day cooldown, 4 max contacts).",
        "Go to Settings > Automation -- enable desired automation jobs.",
        "Go to Mailboxes -- add all sender mailboxes, test connections.",
        "Go to Templates -- create outreach email template with merge fields, activate it.",
        "Go to Campaigns -- create multi-step email sequences for automated follow-up.",
        "Go to Users -- create accounts for operators/admins.",
        "Go to Roles & Permissions -- configure access levels for each role.",
    ])

    add_heading(doc, "A.2 Daily Operation Workflow", level=2)

    doc.add_paragraph("Step 1: Source Leads", style="List Bullet")
    doc.add_paragraph(
        "Dashboard -> 'Run Lead Sourcing' -> Confirm -> Monitor on Pipelines page. "
        "Result: New job postings captured from 10 APIs, companies auto-created. "
        "(Also runs automatically 3x/day if Automation is enabled.)"
    )

    doc.add_paragraph("Step 2: Enrich Contacts", style="List Bullet")
    doc.add_paragraph(
        "Dashboard -> 'Enrich Contacts' -> Select leads -> Run. "
        "Result: Decision-maker contacts discovered with P1-P5 priority levels."
    )

    doc.add_paragraph("Step 3: Validate Emails", style="List Bullet")
    doc.add_paragraph(
        "Dashboard -> 'Validate Emails' -> Select contacts (filter: Pending) -> Run. "
        "Result: Contacts marked as Valid/Invalid/Catch-All/Unknown."
    )

    doc.add_paragraph("Step 4: Send Outreach", style="List Bullet")
    doc.add_paragraph(
        "Option A: Dashboard -> 'Export Mailmerge' for single-send or CSV export. "
        "Option B: Campaigns page -> create/activate campaign for multi-step sequences."
    )

    doc.add_paragraph("Step 5: Monitor Results", style="List Bullet")
    doc.add_paragraph(
        "Dashboard -> Check KPIs, bounce rate, reply rate. "
        "Unified Inbox -> View replies with AI sentiment analysis. "
        "CRM Deals -> Track interested leads through pipeline stages."
    )

    add_heading(doc, "A.3 Weekly Maintenance", level=2)
    add_numbered_list(doc, [
        "Review Dashboard -- Check bounce rate (target < 2%), reply rate, funnel conversion.",
        "Review Warmup -- Check mailbox health scores, DNS records, blacklist status.",
        "Review Analytics -- Compare campaign performance, check cost metrics.",
        "Review Clients -- Check client categories (Regular/Occasional), inactive companies.",
        "Review Automation -- Check event log for any failures or anomalies.",
        "Create Backup -- Go to Backups -> Create Backup (in addition to daily auto-backups).",
    ])


def appendix_b_mapping(doc):
    add_heading(doc, "Appendix B: Original SOP Section Mapping")

    headers = ["Original SOP Section", "System Page(s)"]
    rows = [
        ["Section 1(A): Lead Research & Data Capture", "Dashboard -> Run Lead Sourcing, Leads page, Pipelines"],
        ["Section 1(B): Identify Email-Outreach-Contact", "Dashboard -> Enrich Contacts, Contacts page"],
        ["Section 1(B) Step 2: Identify Decision-Makers", "Contact Enrichment Pipeline (P1-P5 priority)"],
        ["Section 1(B) Step 3: Export Data (CSV)", "Contacts export, Mailmerge Export"],
        ["Section 1(B): Data Quality & Validation", "Dashboard -> Validate Emails, Validation page"],
        ["Section 1(B): Maintain Databases A & B", "Clients page (Regular = Active, Occasional = One-Time)"],
        ["Section 1(C): Email Sending (Mail Merge)", "Campaigns page, Outreach page"],
        ["Section 1(C): MS Word Mail Merge", "Mailmerge Export mode generates CSV for Word"],
        ["What MUST Be Avoided", "Business Rules enforced automatically (Section 24)"],
        ["Section 4: Dashboard Tabs", "Dashboard, Analytics, Leads, Contacts, Outreach pages"],
        ["-- (NEW)", "Campaigns -- multi-step sequences with A/B testing"],
        ["-- (NEW)", "Unified Inbox -- AI sentiment analysis & reply suggestions"],
        ["-- (NEW)", "CRM Deal Pipeline -- Kanban board with auto-advancement"],
        ["-- (NEW)", "Automation Control Center -- 20+ scheduled jobs"],
        ["-- (NEW)", "ICP Wizard -- AI-generated targeting profiles"],
        ["-- (NEW)", "Analytics -- team leaderboard, revenue metrics, cost tracking"],
    ]
    add_styled_table(doc, headers, rows, compact=True)


def appendix_c_architecture(doc):
    add_heading(doc, "Appendix C: Technical Architecture")

    add_heading(doc, "C.1 Technology Stack", level=2)
    headers = ["Layer", "Technology"]
    rows = [
        ["Backend", "FastAPI (Python 3.11) + SQLAlchemy 2.0 ORM + Pydantic"],
        ["Frontend", "Next.js 14 (TypeScript) + Tailwind CSS + Radix UI"],
        ["Database", "MySQL 8.0 (production), SQLite (testing)"],
        ["Auth", "JWT tokens (7-day) + API keys (SHA-256) + RBAC (4 roles)"],
        ["Scheduler", "APScheduler (20+ background jobs)"],
        ["State", "Zustand (auth) + TanStack React Query (server data)"],
        ["Forms", "React Hook Form + Zod validation"],
        ["Charts", "Recharts"],
        ["Icons", "Lucide React"],
        ["Deployment", "Ubuntu 24.04 VPS, systemd services, nginx, Let's Encrypt SSL"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "C.2 Production Deployment", level=2)
    headers = ["Item", "Value"]
    rows = [
        ["Host", "187.124.74.175 (Hostinger, 4 vCPU, 16GB RAM, 193GB disk)"],
        ["Domain", "https://ra.partnerwithus.tech"],
        ["Backend Service", "exzelon-api (4 uvicorn workers, port 8000)"],
        ["Frontend Service", "exzelon-web (Next.js production, port 3000)"],
        ["Reverse Proxy", "nginx (SSL termination, security headers)"],
        ["SSL", "Let's Encrypt (auto-renews via certbot.timer)"],
        ["Database", "MySQL 8.0 (user: ra_user, DB: exzelon_ra_agent)"],
        ["Cache", "Redis 7 (reserved for future use)"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "C.3 System Inventory", level=2)
    headers = ["Category", "Count"]
    rows = [
        ["Dashboard Pages", "20"],
        ["API Endpoint Modules", "32 (140+ routes)"],
        ["Database Tables", "36"],
        ["Services", "70+"],
        ["Adapters", "60+ (13 job, 8 contact, 7 validation, 4 AI, 3 enrichment, 2 CRM, 2 notification)"],
        ["Data Pipelines", "5 (sourcing, enrichment, validation, outreach, warmup)"],
        ["Automation Jobs", "20+ (APScheduler)"],
        ["Backend Tests", "391 (43% coverage)"],
        ["Frontend Tests", "58 (6 suites)"],
        ["E2E Tests", "17 (Playwright smoke tests)"],
    ]
    add_styled_table(doc, headers, rows)


# ---------------------------------------------------------------------------
# Main Generator
# ---------------------------------------------------------------------------
def generate():
    """Generate the complete SOP v1 document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Set paragraph spacing
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Build document
    print("Building SOP v1 document...")

    add_title_page(doc)
    print("  [1/27] Title page")

    add_toc_section(doc)
    print("  [2/27] Table of Contents")

    section_01_overview(doc)
    doc.add_page_break()
    print("  [3/27] System Overview & Login")

    section_02_dashboard(doc)
    doc.add_page_break()
    print("  [4/27] Dashboard")

    section_03_lead_sourcing(doc)
    doc.add_page_break()
    print("  [5/27] Lead Sourcing")

    section_04_contact_enrichment(doc)
    doc.add_page_break()
    print("  [6/27] Contact Enrichment")

    section_05_validation(doc)
    doc.add_page_break()
    print("  [7/27] Email Validation")

    section_06_campaigns(doc)
    doc.add_page_break()
    print("  [8/27] Campaigns")

    section_07_outreach(doc)
    doc.add_page_break()
    print("  [9/27] Outreach")

    section_08_inbox(doc)
    doc.add_page_break()
    print("  [10/27] Unified Inbox")

    section_09_deals(doc)
    doc.add_page_break()
    print("  [11/27] CRM Deal Pipeline")

    section_10_analytics(doc)
    doc.add_page_break()
    print("  [12/27] Analytics")

    section_11_icp_wizard(doc)
    doc.add_page_break()
    print("  [13/27] ICP Wizard")

    section_12_lead_management(doc)
    doc.add_page_break()
    print("  [14/27] Lead Management")

    section_13_clients(doc)
    doc.add_page_break()
    print("  [15/27] Client Management")

    section_14_mailboxes(doc)
    doc.add_page_break()
    print("  [16/27] Mailbox Management")

    section_15_templates(doc)
    doc.add_page_break()
    print("  [17/27] Email Templates")

    section_16_warmup(doc)
    doc.add_page_break()
    print("  [18/27] Warmup Engine")

    section_17_pipelines(doc)
    doc.add_page_break()
    print("  [19/27] Pipeline Orchestration")

    section_18_automation(doc)
    doc.add_page_break()
    print("  [20/27] Automation Control Center")

    section_19_settings(doc)
    doc.add_page_break()
    print("  [21/27] Settings & Configuration")

    section_20_users(doc)
    doc.add_page_break()
    print("  [22/27] User Management")

    section_21_roles(doc)
    doc.add_page_break()
    print("  [23/27] Roles & Permissions")

    section_22_backups(doc)
    doc.add_page_break()
    print("  [24/27] Data Backups")

    section_23_integrations(doc)
    doc.add_page_break()
    print("  [25/27] Integrations & Webhooks")

    section_24_business_rules(doc)
    doc.add_page_break()
    print("  [26/27] Business Rules Reference")

    section_25_glossary(doc)
    doc.add_page_break()
    print("  [27/27] Glossary")

    appendix_a_workflow(doc)
    doc.add_page_break()
    print("  [A] Appendix A: Workflow")

    appendix_b_mapping(doc)
    doc.add_page_break()
    print("  [B] Appendix B: SOP Mapping")

    appendix_c_architecture(doc)
    print("  [C] Appendix C: Architecture")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "--- End of Document ---\n"
        "Exzelon RA Agent | Standard Operating Procedure v1.1 | March 2026 | CONFIDENTIAL\n"
        "System URL: https://ra.partnerwithus.tech"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_GRAY
    run.italic = True

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved to: {OUTPUT_PATH}")
    file_size = os.path.getsize(OUTPUT_PATH)
    print(f"File size: {file_size:,} bytes ({file_size / 1024:.0f} KB)")
    return OUTPUT_PATH


if __name__ == "__main__":
    generate()
