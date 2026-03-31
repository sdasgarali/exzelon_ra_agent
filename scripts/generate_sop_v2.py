"""Generate NeuraLeads AI Agent SOP v2 document (.docx).

Updates from v1:
  - Added Billing & Invoicing module (Section 23, shifted subsequent sections)
  - Updated stats (538 tests, 21 dashboard pages, 37 tables)
  - Updated version to v2.0 / March 2026
  - No multi-tenant terminology anywhere

Run from project root:
    cd backend && python ../scripts/generate_sop_v2.py

Requires: python-docx
Produces: Exzelon_RA_Agent_SOP.docx in parent directory
"""
import os
import sys

# Allow import from same directory
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

from generate_sop_v1 import (
    # Helpers
    set_cell_shading, set_cell_text, add_styled_table, add_callout_box,
    add_screenshot, add_heading, add_bullet_list, add_numbered_list,
    add_toc, add_toc_section,
    # Colors
    COLOR_PRIMARY, COLOR_DARK, COLOR_GREEN, COLOR_RED, COLOR_AMBER,
    COLOR_PURPLE, COLOR_GRAY, COLOR_WHITE, COLOR_HEADER_BG,
    COLOR_GREEN_LIGHT, COLOR_BLUE_LIGHT, COLOR_AMBER_LIGHT,
    COLOR_RED_LIGHT, COLOR_PURPLE_LIGHT, COLOR_GRAY_LIGHT, COLOR_INDIGO_BG,
    # Section functions (reused verbatim)
    section_01_overview, section_02_dashboard, section_03_lead_sourcing,
    section_04_contact_enrichment, section_05_validation,
    section_06_campaigns, section_07_outreach, section_08_inbox,
    section_09_deals, section_10_analytics, section_11_icp_wizard,
    section_12_lead_management, section_13_clients, section_14_mailboxes,
    section_15_templates, section_16_warmup, section_17_pipelines,
    section_18_automation, section_19_settings, section_20_users,
    section_21_roles, section_22_backups,
    section_24_business_rules, section_25_glossary,
    appendix_a_workflow, appendix_b_mapping,
    # Paths
    SCREENSHOT_DIR,
)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Output path — user-specified
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.dirname(PROJECT_ROOT)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "Exzelon_RA_Agent_SOP.docx")


# ---------------------------------------------------------------------------
# Overridden title page (v2 metadata)
# ---------------------------------------------------------------------------
def add_title_page_v2(doc):
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_heading("NeuraLeads AI Agent", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = COLOR_PRIMARY
        run.font.size = Pt(34)

    subtitle = doc.add_heading("Standard Operating Procedure (SOP)", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = COLOR_DARK
        run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Complete End-to-End Operational Guide")
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_GRAY
    run.italic = True

    doc.add_paragraph()

    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("System URL", "https://ra.partnerwithus.tech"),
        ("Version", "2.0 (Updated March 2026)"),
        ("Classification", "CONFIDENTIAL"),
        ("Platform", "FastAPI + Next.js 14 + MySQL"),
        ("Pages", "21 Dashboard Pages | 33 API Modules | 37 Database Tables"),
    ]
    for i, (label, value) in enumerate(meta_data):
        set_cell_text(meta_table.rows[i].cells[0], label, bold=True, size=Pt(10))
        set_cell_text(meta_table.rows[i].cells[1], value, size=Pt(10))
        if i % 2 == 0:
            set_cell_shading(meta_table.rows[i].cells[0], COLOR_GRAY_LIGHT)
            set_cell_shading(meta_table.rows[i].cells[1], COLOR_GRAY_LIGHT)

    doc.add_paragraph()

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("CONFIDENTIAL - INTERNAL USE ONLY")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_RED
    run.bold = True

    doc.add_page_break()


# ---------------------------------------------------------------------------
# NEW: Billing & Invoicing section
# ---------------------------------------------------------------------------
def section_23_billing(doc):
    add_heading(doc, "23. Billing & Invoicing")

    add_callout_box(doc,
        "NEW FEATURE: Complete invoicing and payment management with PDF generation, "
        "Stripe integration, manual payments, automated reminders, and revenue dashboards.",
        style="success"
    )

    add_screenshot(doc, "23_billing.png", "Figure 23.1: Billing Page")

    doc.add_paragraph(
        "The Billing module provides full-lifecycle invoice management: automatic invoice "
        "generation on the 1st of each month, professional PDF invoices, Stripe online "
        "payments, manual payment recording, overdue reminders, and revenue statistics."
    )

    add_heading(doc, "23.1 Invoice Lifecycle", level=2)
    headers = ["Status", "Color", "Meaning"]
    rows = [
        ["Draft", "Gray", "Invoice created but not yet finalized"],
        ["Sent", "Blue", "Invoice sent to client, awaiting payment"],
        ["[GREEN]Paid", "Green", "Payment received and confirmed"],
        ["[RED]Overdue", "Red", "Past due date, reminders being sent"],
        ["[AMBER]Cancelled", "Yellow", "Invoice cancelled before payment"],
        ["Void", "Slate", "Invoice voided after issue"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "23.2 Administrator View (Super Admin)", level=2)
    doc.add_paragraph(
        "Super Admins see the full billing dashboard with four KPI cards, a filterable "
        "invoice list, and a payments history tab."
    )

    add_heading(doc, "23.2.1 Stats Cards", level=3)
    headers = ["Card", "Shows"]
    rows = [
        ["Outstanding", "Total amount of unpaid invoices (sent + overdue)"],
        ["Collected This Month", "Total payments received in the current month"],
        ["Overdue Count", "Number of invoices past their due date"],
        ["MRR", "Monthly Recurring Revenue across all active accounts"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "23.2.2 Invoice Management", level=3)
    headers = ["Action", "Description"]
    rows = [
        ["Bulk Generate", "Generate invoices for multiple accounts at once with period selection"],
        ["Mark as Paid", "Record manual payment: method (manual/bank transfer/check/card), reference, notes"],
        ["Override Amount", "Adjust invoice amount with required reason (draft/sent only)"],
        ["Download PDF", "Download professional PDF invoice with line items and totals"],
        ["Delete", "Soft-delete an invoice (archived, no longer appears in listings)"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "23.2.3 Filtering & Search", level=3)
    add_bullet_list(doc, [
        "Status Filter: Draft, Sent, Paid, Overdue, Cancelled, Void, or All.",
        "Account ID Search: Filter invoices by specific account.",
        "Date Range: Filter by period start/end dates.",
        "Tabs: Switch between Invoices and Payments views.",
    ])

    add_heading(doc, "23.3 User View (Admin / Operator)", level=2)
    doc.add_paragraph(
        "Admin and Operator users see only their own invoices and payment history."
    )
    add_bullet_list(doc, [
        "My Invoices: List of invoices for the user's organization.",
        "Invoice Detail: View line items, amounts, and due dates.",
        "Download PDF: Download own invoice as a professional PDF.",
        "Pay Online: Initiate Stripe checkout for secure online payment (when configured).",
        "Payment History: View all past payments with method, reference, and status.",
    ])

    add_heading(doc, "23.4 Invoice Format (INV-YYYY-NNNN)", level=2)
    doc.add_paragraph(
        "Invoices are numbered sequentially per year: INV-2026-0001, INV-2026-0002, etc. "
        "The number is zero-padded to 4 digits and resets each calendar year."
    )

    add_heading(doc, "23.5 PDF Invoice Layout", level=2)
    doc.add_paragraph(
        "Professional PDF invoices generated via reportlab include:"
    )
    add_bullet_list(doc, [
        "Company header with logo, name, address, phone, email, and website.",
        "Invoice metadata: number, issue date, due date, status.",
        "Bill-to section with client/account details.",
        "Line items table: Description, Quantity, Unit Price, Total.",
        "Subtotal, Tax (calculated automatically), and Total.",
        "Payment terms and footer.",
    ])

    add_heading(doc, "23.6 Payment Methods", level=2)
    headers = ["Method", "Description"]
    rows = [
        ["Stripe", "Online payment via Stripe Checkout (credit/debit card)"],
        ["Manual", "Recorded manually by Super Admin"],
        ["Bank Transfer", "Wire/ACH payment, reference number recorded"],
        ["Check", "Paper check payment, check number recorded"],
        ["Card", "Manual card payment recorded offline"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "23.7 Automated Billing Jobs", level=2)
    headers = ["Job", "Schedule", "Purpose"]
    rows = [
        ["Monthly Invoice Generation", "1st of month at 2:00 AM UTC", "Generate invoices for all active paid accounts"],
        ["Overdue Check", "Daily at 6:00 AM UTC", "Mark sent invoices as overdue if past due date"],
        ["Overdue Reminders", "Daily at 9:00 AM UTC", "Send reminder emails (every 3 days, max 5 reminders)"],
    ]
    add_styled_table(doc, headers, rows)

    add_callout_box(doc,
        "AUTOMATION: All billing jobs respect the global automation master toggle. "
        "Set BILLING_ENABLED=true in environment variables to activate automated billing.",
        style="info"
    )

    add_heading(doc, "23.8 Billing Emails", level=2)
    headers = ["Email", "Trigger", "Includes"]
    rows = [
        ["New Invoice", "When invoice is generated", "Invoice summary + PDF attachment"],
        ["Overdue Reminder", "Every 3 days after due date", "Days overdue count + Pay Now button"],
        ["Payment Acknowledgement", "When payment is recorded", "Payment details + Thank You"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "23.9 Stripe Integration", level=2)
    doc.add_paragraph(
        "When Stripe is configured (STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET, "
        "STRIPE_PUBLISHABLE_KEY), users can pay invoices online via Stripe Checkout. "
        "The webhook endpoint automatically processes checkout.session.completed events "
        "to mark invoices as paid and create payment records."
    )

    add_heading(doc, "23.10 RBAC for Billing", level=2)
    headers = ["Role", "Access"]
    rows = [
        ["Super Admin", "Full: view all invoices, bulk generate, mark paid, override, delete, view stats"],
        ["Admin", "Own: view own invoices, download PDFs, pay online, view payment history"],
        ["Operator", "Own: view own invoices, download PDFs, view payment history"],
        ["Viewer", "No access (403)"],
    ]
    add_styled_table(doc, headers, rows)


# ---------------------------------------------------------------------------
# Updated Integrations section (adds billing webhook)
# ---------------------------------------------------------------------------
def section_23_integrations_v2(doc):
    add_heading(doc, "24. Integrations & Webhooks")

    add_heading(doc, "24.1 Webhook System", level=2)
    doc.add_paragraph(
        "Subscribe to system events with HMAC-SHA256 signed webhook delivery. "
        "Configure target URLs, event filters, and view delivery history with retry tracking."
    )

    headers = ["Event", "Trigger"]
    rows = [
        ["email.sent", "Outreach email delivered"],
        ["email.opened", "Email opened (tracking pixel)"],
        ["email.clicked", "Link clicked (tracking redirect)"],
        ["email.replied", "Reply received"],
        ["email.bounced", "Email bounced"],
        ["contact.unsubscribed", "Contact opted out"],
        ["campaign.completed", "Campaign finished all steps"],
        ["lead.created", "New lead sourced"],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph(
        "Retry policy: Exponential backoff with 3 attempts (1 min, 5 min, 15 min). "
        "Each delivery is signed with X-Webhook-Signature header for verification."
    )

    add_heading(doc, "24.2 CRM Integrations", level=2)
    add_bullet_list(doc, [
        "HubSpot: Push deals, pull contacts, custom property mapping.",
        "Salesforce: OAuth2 flow, Salesforce object sync.",
        "Bidirectional sync with CRM Sync Engine.",
        "Auto-forward interested inbox replies to CRM.",
    ])

    add_heading(doc, "24.3 Zapier/Make Integration", level=2)
    doc.add_paragraph(
        "REST hook subscribe/unsubscribe pattern with sample payloads for Zapier and Make "
        "(Integromat) compatibility. API key authentication via X-API-Key header."
    )

    add_heading(doc, "24.4 Notification Channels", level=2)
    add_bullet_list(doc, [
        "Slack: Webhook-based notifications for campaign events and errors.",
        "Microsoft Teams: Webhook adapter for Teams channels.",
        "Twilio: SMS and calling adapter (stubs for future implementation).",
    ])

    add_heading(doc, "24.5 Stripe Billing Webhook", level=2)
    doc.add_paragraph(
        "Dedicated Stripe webhook endpoint at /billing/webhook/stripe handles "
        "checkout.session.completed events to automatically process online payments, "
        "mark invoices as paid, and send payment acknowledgement emails."
    )


# ---------------------------------------------------------------------------
# Updated glossary (adds billing terms)
# ---------------------------------------------------------------------------
def section_glossary_v2(doc):
    add_heading(doc, "27. Glossary")

    headers = ["Term", "Definition"]
    rows = [
        ["Lead", "A job posting found via job APIs -- represents a hiring opportunity at a company"],
        ["Contact", "A decision-maker at a company (HR manager, recruiter, ops leader, etc.)"],
        ["Client", "A company/organization that has been identified as hiring"],
        ["Pipeline", "An automated multi-step process (sourcing -> enrichment -> validation -> outreach)"],
        ["Campaign", "A multi-step email sequence with A/B testing and conditional branching"],
        ["Mailbox", "A sender email account used to send outreach emails"],
        ["Warmup", "The process of gradually increasing email volume to build domain reputation"],
        ["Bounce", "An email that could not be delivered"],
        ["Catch-All", "A domain that accepts all emails regardless of address existence"],
        ["Dry Run", "A test execution that simulates without actually sending emails"],
        ["Merge Fields", "Template placeholders like {{contact_first_name}} replaced with real data"],
        ["Spintax", "Text variation syntax {option1|option2} for unique email content per recipient"],
        ["P1-P5", "Contact priority levels from highest (Job Poster) to lowest (Functional Manager)"],
        ["ICP", "Ideal Customer Profile -- AI-generated targeting profile"],
        ["Deal", "A CRM sales opportunity tracked through pipeline stages"],
        ["Webhook", "HTTP callback triggered by system events (email.sent, lead.created, etc.)"],
        ["Invoice", "A billing document with line items, amounts, tax, and payment status"],
        ["MRR", "Monthly Recurring Revenue -- total subscription revenue per month"],
        ["KPI", "Key Performance Indicator -- metrics like bounce rate, reply rate, send count"],
        ["RBAC", "Role-Based Access Control -- different users have different permission levels"],
        ["Adapter", "A modular integration component (one per external API/service)"],
        ["Dedup", "Deduplication -- preventing duplicate records from being stored"],
        ["APScheduler", "Background job scheduler that runs automation tasks on a schedule"],
    ]
    add_styled_table(doc, headers, rows, compact=True)


# ---------------------------------------------------------------------------
# Updated Appendix C (updated stats)
# ---------------------------------------------------------------------------
def appendix_c_architecture_v2(doc):
    add_heading(doc, "Appendix C: Technical Architecture")

    add_heading(doc, "C.1 Technology Stack", level=2)
    headers = ["Layer", "Technology"]
    rows = [
        ["Backend", "FastAPI (Python 3.11) + SQLAlchemy 2.0 ORM + Pydantic"],
        ["Frontend", "Next.js 14 (TypeScript) + Tailwind CSS + Radix UI"],
        ["Database", "MySQL 8.0 (production), SQLite (testing)"],
        ["Auth", "JWT tokens (7-day) + API keys (SHA-256) + RBAC (4 roles)"],
        ["Scheduler", "APScheduler (23+ background jobs)"],
        ["State", "Zustand (auth) + TanStack React Query (server data)"],
        ["Forms", "React Hook Form + Zod validation"],
        ["Charts", "Recharts"],
        ["Icons", "Lucide React"],
        ["PDF", "reportlab (invoice PDF generation)"],
        ["Payments", "Stripe SDK (online payments)"],
        ["Deployment", "Ubuntu 24.04 VPS, systemd services, nginx, Let's Encrypt SSL"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "C.2 Production Deployment", level=2)
    headers = ["Item", "Value"]
    rows = [
        ["Host", "187.124.74.175 (Hostinger, 4 vCPU, 16GB RAM, 193GB disk)"],
        ["Domain", "https://ra.partnerwithus.tech"],
        ["Backend Service", "exzelon-api (4 uvicorn workers, port 8000)"],
        ["Frontend Service", "exzelon-web (Next.js production, port 3000)"],
        ["Reverse Proxy", "nginx (SSL termination, security headers)"],
        ["SSL", "Let's Encrypt (auto-renews via certbot.timer)"],
        ["Database", "MySQL 8.0 (user: ra_user, DB: exzelon_ra_agent)"],
        ["Cache", "Redis 7 (reserved for future use)"],
    ]
    add_styled_table(doc, headers, rows)

    add_heading(doc, "C.3 System Inventory", level=2)
    headers = ["Category", "Count"]
    rows = [
        ["Dashboard Pages", "21"],
        ["API Endpoint Modules", "33 (150+ routes)"],
        ["Database Tables", "37"],
        ["Services", "75+"],
        ["Adapters", "60+ (13 job, 8 contact, 7 validation, 4 AI, 3 enrichment, 2 CRM, 2 notification)"],
        ["Data Pipelines", "5 (sourcing, enrichment, validation, outreach, warmup)"],
        ["Automation Jobs", "23+ (APScheduler, includes 3 billing jobs)"],
        ["Backend Tests", "538 (39% coverage)"],
        ["Frontend Tests", "58 (6 suites)"],
        ["E2E Tests", "17 (Playwright smoke tests)"],
    ]
    add_styled_table(doc, headers, rows)


# ---------------------------------------------------------------------------
# Updated Appendix B (adds billing mapping)
# ---------------------------------------------------------------------------
def appendix_b_mapping_v2(doc):
    add_heading(doc, "Appendix B: Original SOP Section Mapping")

    headers = ["Original SOP Section", "System Page(s)"]
    rows = [
        ["Section 1(A): Lead Research & Data Capture", "Dashboard -> Run Lead Sourcing, Leads page, Pipelines"],
        ["Section 1(B): Identify Email-Outreach-Contact", "Dashboard -> Enrich Contacts, Contacts page"],
        ["Section 1(B) Step 2: Identify Decision-Makers", "Contact Enrichment Pipeline (P1-P5 priority)"],
        ["Section 1(B) Step 3: Export Data (CSV)", "Contacts export, Mailmerge Export"],
        ["Section 1(B): Data Quality & Validation", "Dashboard -> Validate Emails, Validation page"],
        ["Section 1(B): Maintain Databases A & B", "Clients page (Regular = Active, Occasional = One-Time)"],
        ["Section 1(C): Email Sending (Mail Merge)", "Campaigns page, Outreach page"],
        ["Section 1(C): MS Word Mail Merge", "Mailmerge Export mode generates CSV for Word"],
        ["What MUST Be Avoided", "Business Rules enforced automatically (Section 26)"],
        ["Section 4: Dashboard Tabs", "Dashboard, Analytics, Leads, Contacts, Outreach pages"],
        ["-- (NEW)", "Campaigns -- multi-step sequences with A/B testing"],
        ["-- (NEW)", "Unified Inbox -- AI sentiment analysis & reply suggestions"],
        ["-- (NEW)", "CRM Deal Pipeline -- Kanban board with auto-advancement"],
        ["-- (NEW)", "Automation Control Center -- 23+ scheduled jobs"],
        ["-- (NEW)", "ICP Wizard -- AI-generated targeting profiles"],
        ["-- (NEW)", "Analytics -- team leaderboard, revenue metrics, cost tracking"],
        ["-- (NEW v2)", "Billing & Invoicing -- PDF invoices, Stripe payments, revenue dashboards"],
    ]
    add_styled_table(doc, headers, rows, compact=True)


# ---------------------------------------------------------------------------
# Updated Appendix A (adds billing to weekly maintenance)
# ---------------------------------------------------------------------------
def appendix_a_workflow_v2(doc):
    add_heading(doc, "Appendix A: Complete Workflow -- Step by Step")

    add_heading(doc, "A.1 Day 1: Initial Setup (One-Time)", level=2)
    add_numbered_list(doc, [
        "Login as Super Admin.",
        "Go to Settings > Job Sources -- configure API keys, target states, industries, job titles.",
        "Go to Settings > Contacts -- configure contact provider API keys.",
        "Go to Settings > Validation -- configure email validation provider API key.",
        "Go to Settings > Outreach -- configure email send mode and SMTP/M365 settings.",
        "Go to Settings > Business Rules -- verify limits (30/day, 10-day cooldown, 4 max contacts).",
        "Go to Settings > Automation -- enable desired automation jobs.",
        "Go to Mailboxes -- add all sender mailboxes, test connections.",
        "Go to Templates -- create outreach email template with merge fields, activate it.",
        "Go to Campaigns -- create multi-step email sequences for automated follow-up.",
        "Go to Users -- create accounts for operators/admins.",
        "Go to Roles & Permissions -- configure access levels for each role.",
        "Go to Billing -- configure billing settings (Stripe keys if using online payments).",
    ])

    add_heading(doc, "A.2 Daily Operation Workflow", level=2)

    doc.add_paragraph("Step 1: Source Leads", style="List Bullet")
    doc.add_paragraph(
        "Dashboard -> 'Run Lead Sourcing' -> Confirm -> Monitor on Pipelines page. "
        "Result: New job postings captured from 10 APIs, companies auto-created. "
        "(Also runs automatically 3x/day if Automation is enabled.)"
    )

    doc.add_paragraph("Step 2: Enrich Contacts", style="List Bullet")
    doc.add_paragraph(
        "Dashboard -> 'Enrich Contacts' -> Select leads -> Run. "
        "Result: Decision-maker contacts discovered with P1-P5 priority levels."
    )

    doc.add_paragraph("Step 3: Validate Emails", style="List Bullet")
    doc.add_paragraph(
        "Dashboard -> 'Validate Emails' -> Select contacts (filter: Pending) -> Run. "
        "Result: Contacts marked as Valid/Invalid/Catch-All/Unknown."
    )

    doc.add_paragraph("Step 4: Send Outreach", style="List Bullet")
    doc.add_paragraph(
        "Option A: Dashboard -> 'Export Mailmerge' for single-send or CSV export. "
        "Option B: Campaigns page -> create/activate campaign for multi-step sequences."
    )

    doc.add_paragraph("Step 5: Monitor Results", style="List Bullet")
    doc.add_paragraph(
        "Dashboard -> Check KPIs, bounce rate, reply rate. "
        "Unified Inbox -> View replies with AI sentiment analysis. "
        "CRM Deals -> Track interested leads through pipeline stages."
    )

    add_heading(doc, "A.3 Weekly Maintenance", level=2)
    add_numbered_list(doc, [
        "Review Dashboard -- Check bounce rate (target < 2%), reply rate, funnel conversion.",
        "Review Warmup -- Check mailbox health scores, DNS records, blacklist status.",
        "Review Analytics -- Compare campaign performance, check cost metrics.",
        "Review Clients -- Check client categories (Regular/Occasional), inactive companies.",
        "Review Automation -- Check event log for any failures or anomalies.",
        "Review Billing -- Check for overdue invoices, verify payments received.",
        "Create Backup -- Go to Backups -> Create Backup (in addition to daily auto-backups).",
    ])

    add_heading(doc, "A.4 Monthly Billing Tasks", level=2)
    add_numbered_list(doc, [
        "Invoices auto-generate on the 1st of each month (if BILLING_ENABLED=true).",
        "Or manually: Billing -> Bulk Generate -> select accounts and period.",
        "Review outstanding invoices -- filter by 'Sent' or 'Overdue' status.",
        "Mark manually-paid invoices as paid with payment method and reference.",
        "Check MRR stats card for revenue trends.",
    ])


# ---------------------------------------------------------------------------
# Main Generator
# ---------------------------------------------------------------------------
def generate():
    """Generate the complete SOP v2 document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Set paragraph spacing
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Build document
    print("Building SOP v2 document...")

    add_title_page_v2(doc)
    print("  [1/30] Title page")

    add_toc_section(doc)
    print("  [2/30] Table of Contents")

    section_01_overview(doc)
    doc.add_page_break()
    print("  [3/30] System Overview & Login")

    section_02_dashboard(doc)
    doc.add_page_break()
    print("  [4/30] Dashboard")

    section_03_lead_sourcing(doc)
    doc.add_page_break()
    print("  [5/30] Lead Sourcing")

    section_04_contact_enrichment(doc)
    doc.add_page_break()
    print("  [6/30] Contact Enrichment")

    section_05_validation(doc)
    doc.add_page_break()
    print("  [7/30] Email Validation")

    section_06_campaigns(doc)
    doc.add_page_break()
    print("  [8/30] Campaigns")

    section_07_outreach(doc)
    doc.add_page_break()
    print("  [9/30] Outreach")

    section_08_inbox(doc)
    doc.add_page_break()
    print("  [10/30] Unified Inbox")

    section_09_deals(doc)
    doc.add_page_break()
    print("  [11/30] CRM Deal Pipeline")

    section_10_analytics(doc)
    doc.add_page_break()
    print("  [12/30] Analytics")

    section_11_icp_wizard(doc)
    doc.add_page_break()
    print("  [13/30] ICP Wizard")

    section_12_lead_management(doc)
    doc.add_page_break()
    print("  [14/30] Lead Management")

    section_13_clients(doc)
    doc.add_page_break()
    print("  [15/30] Client Management")

    section_14_mailboxes(doc)
    doc.add_page_break()
    print("  [16/30] Mailbox Management")

    section_15_templates(doc)
    doc.add_page_break()
    print("  [17/30] Email Templates")

    section_16_warmup(doc)
    doc.add_page_break()
    print("  [18/30] Warmup Engine")

    section_17_pipelines(doc)
    doc.add_page_break()
    print("  [19/30] Pipeline Orchestration")

    section_18_automation(doc)
    doc.add_page_break()
    print("  [20/30] Automation Control Center")

    section_19_settings(doc)
    doc.add_page_break()
    print("  [21/30] Settings & Configuration")

    section_20_users(doc)
    doc.add_page_break()
    print("  [22/30] User Management")

    section_21_roles(doc)
    doc.add_page_break()
    print("  [23/30] Roles & Permissions")

    section_22_backups(doc)
    doc.add_page_break()
    print("  [24/30] Data Backups")

    # NEW: Billing section
    section_23_billing(doc)
    doc.add_page_break()
    print("  [25/30] Billing & Invoicing (NEW)")

    # Integrations (renumbered to 24)
    section_23_integrations_v2(doc)
    doc.add_page_break()
    print("  [26/30] Integrations & Webhooks")

    # Business Rules (renumbered to 25)
    section_24_business_rules(doc)
    doc.add_page_break()
    print("  [27/30] Business Rules Reference")

    # Glossary (renumbered to 26, updated)
    section_glossary_v2(doc)
    doc.add_page_break()
    print("  [28/30] Glossary")

    appendix_a_workflow_v2(doc)
    doc.add_page_break()
    print("  [A] Appendix A: Workflow")

    appendix_b_mapping_v2(doc)
    doc.add_page_break()
    print("  [B] Appendix B: SOP Mapping")

    appendix_c_architecture_v2(doc)
    print("  [C] Appendix C: Architecture")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "--- End of Document ---\n"
        "NeuraLeads AI Agent | Standard Operating Procedure v2.0 | March 2026 | CONFIDENTIAL\n"
        "System URL: https://ra.partnerwithus.tech"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_GRAY
    run.italic = True

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved to: {OUTPUT_PATH}")
    file_size = os.path.getsize(OUTPUT_PATH)
    print(f"File size: {file_size:,} bytes ({file_size / 1024:.0f} KB)")
    return OUTPUT_PATH


if __name__ == "__main__":
    generate()
